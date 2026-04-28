from flask import render_template, redirect, url_for, jsonify, request, make_response, flash, abort, current_app, send_file, session
from flask_login import login_required, current_user, logout_user
from app.main import bp
from app.models import (
    AIUsageEvent,
    ComplianceFrameworkVersion,
    ComplianceRequirement,
    Document,
    DocumentTag,
    Organization,
    OrganizationAISettings,
    OrganizationMembership,
    OrganizationRequirementAssessment,
    RequirementEvidenceLink,
    User,
)
from app import db, mail, limiter, invalidate_org_switcher_context_cache
from app.services.azure_data_service import azure_data_service
from app.services.analytics_service import analytics_service
from app.services.azure_openai_policy_service import azure_openai_policy_service
from app.services.policy_draft_service import policy_draft_service
from app.services.compliance_scoring_service import compliance_scoring_service
from app.services.notification_service import notification_service
from app.services.rag_query_service import rag_query_service
from app.services.document_analysis_service import document_analysis_service
from app.services.billing_service import billing_service
from sqlalchemy import and_, or_, func

import threading
import time
from datetime import datetime, timezone, timedelta

from flask_mail import Message
from itsdangerous import URLSafeTimedSerializer
from werkzeug.exceptions import NotFound

import os
import hashlib
import json
import csv
import io
import zipfile


_RESEND_ORG_INVITE_COOLDOWN_SECONDS = 60 * 5

_ORG_INVITE_TOKEN_SALT = 'org-invite'


_ORG_LOGO_CACHE: dict[tuple[int, str], tuple[float, bytes, str | None]] = {}
_ORG_LOGO_CACHE_LOCK = threading.Lock()
_DOCUMENTS_SEARCH_TEXT_AVAILABLE: bool | None = None


def _safe_int_env(name: str, default: int) -> int:
    try:
        return int(os.environ.get(name) or default)
    except Exception:
        return default


def _clamp_int(value, *, default: int, minimum: int, maximum: int) -> int:
    try:
        parsed = int(value)
    except Exception:
        parsed = default
    return max(minimum, min(maximum, parsed))


def _limit_text(text: str, *, max_chars: int) -> str:
    value = (text or '').strip()
    if max_chars <= 0:
        return value
    return value[:max_chars]


def _log_ai_call(event_name: str, *, org_id: int, mode: str, provider: str, model: str, usage: dict | None, latency_ms: int):
    usage = usage or {}
    current_app.logger.info(
        '[AI_USAGE] %s',
        json.dumps(
            {
                'event': event_name,
                'org_id': int(org_id),
                'mode': mode,
                'provider': provider,
                'model': model,
                'prompt_tokens': int(usage.get('prompt_tokens') or 0),
                'completion_tokens': int(usage.get('completion_tokens') or 0),
                'total_tokens': int(usage.get('total_tokens') or 0),
                'latency_ms': int(latency_ms),
            },
            sort_keys=True,
        ),
    )

    try:
        user_id = int(current_user.id) if getattr(current_user, 'is_authenticated', False) else None
        db.session.add(
            AIUsageEvent(
                organization_id=int(org_id),
                user_id=user_id,
                event=(event_name or '').strip() or 'unknown',
                mode=(mode or '').strip() or 'unknown',
                provider=(provider or '').strip() or 'unknown',
                model=(model or '').strip() or 'unknown',
                prompt_tokens=int(usage.get('prompt_tokens') or 0),
                completion_tokens=int(usage.get('completion_tokens') or 0),
                total_tokens=int(usage.get('total_tokens') or 0),
                latency_ms=int(latency_ms or 0),
            )
        )
        db.session.commit()
    except Exception:
        db.session.rollback()
        current_app.logger.exception('Failed to persist AI usage event')


def _ai_rate_limit_key() -> str:
    try:
        user_id = int(current_user.id) if getattr(current_user, 'is_authenticated', False) else 0
    except Exception:
        user_id = 0
    try:
        org_id = int(_active_org_id() or 0)
    except Exception:
        org_id = 0

    ip = (request.headers.get('X-Forwarded-For') or request.remote_addr or 'unknown').strip()
    return f'org:{org_id}:user:{user_id}:ip:{ip}'


def _get_org_ai_settings(org_id: int | None) -> OrganizationAISettings | None:
    try:
        if not org_id:
            return None
        return OrganizationAISettings.query.filter_by(organization_id=int(org_id)).first()
    except Exception:
        return None


def _effective_ai_setting(org_id: int | None, key: str, default):
    settings = _get_org_ai_settings(org_id)
    if settings is None:
        return default
    value = getattr(settings, key, None)
    return default if value is None else value


def _rag_rate_limit_value() -> str:
    org_id = _active_org_id()
    default = current_app.config.get('AI_RAG_RATE_LIMIT') or '20 per minute'
    return str(_effective_ai_setting(org_id, 'rag_rate_limit', default) or default)


def _policy_rate_limit_value() -> str:
    org_id = _active_org_id()
    default = current_app.config.get('AI_POLICY_RATE_LIMIT') or '10 per minute'
    return str(_effective_ai_setting(org_id, 'policy_rate_limit', default) or default)


def _org_invite_token_ttl_seconds() -> int:
    # Keep in sync with auth.routes._org_invite_token_ttl_seconds
    return max(60, _safe_int_env('ORG_INVITE_TOKEN_TTL_SECONDS', 60 * 60 * 24))


def _format_duration_seconds(seconds: int) -> str:
    seconds = int(seconds or 0)
    if seconds <= 0:
        return 'a short time'
    if seconds % (60 * 60 * 24) == 0:
        days = seconds // (60 * 60 * 24)
        return f'{days} day' if days == 1 else f'{days} days'
    if seconds % (60 * 60) == 0:
        hours = seconds // (60 * 60)
        return f'{hours} hour' if hours == 1 else f'{hours} hours'
    minutes = max(1, seconds // 60)
    return f'{minutes} minute' if minutes == 1 else f'{minutes} minutes'


def _etag_matches_if_none_match(if_none_match: str | None, etag: str) -> bool:
    if not if_none_match:
        return False
    value = if_none_match.strip()
    if value == '*':
        return True
    candidates = [part.strip() for part in value.split(',') if part.strip()]
    strong_etag = etag[2:] if etag.startswith('W/') else etag
    return (etag in candidates) or (strong_etag in candidates)


def _documents_search_text_available() -> bool:
    global _DOCUMENTS_SEARCH_TEXT_AVAILABLE
    if _DOCUMENTS_SEARCH_TEXT_AVAILABLE is not None:
        return bool(_DOCUMENTS_SEARCH_TEXT_AVAILABLE)
    try:
        inspector = db.inspect(db.engine)
        cols = [c.get('name') for c in inspector.get_columns('documents')]
        _DOCUMENTS_SEARCH_TEXT_AVAILABLE = 'search_text' in cols
    except Exception:
        _DOCUMENTS_SEARCH_TEXT_AVAILABLE = False
    return bool(_DOCUMENTS_SEARCH_TEXT_AVAILABLE)


def _get_cached_org_logo(org_id: int, blob_name: str) -> tuple[bytes, str | None] | None:
    now = time.monotonic()
    with _ORG_LOGO_CACHE_LOCK:
        cached = _ORG_LOGO_CACHE.get((org_id, blob_name))
        if not cached:
            return None
        expires_at, data, content_type = cached
        if now >= expires_at:
            try:
                del _ORG_LOGO_CACHE[(org_id, blob_name)]
            except KeyError:
                pass
            return None
        return data, content_type


def _set_cached_org_logo(org_id: int, blob_name: str, data: bytes, content_type: str | None, ttl_seconds: int) -> None:
    if ttl_seconds <= 0:
        return
    expires_at = time.monotonic() + ttl_seconds
    with _ORG_LOGO_CACHE_LOCK:
        _ORG_LOGO_CACHE[(org_id, blob_name)] = (expires_at, data, content_type)


def _org_logo_disk_cache_paths(org_id: int, blob_name: str) -> tuple[str, str]:
    digest = hashlib.sha256(blob_name.encode('utf-8')).hexdigest()
    base_dir = os.path.join(current_app.instance_path, 'cache', 'org_logos')
    return (
        os.path.join(base_dir, f'{org_id}_{digest}.bin'),
        os.path.join(base_dir, f'{org_id}_{digest}.json'),
    )


def _get_disk_cached_org_logo(org_id: int, blob_name: str) -> tuple[bytes, str | None] | None:
    try:
        ttl_seconds = int(current_app.config.get('ORG_LOGO_DISK_CACHE_SECONDS') or 86400)
    except Exception:
        ttl_seconds = 86400
    if ttl_seconds <= 0:
        return None

    data_path, meta_path = _org_logo_disk_cache_paths(org_id, blob_name)
    try:
        if not os.path.exists(data_path) or not os.path.exists(meta_path):
            return None
        with open(meta_path, 'r', encoding='utf-8') as f:
            meta = json.load(f) or {}
        created_at = float(meta.get('created_at') or 0)
        if (time.time() - created_at) > ttl_seconds:
            return None
        with open(data_path, 'rb') as f:
            data = f.read()
        if not data:
            return None
        return data, (meta.get('content_type') or None)
    except Exception:
        return None


def _set_disk_cached_org_logo(org_id: int, blob_name: str, data: bytes, content_type: str | None) -> None:
    try:
        ttl_seconds = int(current_app.config.get('ORG_LOGO_DISK_CACHE_SECONDS') or 86400)
    except Exception:
        ttl_seconds = 86400
    if ttl_seconds <= 0:
        return
    try:
        data_path, meta_path = _org_logo_disk_cache_paths(org_id, blob_name)
        os.makedirs(os.path.dirname(data_path), exist_ok=True)
        tmp_data = data_path + '.tmp'
        tmp_meta = meta_path + '.tmp'
        with open(tmp_data, 'wb') as f:
            f.write(data)
        with open(tmp_meta, 'w', encoding='utf-8') as f:
            json.dump({'created_at': time.time(), 'content_type': content_type}, f)
        os.replace(tmp_data, data_path)
        os.replace(tmp_meta, meta_path)
    except Exception:
        return


@bp.route('/terms')
def terms():
    return render_template('legal/terms.html', title='Terms and Conditions')


@bp.route('/privacy')
def privacy():
    return render_template('legal/privacy.html', title='Privacy Policy')


@bp.route('/disclaimer')
def disclaimer():
    return render_template('legal/disclaimer.html', title='Disclaimer')


def _active_org_id() -> int | None:
    org_id = getattr(current_user, 'organization_id', None)
    return int(org_id) if org_id else None


def _require_active_org():
    org_id = _active_org_id()
    if not org_id:
        flash('Please select an organisation to continue.', 'info')
        return redirect(url_for('onboarding.organization'))

    membership = (
        OrganizationMembership.query
        .filter_by(user_id=int(current_user.id), organization_id=int(org_id), is_active=True)
        .first()
    )
    if not membership:
        flash('You do not have access to that organisation.', 'error')
        return redirect(url_for('onboarding.organization'))
    return None


def _require_org_admin():
    maybe = _require_active_org()
    if maybe is not None:
        return maybe
    if not current_user.has_permission('users.manage', org_id=_active_org_id()):
        abort(403)
    return None


def _require_org_permission(permission_code: str):
    maybe = _require_active_org()
    if maybe is not None:
        return maybe
    if not current_user.has_permission(permission_code, org_id=_active_org_id()):
        abort(403)
    return None


def _plan_enforcement_enabled() -> bool:
    # Keep tests deterministic and avoid forcing test fixtures to set plan metadata.
    if bool(current_app.config.get('TESTING')):
        return False
    return True


def _is_super_admin_user() -> bool:
    return bool(billing_service.is_super_admin_email(getattr(current_user, 'email', None)))


def _require_plan_feature(feature_key: str, *, fallback_endpoint: str = 'main.plans_preview'):
    if not _plan_enforcement_enabled():
        return None

    maybe = _require_active_org()
    if maybe is not None:
        return maybe

    org_id = _active_org_id()
    if not org_id:
        return redirect(url_for(fallback_endpoint))

    organization = db.session.get(Organization, int(org_id))
    if not organization:
        abort(404)

    billing_state = billing_service.resolve_entitlements(
        organization,
        actor_email=getattr(current_user, 'email', None),
    )
    feature_access = dict(billing_state.get('feature_access') or {})
    allowed = bool(feature_access.get((feature_key or '').strip().lower()))
    if allowed:
        return None

    min_plan = billing_service.FEATURE_MIN_PLAN.get((feature_key or '').strip().lower(), 'higher')
    flash(
        f"This feature requires the {str(min_plan).capitalize()} plan or above. "
        'Use Plans to switch the test plan and try again.',
        'warning',
    )
    return redirect(url_for(fallback_endpoint))


def _membership_has_permission(membership: OrganizationMembership, code: str) -> bool:
    if not membership or not membership.is_active:
        return False

    if membership.rbac_role:
        try:
            return code in membership.rbac_role.effective_permission_codes()
        except Exception:
            return False

    # Legacy fallback: only supports basic admin mapping.
    if code == 'users.manage':
        return (membership.role or '').strip().lower() in {
            'admin',
            'organisation administrator',
            'organization administrator',
        }

    return False


def _serializer() -> URLSafeTimedSerializer:
    return URLSafeTimedSerializer(current_app.config['SECRET_KEY'])


def _mail_configured() -> bool:
    return bool(current_app.config.get('MAIL_SERVER') and current_app.config.get('MAIL_DEFAULT_SENDER'))


def _password_reset_token(user: User) -> str:
    # Must match the implementation in auth/routes.py
    return _serializer().dumps({'user_id': user.id, 'email': user.email}, salt='password-reset')


def _org_invite_token(user: User) -> str:
    # Must match the implementation in auth/routes.py
    return _serializer().dumps({'user_id': user.id, 'email': user.email}, salt=_ORG_INVITE_TOKEN_SALT)


def _send_invite_email(user: User, reset_url: str, organization: Organization) -> None:
    if not _mail_configured():
        current_app.logger.warning('MAIL not configured; invite reset URL: %s', reset_url)
        return

    subject = f"You're invited to {organization.name}"
    expiry_duration = _format_duration_seconds(_org_invite_token_ttl_seconds())
    
    # Plain text version
    body = (
        f"You've been invited to join {organization.name} on Cenaris.\n\n"
        f"Set your password here: {reset_url}\n\n"
        f"This link expires in {expiry_duration}.\n\n"
        "If you weren't expecting this invite, you can ignore this email."
    )
    
    # HTML version with template
    try:
        from flask import render_template
        html = render_template('email/invite.html', user=user, reset_url=reset_url, organization=organization, expiry_duration=expiry_duration)
    except Exception:
        html = None
    
    try:
        from app.auth.routes import _send_email, _send_email_html
        if html:
            _send_email_html(user.email, subject, body, html)
        else:
            _send_email(user.email, subject, body)
    except Exception:
        current_app.logger.exception('Failed to send invite email to %s (org_id=%s)', user.email, getattr(organization, 'id', None))
        raise


def _is_pending_org_invite(membership: OrganizationMembership, user: User) -> bool:
    # In this app, org "invites" create an inactive-password user and an org membership.
    # A "pending invite" is specifically a membership that was invited (invited_at set),
    # has not been accepted yet, and the user still has no password set.
    # OAuth users may not have a password_hash, so we must not treat them as pending unless
    # the membership is actually invite-tracked.
    return bool(
        membership
        and user
        and membership.invited_at is not None
        and membership.invite_accepted_at is None
        and membership.invite_revoked_at is None
        and not bool(user.password_hash)
    )


def _update_organization_logo(organization: Organization, logo_file) -> tuple[bool, str]:
    """
    Unified logo upload handler for both onboarding and settings.
    Deletes old logo before uploading new one to prevent orphaned files.
    
    Args:
        organization: Organization object to update
        logo_file: FileStorage object from form
        
    Returns:
        Tuple of (success: bool, message: str)
    """
    import uuid
    from app.services.azure_storage_service import azure_storage_service
    
    if not logo_file or not getattr(logo_file, 'filename', ''):
        return False, 'No logo file selected'
    
    # Validate file extension
    ext = (logo_file.filename.rsplit('.', 1)[-1] or '').lower()
    safe_ext = ext if ext in {'png', 'jpg', 'jpeg', 'webp'} else 'png'
    
    # Generate new blob name
    unique = uuid.uuid4().hex
    new_blob_name = f"organizations/{organization.id}/branding/logo_{unique}.{safe_ext}"
    content_type = getattr(logo_file, 'mimetype', None)
    
    # Delete old logo if exists (prevents orphaned files)
    if organization.logo_blob_name:
        try:
            old_blob = organization.logo_blob_name
            # Remove org prefix if it's already in the blob name
            if old_blob.startswith('org_'):
                old_blob = old_blob[len(f'org_{organization.id}/'):]
            azure_storage_service.delete_blob(old_blob, organization_id=int(organization.id))
            current_app.logger.info(f'Deleted old logo: {organization.logo_blob_name}')
        except Exception as e:
            current_app.logger.warning(f'Could not delete old logo {organization.logo_blob_name}: {e}')
            # Continue anyway - old logo deletion failure shouldn't block new upload
    
    # Upload new logo
    data = logo_file.read()
    if not azure_storage_service.upload_blob(new_blob_name, data, content_type=content_type, organization_id=int(organization.id)):
        return False, 'Logo upload failed. Check Azure Storage configuration.'
    
    # Update organization record
    organization.logo_blob_name = new_blob_name
    organization.logo_content_type = content_type
    
    return True, 'Logo uploaded successfully'


@bp.route('/org/switch', methods=['POST'])
@login_required
def switch_organization():
    """Switch the active organization for the current user."""
    org_id_raw = (request.form.get('organization_id') or '').strip()
    if not org_id_raw.isdigit():
        flash('Invalid organisation.', 'error')
        return redirect(url_for('main.dashboard'))

    org_id = int(org_id_raw)
    membership = (
        OrganizationMembership.query
        .filter_by(user_id=int(current_user.id), organization_id=org_id, is_active=True)
        .first()
    )
    if not membership:
        flash('You do not have access to that organization.', 'error')
        return redirect(url_for('main.dashboard'))

    # Query the actual user object from the database to ensure changes persist
    user = db.session.get(User, int(current_user.id))
    user.organization_id = org_id
    db.session.commit()
    flash('Organisation switched.', 'success')
    return redirect(request.referrer or url_for('main.dashboard'))


@bp.route('/org/admin')
@login_required
def org_admin_dashboard():
    """Organisation admin overview."""
    maybe = _require_org_admin()
    if maybe is not None:
        return maybe

    from app.main.forms import (
        InviteMemberForm,
        MembershipActionForm,
        PendingInviteResendForm,
        PendingInviteRevokeForm,
        UpdateMemberRoleForm,
        UpdateMemberDepartmentForm,
    )
    from app.models import Department

    org_id = _active_org_id()
    organization = db.session.get(Organization, int(org_id))
    if not organization:
        abort(404)

    # Ensure RBAC defaults exist so UI can show role names reliably.
    try:
        from app.services.rbac import ensure_rbac_seeded_for_org

        ensure_rbac_seeded_for_org(int(org_id))
        db.session.commit()
    except Exception:
        db.session.rollback()

    members = (
        OrganizationMembership.query
        .filter_by(organization_id=int(org_id))
        .join(User, User.id == OrganizationMembership.user_id)
        .order_by(OrganizationMembership.is_active.desc(), User.email.asc())
        .all()
    )

    pending_invites = [m for m in members if _is_pending_org_invite(m, m.user)]

    # Used to determine whether the current user (if admin) can remove their own membership.
    def _can_manage_users(m: OrganizationMembership) -> bool:
        return _membership_has_permission(m, 'users.manage')

    active_admin_count = sum(1 for m in members if _can_manage_users(m))
    current_membership = next((m for m in members if int(m.user_id) == int(current_user.id)), None)
    current_is_active_admin = bool(current_membership and _can_manage_users(current_membership))
    can_current_user_leave_org = (not current_is_active_admin) or (active_admin_count > 1)

    user_count = sum(1 for m in members if bool(m.is_active))
    document_count = Document.query.filter_by(organization_id=int(org_id), is_active=True).count()

    invite_form = InviteMemberForm()
    try:
        from app.models import RBACRole

        roles = (
            RBACRole.query
            .filter_by(organization_id=int(org_id))
            .order_by(RBACRole.name.asc())
            .all()
        )
        invite_form.role.choices = [(str(r.id), r.name) for r in roles]
    except Exception:
        invite_form.role.choices = []
    departments = (
        Department.query
        .filter_by(organization_id=int(org_id))
        .order_by(Department.name.asc())
        .all()
    )
    invite_form.department_id.choices = [('', 'Select department')] + [
        (str(d.id), d.name) for d in departments
    ]
    member_action_form = MembershipActionForm()
    update_role_form = UpdateMemberRoleForm()
    update_department_form = UpdateMemberDepartmentForm()
    pending_invite_resend_form = PendingInviteResendForm()
    pending_invite_revoke_form = PendingInviteRevokeForm()

    # Populate role choices for role-update form.
    available_roles = []
    try:
        from app.models import RBACRole

        roles = (
            RBACRole.query
            .filter_by(organization_id=int(org_id))
            .order_by(RBACRole.name.asc())
            .all()
        )
        available_roles = roles
        update_role_form.role_id.choices = [(str(r.id), r.name) for r in roles]
    except Exception:
        update_role_form.role_id.choices = []

    return render_template(
        'main/org_admin_dashboard.html',
        title='Team Management',
        organization=organization,
        members=members,
        pending_invites=pending_invites,
        active_admin_count=active_admin_count,
        can_current_user_leave_org=can_current_user_leave_org,
        user_count=user_count,
        document_count=document_count,
        invite_expires_in=_format_duration_seconds(_org_invite_token_ttl_seconds()),
        invite_form=invite_form,
        member_action_form=member_action_form,
        update_role_form=update_role_form,
        update_department_form=update_department_form,
        pending_invite_resend_form=pending_invite_resend_form,
        pending_invite_revoke_form=pending_invite_revoke_form,
        departments=departments,
        available_roles=available_roles,
    )


@bp.route('/org/admin/compliance/initialize', methods=['POST'])
@login_required
def org_admin_initialize_compliance_data():
    """Initialize NDIS mapping data for the active organisation."""
    maybe = _require_org_admin()
    if maybe is not None:
        return maybe

    from app.main.forms import InitializeComplianceDataForm
    from app.services.compliance_mapping_service import ComplianceMappingImportError, compliance_mapping_service

    form = InitializeComplianceDataForm()
    if not form.validate_on_submit():
        flash('Unable to initialize compliance data. Please refresh and try again.', 'error')
        return redirect(url_for('main.org_admin_dashboard'))

    org_id = _active_org_id()
    mapping_dir = os.path.abspath(
        os.path.join(
            current_app.root_path,
            os.pardir,
            'data',
            'sources',
            'ndis',
            'mapping',
        )
    )
    mapping_csv_path = os.path.join(mapping_dir, 'MASTER Cenaris_NDIS_Audit_Master_Mapping_v1.csv')
    mapping_xlsx_path = os.path.join(mapping_dir, 'MASTER Cenaris_NDIS_Audit_Master_Mapping_v1.xlsx')
    mapping_file_path = mapping_csv_path if os.path.exists(mapping_csv_path) else mapping_xlsx_path

    if not os.path.exists(mapping_file_path):
        flash('NDIS mapping file is missing. Please upload or restore it in data/sources/ndis/mapping.', 'error')
        return redirect(url_for('main.org_admin_dashboard'))

    try:
        result = compliance_mapping_service.import_master_mapping(
            mapping_file_path,
            organization_id=int(org_id),
            imported_by_user_id=int(current_user.id),
            version_label='v1.0',
        )
        flash(
            f'NDIS mapping initialized. Loaded {result.imported_requirements} requirements for this organisation.',
            'success',
        )
    except ComplianceMappingImportError as e:
        flash(f'Initialization failed: {e}', 'error')
    except Exception as e:
        db.session.rollback()
        current_app.logger.exception('Failed to initialize NDIS mapping for org %s', org_id)
        flash(f'Initialization failed: {e}', 'error')

    return redirect(url_for('main.org_admin_dashboard'))


@bp.route('/org/admin/members/department', methods=['POST'])
@login_required
def org_admin_update_member_department():
    """Update a member's department assignment."""
    maybe = _require_org_permission('users.manage')
    if maybe is not None:
        return maybe

    from flask import request, jsonify
    from app.main.forms import UpdateMemberDepartmentForm
    from app.models import Department

    def _wants_json() -> bool:
        return (request.headers.get('X-Requested-With') == 'fetch') or (request.accept_mimetypes.best == 'application/json')

    org_id = _active_org_id()
    organization = db.session.get(Organization, int(org_id))
    if not organization:
        if _wants_json():
            return jsonify(success=False, error='Organisation not found.'), 404
        flash('Organisation not found.', 'error')
        return redirect(url_for('main.org_admin_dashboard'))

    form = UpdateMemberDepartmentForm()

    # Populate choices so WTForms validates the selection.
    departments = (
        Department.query
        .filter_by(organization_id=int(org_id))
        .order_by(Department.name.asc())
        .all()
    )
    form.department_id.choices = [('', 'Unassigned')] + [(str(d.id), d.name) for d in departments]

    if not form.validate_on_submit():
        if _wants_json():
            return jsonify(success=False, error='Invalid request.'), 400
        flash('Invalid request.', 'error')
        return redirect(url_for('main.org_admin_dashboard'))

    membership_id_raw = (form.membership_id.data or '').strip()
    dept_id_raw = (form.department_id.data or '').strip()

    if not membership_id_raw.isdigit():
        if _wants_json():
            return jsonify(success=False, error='Invalid request.'), 400
        flash('Invalid request.', 'error')
        return redirect(url_for('main.org_admin_dashboard'))

    membership = db.session.get(OrganizationMembership, int(membership_id_raw))
    if not membership or int(membership.organization_id) != int(org_id):
        if _wants_json():
            return jsonify(success=False, error='Membership not found.'), 404
        flash('Membership not found.', 'error')
        return redirect(url_for('main.org_admin_dashboard'))

    new_dept = None
    if dept_id_raw:
        if not dept_id_raw.isdigit():
            if _wants_json():
                return jsonify(success=False, error='Invalid department.'), 400
            flash('Invalid department.', 'error')
            return redirect(url_for('main.org_admin_dashboard'))

        new_dept = db.session.get(Department, int(dept_id_raw))
        if not new_dept or int(new_dept.organization_id) != int(org_id):
            if _wants_json():
                return jsonify(success=False, error='Department not found.'), 404
            flash('Department not found.', 'error')
            return redirect(url_for('main.org_admin_dashboard'))

    try:
        membership.department_id = int(new_dept.id) if new_dept else None
        db.session.commit()

        if _wants_json():
            return jsonify(
                success=True,
                membership_id=int(membership.id),
                department={
                    'id': int(new_dept.id),
                    'name': new_dept.name,
                    'color': new_dept.color,
                } if new_dept else None,
            )

        flash('Department updated.', 'success')
    except Exception:
        db.session.rollback()
        current_app.logger.exception('Failed updating member department')
        if _wants_json():
            return jsonify(success=False, error='Failed to update department. Please try again.'), 500
        flash('Failed to update department. Please try again.', 'error')

    return redirect(url_for('main.org_admin_dashboard'))


@bp.route('/org/admin/members/role', methods=['POST'])
@login_required
def org_admin_update_member_role():
    """Update a member's org-scoped RBAC role."""
    maybe = _require_org_permission('roles.manage')
    if maybe is not None:
        return maybe

    from flask import request, jsonify

    def _wants_json() -> bool:
        return (request.headers.get('X-Requested-With') == 'fetch') or (request.accept_mimetypes.best == 'application/json')

    from app.main.forms import UpdateMemberRoleForm
    from app.models import RBACRole

    org_id = _active_org_id()
    organization = db.session.get(Organization, int(org_id))
    if not organization:
        if _wants_json():
            return jsonify(success=False, error='Organisation not found.'), 404
        flash('Organisation not found.', 'error')
        return redirect(url_for('main.org_admin_dashboard'))

    form = UpdateMemberRoleForm()

    # Populate role choices so WTForms validates the selection.
    try:
        roles = (
            RBACRole.query
            .filter_by(organization_id=int(org_id))
            .order_by(RBACRole.name.asc())
            .all()
        )
        form.role_id.choices = [(str(r.id), r.name) for r in roles]
    except Exception:
        form.role_id.choices = []

    if not form.validate_on_submit():
        if _wants_json():
            return jsonify(success=False, error='Invalid request.'), 400
        flash('Invalid request.', 'error')
        return redirect(url_for('main.org_admin_dashboard'))

    membership_id_raw = (form.membership_id.data or '').strip()
    role_id_raw = (form.role_id.data or '').strip()

    if not membership_id_raw.isdigit() or not role_id_raw.isdigit():
        if _wants_json():
            return jsonify(success=False, error='Invalid request.'), 400
        flash('Invalid request.', 'error')
        return redirect(url_for('main.org_admin_dashboard'))

    membership = db.session.get(OrganizationMembership, int(membership_id_raw))
    if not membership or int(membership.organization_id) != int(org_id):
        if _wants_json():
            return jsonify(success=False, error='Membership not found.'), 404
        flash('Membership not found.', 'error')
        return redirect(url_for('main.org_admin_dashboard'))

    target_role = db.session.get(RBACRole, int(role_id_raw))
    if not target_role or int(target_role.organization_id) != int(org_id):
        if _wants_json():
            return jsonify(success=False, error='Role not found.'), 404
        flash('Role not found.', 'error')
        return redirect(url_for('main.org_admin_dashboard'))

    currently_admin = _membership_has_permission(membership, 'users.manage')
    try:
        new_admin = 'users.manage' in target_role.effective_permission_codes()
    except Exception:
        new_admin = False

    # Guard: never demote the last active admin.
    if membership.is_active and currently_admin and not new_admin:
        active_memberships = (
            OrganizationMembership.query
            .filter_by(organization_id=int(org_id), is_active=True)
            .all()
        )
        active_admins = sum(1 for m in active_memberships if _membership_has_permission(m, 'users.manage'))
        if active_admins <= 1:
            if _wants_json():
                return jsonify(success=False, error='Cannot change role: you would remove the last admin.'), 400
            flash('Cannot change role: you would remove the last admin.', 'error')
            return redirect(url_for('main.org_admin_dashboard'))

    try:
        membership.role_id = int(target_role.id)
        # Keep legacy string role in sync during transition.
        membership.role = 'Admin' if new_admin else 'User'
        db.session.commit()

        # Invalidate cached navigation context (role badge/permissions) so the
        # change is reflected immediately for the affected user.
        try:
            from app import invalidate_org_switcher_context_cache
            invalidate_org_switcher_context_cache(membership.user_id, membership.organization_id)
            invalidate_org_switcher_context_cache(current_user.id, membership.organization_id)
        except Exception:
            pass
        
        # Force SQLAlchemy to reload the rbac_role relationship from database
        # This ensures display_role_name shows the correct new role
        db.session.expire(membership, ['rbac_role'])
        db.session.refresh(membership)
        role_name = target_role.name
        if _wants_json():
            return jsonify(
                success=True,
                membership_id=int(membership.id),
                user_id=int(membership.user_id),
                new_role_name=role_name,
                is_current_user=(int(membership.user_id) == int(current_user.id)),
            )

        flash('Role updated.', 'success')
    except Exception:
        db.session.rollback()
        current_app.logger.exception('Failed updating member role')
        if _wants_json():
            return jsonify(success=False, error='Failed to update role. Please try again.'), 500

        flash('Failed to update role. Please try again.', 'error')

    return redirect(url_for('main.org_admin_dashboard'))


@bp.route('/org/admin/invite', methods=['POST'])
@login_required
def org_admin_invite_member():
    """Invite/add a user to the active organization by email."""
    maybe = _require_org_permission('users.invite')
    if maybe is not None:
        return maybe

    from app.main.forms import InviteMemberForm
    from datetime import datetime, timezone
    from sqlalchemy import func
    from app.models import Department

    org_id = _active_org_id()
    organization = db.session.get(Organization, int(org_id))
    if not organization:
        abort(404)

    form = InviteMemberForm()

    # Seed RBAC so role selection works.
    try:
        from app.services.rbac import ensure_rbac_seeded_for_org

        ensure_rbac_seeded_for_org(int(org_id))
        db.session.flush()
    except Exception:
        db.session.rollback()
    # Populate department choices (so WTForms validates select value).
    departments = (
        Department.query
        .filter_by(organization_id=int(org_id))
        .order_by(Department.name.asc())
        .all()
    )
    form.department_id.choices = [('', 'Select department')] + [(str(d.id), d.name) for d in departments]

    # Populate role choices from RBAC roles.
    try:
        from app.models import RBACRole

        roles = (
            RBACRole.query
            .filter_by(organization_id=int(org_id))
            .order_by(RBACRole.name.asc())
            .all()
        )
        form.role.choices = [(str(r.id), r.name) for r in roles]
    except Exception:
        form.role.choices = []
    if not form.validate_on_submit():
        if getattr(form, 'department_id', None) is not None and getattr(form.department_id, 'errors', None):
            flash(form.department_id.errors[0], 'error')
        else:
            flash('Please correct the invite form errors and try again.', 'error')
        return redirect(url_for('main.org_admin_dashboard'))

    email = (form.email.data or '').strip().lower()

    role_id = None
    selected_role = None
    role_raw = (form.role.data or '').strip()
    if role_raw.isdigit():
        role_id = int(role_raw)
        try:
            from app.models import RBACRole

            selected_role = db.session.get(RBACRole, int(role_id))
            if not selected_role or int(selected_role.organization_id) != int(org_id):
                selected_role = None
        except Exception:
            selected_role = None

    if not selected_role:
        try:
            from app.models import RBACRole
            from app.services.rbac import BUILTIN_ROLE_KEYS

            selected_role = (
                RBACRole.query
                .filter_by(organization_id=int(org_id), name=BUILTIN_ROLE_KEYS.MEMBER)
                .first()
            )
        except Exception:
            selected_role = None

    selected_role_id = int(selected_role.id) if selected_role else None

    # Department: either select existing OR create new.
    department = None
    new_dept_name = (form.new_department_name.data or '').strip()
    new_dept_color = (form.new_department_color.data or 'primary').strip() or 'primary'
    allowed_colors = {'primary', 'secondary', 'success', 'info', 'warning', 'danger', 'dark'}
    if new_dept_color not in allowed_colors:
        new_dept_color = 'primary'

    if new_dept_name:
        # Try to find existing (case-insensitive) department in this org.
        department = (
            Department.query
            .filter(Department.organization_id == int(org_id))
            .filter(func.lower(Department.name) == func.lower(new_dept_name))
            .first()
        )
        if not department:
            department = Department(
                organization_id=int(org_id),
                name=new_dept_name,
                color=new_dept_color,
            )
            db.session.add(department)
            db.session.flush()
    else:
        dept_id_raw = (form.department_id.data or '').strip()
        if dept_id_raw.isdigit():
            department = db.session.get(Department, int(dept_id_raw))
            if department and int(department.organization_id) != int(org_id):
                department = None

    # Check if user already has an active membership in this org
    user = User.query.filter_by(email=email).first()
    if user:
        existing_membership = (
            OrganizationMembership.query
            .filter_by(organization_id=int(org_id), user_id=int(user.id))
            .first()
        )
        if existing_membership and existing_membership.is_active:
            if not bool(user.password_hash):
                flash(f'An invitation has already been sent to {email}. You can resend it from the pending invites section below.', 'warning')
            else:
                flash(f'{email} is already a member of this organization.', 'warning')
            return redirect(url_for('main.org_admin_dashboard'))

    created_user = False
    try:
        if not user:
            user = User(
                email=email,
                email_verified=False,
                is_active=True,
                created_at=datetime.now(timezone.utc),
                organization_id=int(org_id),
            )
            db.session.add(user)
            db.session.flush()
            created_user = True

        membership = (
            OrganizationMembership.query
            .filter_by(organization_id=int(org_id), user_id=int(user.id))
            .first()
        )
        
        if membership:
            # Re-adding a previously removed member - treat as new invite
            membership.is_active = True
            membership.role_id = selected_role_id
            membership.department_id = int(department.id) if department else None
            # Reset invite acceptance tracking
            membership.invite_accepted_at = None
        else:
            membership = OrganizationMembership(
                organization_id=int(org_id),
                user_id=int(user.id),
                role_id=selected_role_id,
                is_active=True,
                department_id=(int(department.id) if department else None),
            )
            db.session.add(membership)

        # Keep legacy role string compatible with existing admin checks.
        try:
            from app.services.rbac import BUILTIN_ROLE_KEYS

            if selected_role and (selected_role.name or '').strip() == BUILTIN_ROLE_KEYS.ORG_ADMIN:
                membership.role = 'Admin'
            else:
                membership.role = 'User'
        except Exception:
            membership.role = membership.role or 'User'

        # Track invite details - send to all newly added/re-added members
        now = datetime.now(timezone.utc)
        membership.invited_at = membership.invited_at or now
        membership.invited_by_user_id = int(getattr(current_user, 'id', 0) or 0) or None
        membership.invite_last_sent_at = now
        membership.invite_send_count = int(membership.invite_send_count or 0) + 1
        membership.invite_revoked_at = None

        # Set the user's active organization to the one they're being invited to
        user.organization_id = int(org_id)

        db.session.commit()
    except Exception:
        db.session.rollback()
        flash('Failed to invite member. Please try again.', 'error')
        current_app.logger.exception('Failed inviting member')
        return redirect(url_for('main.org_admin_dashboard'))

    # Send invite email - for users without password, they set it via reset link
    # For users with existing password, they can use their existing password to login
    email_sent = False
    try:
        token = _org_invite_token(user)
        reset_url = url_for('auth.reset_password', token=token, _external=True)
        _send_invite_email(user, reset_url, organization)
        email_sent = True
    except Exception as e:
        current_app.logger.exception('Failed to send invite email')
        flash(f'User invited but email could not be sent. Error: {str(e)}. Please configure email settings.', 'warning')

    if created_user:
        if email_sent:
            flash(
                f'Invitation sent to {email}! The link expires in {_format_duration_seconds(_org_invite_token_ttl_seconds())}.',
                'success',
            )
        else:
            flash(f'User created but email not configured. Share this invite link manually with {email}.', 'warning')
    else:
        if email_sent:
            flash(
                f'User re-invited. The link expires in {_format_duration_seconds(_org_invite_token_ttl_seconds())}.',
                'success',
            )
        else:
            flash('User added to the organisation.', 'success')
    return redirect(url_for('main.org_admin_dashboard'))


@bp.route('/org/admin/departments/create', methods=['POST'])
@login_required
def org_admin_create_department():
    """Create a department for the active organization (AJAX helper)."""
    maybe = _require_org_permission('departments.manage')
    if maybe is not None:
        return maybe

    from app.main.forms import CreateDepartmentForm
    from app.models import Department
    from sqlalchemy import func

    org_id = _active_org_id()
    if not org_id:
        return jsonify({'success': False, 'error': 'No active organization'}), 400

    form = CreateDepartmentForm()
    if not form.validate_on_submit():
        # Keep response simple for UI.
        msg = 'Invalid department details.'
        if form.name.errors:
            msg = form.name.errors[0]
        elif form.color.errors:
            msg = form.color.errors[0]
        return jsonify({'success': False, 'error': msg}), 400

    name = (form.name.data or '').strip()
    color = (form.color.data or 'primary').strip() or 'primary'
    allowed_colors = {'primary', 'secondary', 'success', 'info', 'warning', 'danger', 'dark'}
    if color not in allowed_colors:
        color = 'primary'

    # Case-insensitive de-dupe by name within org.
    existing = (
        Department.query
        .filter(Department.organization_id == int(org_id))
        .filter(func.lower(Department.name) == func.lower(name))
        .first()
    )
    if existing:
        return jsonify({
            'success': True,
            'created': False,
            'department': {'id': int(existing.id), 'name': existing.name, 'color': existing.color},
        })

    try:
        dept = Department(organization_id=int(org_id), name=name, color=color)
        db.session.add(dept)
        db.session.commit()
        return jsonify({
            'success': True,
            'created': True,
            'department': {'id': int(dept.id), 'name': dept.name, 'color': dept.color},
        })
    except Exception:
        db.session.rollback()
        current_app.logger.exception('Failed creating department')
        return jsonify({'success': False, 'error': 'Failed to create department.'}), 500


@bp.route('/org/admin/departments/<int:dept_id>/edit', methods=['POST'])
@login_required
def org_admin_edit_department(dept_id):
    """Edit a department (AJAX helper)."""
    maybe = _require_org_permission('departments.manage')
    if maybe is not None:
        return maybe

    from app.main.forms import EditDepartmentForm
    from app.models import Department
    from sqlalchemy import func

    org_id = _active_org_id()
    if not org_id:
        return jsonify({'success': False, 'error': 'No active organization'}), 400

    dept = Department.query.filter_by(id=dept_id, organization_id=int(org_id)).first()
    if not dept:
        return jsonify({'success': False, 'error': 'Department not found'}), 404

    form = EditDepartmentForm()
    if not form.validate_on_submit():
        msg = 'Invalid department details.'
        if form.name.errors:
            msg = form.name.errors[0]
        elif form.color.errors:
            msg = form.color.errors[0]
        return jsonify({'success': False, 'error': msg}), 400

    name = (form.name.data or '').strip()
    color = (form.color.data or 'primary').strip() or 'primary'
    allowed_colors = {'primary', 'secondary', 'success', 'info', 'warning', 'danger', 'dark'}
    if color not in allowed_colors:
        color = 'primary'

    # Check for name conflict (case-insensitive, excluding current dept).
    conflict = (
        Department.query
        .filter(Department.organization_id == int(org_id))
        .filter(Department.id != dept_id)
        .filter(func.lower(Department.name) == func.lower(name))
        .first()
    )
    if conflict:
        return jsonify({'success': False, 'error': 'A department with this name already exists'}), 400

    try:
        dept.name = name
        dept.color = color
        db.session.commit()
        return jsonify({
            'success': True,
            'department': {'id': int(dept.id), 'name': dept.name, 'color': dept.color},
        })
    except Exception:
        db.session.rollback()
        current_app.logger.exception('Failed editing department')
        return jsonify({'success': False, 'error': 'Failed to edit department.'}), 500


@bp.route('/org/admin/departments/<int:dept_id>/delete', methods=['POST'])
@login_required
def org_admin_delete_department(dept_id):
    """Delete a department (AJAX helper). Members assigned to this department will have it set to NULL."""
    maybe = _require_org_permission('departments.manage')
    if maybe is not None:
        return maybe

    from app.main.forms import DeleteDepartmentForm
    from app.models import Department, OrganizationMembership

    org_id = _active_org_id()
    if not org_id:
        return jsonify({'success': False, 'error': 'No active organization'}), 400

    form = DeleteDepartmentForm()
    if not form.validate_on_submit():
        return jsonify({'success': False, 'error': 'Invalid request'}), 400

    dept = Department.query.filter_by(id=dept_id, organization_id=int(org_id)).first()
    if not dept:
        return jsonify({'success': False, 'error': 'Department not found'}), 404

    try:
        # Unassign members from this department.
        OrganizationMembership.query.filter_by(department_id=dept_id).update({'department_id': None})
        db.session.delete(dept)
        db.session.commit()
        return jsonify({'success': True})
    except Exception:
        db.session.rollback()
        current_app.logger.exception('Failed deleting department')
        return jsonify({'success': False, 'error': 'Failed to delete department.'}), 500


@bp.route('/org/admin/invite/resend', methods=['POST'])
@login_required
def org_admin_resend_invite():
    """Resend an invite email to a pending invited user (cooldown enforced)."""
    maybe = _require_org_permission('users.invite')
    if maybe is not None:
        return maybe

    from app.main.forms import PendingInviteResendForm
    from datetime import datetime, timezone

    org_id = _active_org_id()
    organization = db.session.get(Organization, int(org_id))
    if not organization:
        abort(404)

    form = PendingInviteResendForm()
    if not form.validate_on_submit():
        flash('Invalid request.', 'error')
        return redirect(url_for('main.org_admin_dashboard'))

    membership_id_raw = (form.membership_id.data or '').strip()
    if not membership_id_raw.isdigit():
        flash('Invalid request.', 'error')
        return redirect(url_for('main.org_admin_dashboard'))

    membership = db.session.get(OrganizationMembership, int(membership_id_raw))
    if not membership or int(membership.organization_id) != int(org_id):
        flash('Invite not found.', 'error')
        return redirect(url_for('main.org_admin_dashboard'))

    # Ensure we are reading the latest invite tracking fields in case this row
    # was updated in another session/process (or earlier request).
    try:
        db.session.refresh(membership)
    except Exception:
        pass

    user = db.session.get(User, int(membership.user_id)) if membership else None
    if not _is_pending_org_invite(membership, user):
        flash('That invite is no longer pending.', 'info')
        return redirect(url_for('main.org_admin_dashboard'))

    now = datetime.now(timezone.utc)
    last_sent = membership.invite_last_sent_at
    if last_sent:
        # SQLite can return naive datetimes; normalize to UTC-aware for arithmetic.
        if getattr(last_sent, 'tzinfo', None) is None:
            last_sent = last_sent.replace(tzinfo=timezone.utc)
        wait_seconds = _RESEND_ORG_INVITE_COOLDOWN_SECONDS - int((now - last_sent).total_seconds())
        if wait_seconds > 0:
            flash(f'Please wait {wait_seconds} seconds before resending this invite.', 'warning')
            return redirect(url_for('main.org_admin_dashboard'))

    try:
        membership.invite_last_sent_at = now
        membership.invite_send_count = int(membership.invite_send_count or 0) + 1
        db.session.commit()
    except Exception:
        db.session.rollback()
        current_app.logger.exception('Failed updating invite tracking')
        flash('Failed to resend invite. Please try again.', 'error')
        return redirect(url_for('main.org_admin_dashboard'))

    try:
        token = _org_invite_token(user)
        reset_url = url_for('auth.reset_password', token=token, _external=True)
        _send_invite_email(user, reset_url, organization)
    except Exception:
        current_app.logger.exception('Failed to send invite email')

    flash(f'Invite resent. The link expires in {_format_duration_seconds(_org_invite_token_ttl_seconds())}.', 'success')
    return redirect(url_for('main.org_admin_dashboard'))


@bp.route('/org/admin/invite/revoke', methods=['POST'])
@login_required
def org_admin_revoke_invite():
    """Revoke a pending invite by disabling the membership."""
    maybe = _require_org_permission('users.invite')
    if maybe is not None:
        return maybe

    from app.main.forms import PendingInviteRevokeForm
    from datetime import datetime, timezone

    org_id = _active_org_id()

    form = PendingInviteRevokeForm()
    if not form.validate_on_submit():
        flash('Invalid request.', 'error')
        return redirect(url_for('main.org_admin_dashboard'))

    membership_id_raw = (form.membership_id.data or '').strip()
    if not membership_id_raw.isdigit():
        flash('Invalid request.', 'error')
        return redirect(url_for('main.org_admin_dashboard'))

    membership = db.session.get(OrganizationMembership, int(membership_id_raw))
    if not membership or int(membership.organization_id) != int(org_id):
        flash('Invite not found.', 'error')
        return redirect(url_for('main.org_admin_dashboard'))

    # Guard: never allow an admin to revoke themselves via the invite flow.
    if int(membership.user_id) == int(current_user.id):
        flash('You cannot revoke your own access.', 'error')
        return redirect(url_for('main.org_admin_dashboard'))

    user = db.session.get(User, int(membership.user_id)) if membership else None
    if not _is_pending_org_invite(membership, user):
        flash('That invite is no longer pending.', 'info')
        return redirect(url_for('main.org_admin_dashboard'))

    try:
        membership.is_active = False
        membership.invite_revoked_at = datetime.now(timezone.utc)
        db.session.commit()
    except Exception:
        db.session.rollback()
        current_app.logger.exception('Failed revoking invite')
        flash('Failed to revoke invite. Please try again.', 'error')
        return redirect(url_for('main.org_admin_dashboard'))

    flash('Invite revoked.', 'success')
    return redirect(url_for('main.org_admin_dashboard'))


@bp.route('/org/admin/members/remove', methods=['POST'])
@login_required
def org_admin_remove_member():
    """Remove or disable a user's membership from the active organization."""
    maybe = _require_org_admin()
    if maybe is not None:
        return maybe

    from app.main.forms import MembershipActionForm

    org_id = _active_org_id()
    form = MembershipActionForm()
    if not form.validate_on_submit():
        flash('Invalid request.', 'error')
        return redirect(url_for('main.org_admin_dashboard'))

    membership_id_raw = (form.membership_id.data or '').strip()
    action = (form.action.data or '').strip().lower()
    
    if not membership_id_raw.isdigit():
        flash('Invalid membership.', 'error')
        return redirect(url_for('main.org_admin_dashboard'))
    
    if action not in ('disable', 'delete'):
        flash('Invalid action.', 'error')
        return redirect(url_for('main.org_admin_dashboard'))

    membership_id = int(membership_id_raw)
    membership = db.session.get(OrganizationMembership, membership_id)
    if not membership or int(membership.organization_id) != int(org_id):
        flash('Membership not found.', 'error')
        return redirect(url_for('main.org_admin_dashboard'))

    # Guard: do not remove/disable the last active user-manager.
    is_admin = _membership_has_permission(membership, 'users.manage')
    if is_admin and membership.is_active:
        active_memberships = (
            OrganizationMembership.query
            .filter_by(organization_id=int(org_id), is_active=True)
            .all()
        )
        active_admins = sum(1 for m in active_memberships if _membership_has_permission(m, 'users.manage'))
        if active_admins <= 1:
            flash('Cannot remove the last admin. Promote another member to admin first.', 'error')
            return redirect(url_for('main.org_admin_dashboard'))

    # Allow self-removal only when there is another active admin (if the user is an admin).
    if int(membership.user_id) == int(current_user.id):
        if is_admin and membership.is_active:
            active_memberships = (
                OrganizationMembership.query
                .filter_by(organization_id=int(org_id), is_active=True)
                .all()
            )
            active_admins = sum(1 for m in active_memberships if _membership_has_permission(m, 'users.manage'))
            if active_admins <= 1:
                flash('You are the only admin. Promote another admin before leaving the organisation.', 'error')
                return redirect(url_for('main.org_admin_dashboard'))

    try:
        if action == 'delete':
            # Completely delete the membership
            db.session.delete(membership)
            db.session.commit()
            flash('Member permanently removed from the organisation.', 'success')
        else:  # disable
            # Just deactivate
            membership.is_active = False
            db.session.commit()
            flash('Member disabled. You can re-enable them later if needed.', 'success')
    except Exception:
        db.session.rollback()
        flash('Failed to remove member. Please try again.', 'error')
        current_app.logger.exception('Failed removing member')

    return redirect(url_for('main.org_admin_dashboard'))


@bp.route('/theme', methods=['POST'])
def set_theme():
    """Persist theme preference in a cookie (light/dark)."""
    theme = (request.form.get('theme') or '').strip().lower()
    if theme not in {'light', 'dark'}:
        theme = 'light'

    # Redirect back to the originating page when possible.
    next_url = (request.form.get('next') or '').strip()
    if next_url and next_url.startswith('/'):
        redirect_target = next_url
    elif request.referrer:
        redirect_target = request.referrer
    else:
        redirect_target = url_for('main.dashboard') if current_user.is_authenticated else url_for('main.index')

    resp = make_response(redirect(redirect_target))
    resp.set_cookie(
        'theme',
        theme,
        max_age=60 * 60 * 24 * 365,  # 1 year
        samesite='Lax',
        secure=bool(request.is_secure),
    )
    return resp

@bp.route('/')
def index():
    """Home page route."""
    # If user is logged in and explicitly wants to switch accounts, show option
    if current_user.is_authenticated:
        # Check if user wants to see login/signup options (for account switching)
        if request.args.get('switch_account') == '1':
            flash('You are currently logged in. To switch accounts, please logout first.', 'info')
            # Don't auto-redirect, show home page with logout option
            return render_template('main/index.html', title='Home', show_logout=True)
        return redirect(url_for('main.dashboard'))
    return render_template('main/index.html', title='Home')


def _review_frequency_to_days(value: str | None) -> int | None:
    text = (value or '').strip().lower()
    if not text:
        return None

    if 'fortnight' in text:
        return 14
    if 'quarter' in text:
        return 90
    if 'month' in text and 'bi' in text:
        return 60
    if 'month' in text:
        return 30
    if 'annual' in text or 'year' in text:
        return 365
    if 'week' in text:
        return 7
    if 'day' in text:
        return 1

    import re

    match = re.search(r'(\d+)\s*(day|days|week|weeks|month|months|year|years)', text)
    if not match:
        return None

    amount = int(match.group(1))
    unit = match.group(2)
    if unit.startswith('day'):
        return amount
    if unit.startswith('week'):
        return amount * 7
    if unit.startswith('month'):
        return amount * 30
    if unit.startswith('year'):
        return amount * 365
    return None


def _build_dashboard_deadlines(*, org_id: int, limit: int = 3) -> list[dict]:
    now = datetime.now(timezone.utc)
    linked_requirement_ids = {
        int(req_id)
        for (req_id,) in (
            db.session.query(RequirementEvidenceLink.requirement_id)
            .filter(RequirementEvidenceLink.organization_id == int(org_id))
            .distinct()
            .all()
        )
        if req_id is not None
    }

    if not linked_requirement_ids:
        return []

    query_rows = (
        db.session.query(
            OrganizationRequirementAssessment,
            ComplianceRequirement,
            ComplianceFrameworkVersion,
        )
        .join(ComplianceRequirement, ComplianceRequirement.id == OrganizationRequirementAssessment.requirement_id)
        .join(ComplianceFrameworkVersion, ComplianceFrameworkVersion.id == ComplianceRequirement.framework_version_id)
        .filter(
            OrganizationRequirementAssessment.organization_id == int(org_id),
            ComplianceFrameworkVersion.is_active.is_(True),
            or_(
                ComplianceFrameworkVersion.organization_id.is_(None),
                ComplianceFrameworkVersion.organization_id == int(org_id),
            ),
        )
        .all()
    )

    candidates: list[dict] = []
    for assessment, requirement, framework in query_rows:
        if int(getattr(requirement, 'id', 0) or 0) not in linked_requirement_ids:
            continue

        review_days = _review_frequency_to_days(getattr(requirement, 'review_frequency', None))
        if not review_days:
            continue

        # Deadlines are meaningful only after a real assessment event.
        base_dt = assessment.last_assessed_at
        if not base_dt:
            continue
        if base_dt.tzinfo is None:
            base_dt = base_dt.replace(tzinfo=timezone.utc)

        due_dt = base_dt + timedelta(days=int(review_days))
        days_left = (due_dt.date() - now.date()).days

        # Show due/overdue items within a practical horizon.
        if days_left > 90:
            continue

        severity = 'info'
        if days_left <= 7:
            severity = 'danger'
        elif days_left <= 21:
            severity = 'warning'

        if days_left < 0:
            due_label = f'Overdue by {abs(days_left)} day(s)'
        elif days_left == 0:
            due_label = 'Due today'
        else:
            due_label = f'Due in {days_left} day(s)'

        scheme = (framework.scheme or 'Framework').strip()
        req_label = (requirement.requirement_id or requirement.quality_indicator_code or requirement.standard_name or 'Requirement').strip()
        title = f'{scheme}: {req_label}'

        candidates.append(
            {
                'title': title,
                'due_label': due_label,
                'days_left': int(days_left),
                'severity': severity,
                'review_frequency': (requirement.review_frequency or '').strip(),
            }
        )

    candidates.sort(key=lambda item: item['days_left'])
    return candidates[: max(1, int(limit))]


def _build_dashboard_bridge_stats(*, org_id: int) -> dict:
    """Build workflow bridge metrics from document review -> requirement assessment."""
    linked_doc_ids = {
        int(doc_id)
        for (doc_id,) in (
            db.session.query(RequirementEvidenceLink.document_id)
            .filter(RequirementEvidenceLink.organization_id == int(org_id))
            .distinct()
            .all()
        )
        if doc_id is not None
    }

    reviewed_doc_ids = {
        int(doc_id)
        for (doc_id,) in (
            db.session.query(Document.id)
            .filter(
                Document.organization_id == int(org_id),
                Document.is_active.is_(True),
                Document.ai_analysis_at.isnot(None),
            )
            .all()
        )
        if doc_id is not None
    }

    linked_requirement_ids = {
        int(req_id)
        for (req_id,) in (
            db.session.query(RequirementEvidenceLink.requirement_id)
            .filter(RequirementEvidenceLink.organization_id == int(org_id))
            .distinct()
            .all()
        )
        if req_id is not None
    }

    pending_assessment_count = 0
    if linked_requirement_ids:
        pending_assessment_count = (
            db.session.query(OrganizationRequirementAssessment)
            .filter(
                OrganizationRequirementAssessment.organization_id == int(org_id),
                OrganizationRequirementAssessment.requirement_id.in_(linked_requirement_ids),
                or_(
                    OrganizationRequirementAssessment.computed_flag.is_(None),
                    OrganizationRequirementAssessment.computed_flag.in_(['', 'Not assessed']),
                ),
            )
            .count()
        )

    return {
        'reviewed_docs': len(reviewed_doc_ids),
        'reviewed_docs_not_linked': len(reviewed_doc_ids - linked_doc_ids),
        'linked_requirements': len(linked_requirement_ids),
        'linked_requirements_pending_assessment': int(pending_assessment_count),
    }


@bp.route('/compliance-journey')
@login_required
def compliance_journey():
    """Legacy route retained for compatibility."""
    maybe = _require_active_org()
    if maybe is not None:
        return maybe

    flash('Compliance Journey has been retired. Use Requirements instead.', 'info')
    return redirect(url_for('main.compliance_requirements'))


@bp.route('/compliance-journey/persona', methods=['POST'])
@login_required
def compliance_journey_set_persona():
    """Legacy persona endpoint retained for compatibility."""
    maybe = _require_active_org()
    if maybe is not None:
        return maybe

    flash('Compliance Journey preferences are no longer required.', 'info')
    return redirect(url_for('main.compliance_requirements'))

@bp.route('/dashboard')
@login_required
def dashboard():
    """Dashboard route for authenticated users."""
    maybe = _require_active_org()
    if maybe is not None:
        return maybe

    org_id = _active_org_id()
    if not current_user.has_permission('documents.view', org_id=int(org_id)):
        abort(403)

    # Parallel non-blocking queries: recent docs + count.
    from sqlalchemy.orm import joinedload
    recent_documents = (
        Document.query
        .options(joinedload(Document.uploader))
        .filter_by(organization_id=org_id, is_active=True)
        .order_by(Document.uploaded_at.desc())
        .limit(5)
        .all()
    )
    recent_analysed_documents = (
        Document.query
        .options(joinedload(Document.uploader))
        .filter(
            Document.organization_id == int(org_id),
            Document.is_active.is_(True),
            Document.ai_analysis_at.isnot(None),
        )
        .order_by(Document.ai_analysis_at.desc())
        .limit(5)
        .all()
    )
    # Avoid full table scan for count; use an approximate or limit scope.
    total_documents = Document.query.filter_by(organization_id=org_id, is_active=True).limit(1000).count()
    reviewed_documents_count = (
        Document.query
        .filter(
            Document.organization_id == int(org_id),
            Document.is_active.is_(True),
            Document.ai_analysis_at.isnot(None),
        )
        .limit(1000)
        .count()
    )

    analytics_payload = analytics_service.build_dashboard_payload(organization_id=int(org_id))
    dashboard_summary = analytics_payload.get('summary') or {}
    dashboard_frameworks = (analytics_payload.get('framework_analytics') or [])[:6]
    dashboard_deadlines = _build_dashboard_deadlines(org_id=int(org_id), limit=3)
    dashboard_bridge = _build_dashboard_bridge_stats(org_id=int(org_id))

    return render_template('main/dashboard.html', 
                         title='Dashboard',
                         recent_documents=recent_documents,
                         recent_analysed_documents=recent_analysed_documents,
                         total_documents=total_documents,
                         dashboard_reviewed_documents=reviewed_documents_count,
                         dashboard_summary=dashboard_summary,
                         dashboard_frameworks=dashboard_frameworks,
                         dashboard_deadlines=dashboard_deadlines,
                         dashboard_bridge=dashboard_bridge)

@bp.route('/upload')
@login_required
def upload():
    """Upload page route."""
    maybe = _require_active_org()
    if maybe is not None:
        return maybe
    return render_template('main/upload.html', title='Upload Document')

@bp.route('/documents')
@login_required
def documents():
    """Documents listing route."""
    maybe = _require_active_org()
    if maybe is not None:
        return maybe

    org_id = _active_org_id()
    if not current_user.has_permission('documents.view', org_id=int(org_id)):
        abort(403)
    query = Document.query.filter_by(organization_id=org_id, is_active=True)
    user_documents = query.order_by(Document.uploaded_at.desc()).all()
    return render_template('main/documents.html', 
                         title='My Documents',
                         documents=user_documents)

@bp.route('/evidence-repository')
@login_required
def evidence_repository():
    """Evidence repository route to display all documents."""
    maybe = _require_active_org()
    if maybe is not None:
        return maybe

    org_id = _active_org_id()
    if not current_user.has_permission('documents.view', org_id=int(org_id)):
        abort(403)
    
    # Filters
    query_text = (request.args.get('q') or '').strip()
    selected_tag = (request.args.get('tag') or '').strip()
    file_type = (request.args.get('file_type') or '').strip().lower()
    date_from = (request.args.get('date_from') or '').strip()
    date_to = (request.args.get('date_to') or '').strip()

    # Pagination to avoid loading thousands of documents at once.
    page = request.args.get('page', 1, type=int)
    per_page = int(request.args.get('per_page', '50') or 50)
    per_page = min(max(per_page, 10), 200)  # clamp between 10-200
    
    # Use options to eager-load relationships and avoid N+1 queries
    from sqlalchemy.orm import joinedload
    query = (
        Document.query
        .options(joinedload(Document.uploader))
        .filter_by(organization_id=org_id, is_active=True)
    )

    if query_text or selected_tag:
        query = query.outerjoin(Document.tags)

    if query_text:
        like = f'%{query_text}%'
        text_filter = or_(
            Document.filename.ilike(like),
            Document.content_type.ilike(like),
            Document.search_text.ilike(like),
            DocumentTag.name.ilike(like),
        )
        try:
            bind = db.session.get_bind()
            if bind and bind.dialect.name == 'postgresql':
                text_filter = or_(
                    text_filter,
                    func.to_tsvector('simple', func.coalesce(Document.search_text, '')).op('@@')(
                        func.plainto_tsquery('simple', query_text)
                    ),
                )
        except Exception:
            pass
        query = query.filter(text_filter)

    if selected_tag:
        query = query.filter(func.lower(DocumentTag.name) == selected_tag.lower())

    if file_type == 'pdf':
        query = query.filter(Document.content_type.ilike('application/pdf%'))
    elif file_type == 'image':
        query = query.filter(Document.content_type.ilike('image/%'))
    elif file_type == 'word':
        query = query.filter(or_(Document.filename.ilike('%.doc'), Document.filename.ilike('%.docx')))

    if date_from:
        try:
            from_dt = datetime.strptime(date_from, '%Y-%m-%d')
            query = query.filter(Document.uploaded_at >= from_dt)
        except Exception:
            flash('Invalid From date filter.', 'warning')

    if date_to:
        try:
            to_dt = datetime.strptime(date_to, '%Y-%m-%d')
            to_dt_exclusive = datetime(to_dt.year, to_dt.month, to_dt.day, 23, 59, 59)
            query = query.filter(Document.uploaded_at <= to_dt_exclusive)
        except Exception:
            flash('Invalid To date filter.', 'warning')

    query = query.distinct()

    pagination = query.order_by(Document.uploaded_at.desc()).paginate(
        page=page, per_page=per_page, error_out=False
    )
    documents = pagination.items

    available_tags = (
        DocumentTag.query
        .filter_by(organization_id=int(org_id))
        .order_by(DocumentTag.name.asc())
        .all()
    )

    return render_template('main/evidence_repository.html', 
                         title='Evidence Repository',
                         documents=documents,
                         pagination=pagination,
                         available_tags=available_tags,
                         q=query_text,
                         selected_tag=selected_tag,
                         file_type=file_type,
                         date_from=date_from,
                         date_to=date_to)

def _authorized_org_document_or_404(doc_id: int) -> Document:
    if not getattr(current_user, 'is_authenticated', False):
        abort(404)

    org_id = _active_org_id()
    if not org_id:
        abort(404)

    membership = (
        OrganizationMembership.query
        .filter_by(user_id=int(current_user.id), organization_id=int(org_id), is_active=True)
        .first()
    )
    if not membership:
        abort(404)

    if not current_user.has_permission('documents.view', org_id=int(org_id)):
        abort(404)

    document = db.session.get(Document, int(doc_id))
    if not document or not getattr(document, 'is_active', True):
        abort(404)
    if int(document.organization_id) != int(org_id):
        abort(404)
    return document


def _is_previewable_document(document: Document) -> bool:
    content_type = (document.content_type or '').lower()
    filename = (document.filename or '').lower()
    if content_type.startswith('image/'):
        return True
    if content_type.startswith('text/'):
        return True
    if content_type.startswith('application/pdf'):
        return True
    return filename.endswith(('.pdf', '.png', '.jpg', '.jpeg', '.gif', '.bmp', '.txt'))


def _normalize_tags(raw_tags: str) -> list[str]:
    parts = [p.strip() for p in (raw_tags or '').replace(';', ',').split(',')]
    cleaned = []
    seen = set()
    for item in parts:
        if not item:
            continue
        normalized = ' '.join(item.split())[:64]
        key = normalized.lower()
        if not normalized or key in seen:
            continue
        seen.add(key)
        cleaned.append(normalized)
    return cleaned


def _refresh_document_search_text(document: Document):
    if not _documents_search_text_available():
        return
    base = [document.filename or '', document.content_type or '', document.extracted_text or '']
    tag_names = [t.name for t in (document.tags or []) if (t.name or '').strip()]
    base.extend(tag_names)
    document.search_text = ' '.join([b.strip() for b in base if (b or '').strip()])


@bp.route('/document/<int:doc_id>/download')
def download_document(doc_id):
    """Download a document."""
    from flask import send_file, abort
    from app.services.azure_storage import AzureBlobStorageService
    import io

    try:
        document = _authorized_org_document_or_404(int(doc_id))
    except NotFound:
        return ('', 404)
    
    try:
        storage_service = AzureBlobStorageService()
        result = storage_service.download_file(document.blob_name)
        if not result.get('success'):
            if result.get('error_code') == 'FILE_NOT_FOUND':
                abort(404)
            abort(500)

        blob_data = result.get('data')
        if not blob_data:
            abort(404)
        
        # Create file-like object
        file_stream = io.BytesIO(blob_data)
        file_stream.seek(0)
        
        # Send file to user
        return send_file(
            file_stream,
            mimetype=document.content_type,
            as_attachment=True,
            download_name=document.filename
        )
    except Exception as e:
        current_app.logger.exception(f'Error downloading document {doc_id}: {e}')
        abort(500)


@bp.route('/document/<int:doc_id>/preview')
@login_required
def preview_document(doc_id):
    """Stream a secure inline preview for supported file types."""
    from app.services.azure_storage import AzureBlobStorageService

    document = _authorized_org_document_or_404(int(doc_id))
    if not _is_previewable_document(document):
        flash('Preview is only available for PDF, image, and text documents.', 'info')
        return redirect(url_for('main.document_details', doc_id=int(document.id)))

    try:
        storage_service = AzureBlobStorageService()
        result = storage_service.download_file(document.blob_name)
        if not result.get('success'):
            abort(404)
        blob_data = result.get('data')
        if not blob_data:
            abort(404)

        file_stream = io.BytesIO(blob_data)
        file_stream.seek(0)
        response = send_file(
            file_stream,
            mimetype=document.content_type or 'application/octet-stream',
            as_attachment=False,
            download_name=document.filename,
        )
        response.headers['X-Content-Type-Options'] = 'nosniff'
        return response
    except Exception:
        current_app.logger.exception('Failed to preview document %s', doc_id)
        abort(500)


@bp.route('/documents/download-zip', methods=['POST'])
@login_required
def download_documents_zip():
    """Bulk download selected documents as a ZIP archive."""
    from app.services.azure_storage import AzureBlobStorageService

    maybe = _require_active_org()
    if maybe is not None:
        return maybe

    org_id = _active_org_id()
    if not current_user.has_permission('documents.view', org_id=int(org_id)):
        abort(403)

    requested_ids = [int(v) for v in request.form.getlist('doc_ids') if str(v).isdigit()]
    if not requested_ids:
        flash('Select at least one document to download.', 'warning')
        return redirect(url_for('main.evidence_repository'))

    documents = (
        Document.query
        .filter(
            Document.organization_id == int(org_id),
            Document.is_active.is_(True),
            Document.id.in_(requested_ids),
        )
        .order_by(Document.uploaded_at.desc())
        .all()
    )
    if not documents:
        flash('No valid documents found for ZIP export.', 'error')
        return redirect(url_for('main.evidence_repository'))

    storage_service = AzureBlobStorageService()
    zip_buffer = io.BytesIO()
    used_names = set()
    added = 0

    with zipfile.ZipFile(zip_buffer, mode='w', compression=zipfile.ZIP_DEFLATED) as archive:
        for document in documents:
            if not (document.blob_name or '').strip():
                continue
            result = storage_service.download_file(document.blob_name)
            if not result.get('success'):
                continue
            blob_data = result.get('data')
            if not blob_data:
                continue

            base_name = (document.filename or f'document_{document.id}').strip()
            safe_name = base_name
            suffix = 1
            while safe_name in used_names:
                name, ext = os.path.splitext(base_name)
                safe_name = f"{name} ({suffix}){ext}"
                suffix += 1
            used_names.add(safe_name)
            archive.writestr(safe_name, blob_data)
            added += 1

    if added == 0:
        flash('Unable to package selected documents.', 'error')
        return redirect(url_for('main.evidence_repository'))

    zip_buffer.seek(0)
    filename = f"documents_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip"
    return send_file(
        zip_buffer,
        mimetype='application/zip',
        as_attachment=True,
        download_name=filename,
    )


@bp.route('/document/<int:doc_id>/tags', methods=['POST'])
@login_required
def add_document_tags(doc_id):
    """Attach one or more tags to a document."""
    maybe = _require_active_org()
    if maybe is not None:
        return maybe

    org_id = _active_org_id()
    if not current_user.has_permission('documents.view', org_id=int(org_id)):
        abort(403)

    document = db.session.get(Document, int(doc_id))
    if not document or int(document.organization_id) != int(org_id):
        abort(404)

    parsed_tags = _normalize_tags(request.form.get('tags') or '')
    if not parsed_tags:
        flash('Add at least one valid tag.', 'warning')
        return redirect(url_for('main.document_details', doc_id=int(doc_id)))

    added_count = 0
    for tag_name in parsed_tags:
        normalized_name = tag_name.lower()
        tag = DocumentTag.query.filter_by(
            organization_id=int(org_id),
            normalized_name=normalized_name,
        ).first()
        if not tag:
            tag = DocumentTag(
                organization_id=int(org_id),
                name=tag_name,
                normalized_name=normalized_name,
            )
            db.session.add(tag)
            db.session.flush()

        if not any(int(t.id) == int(tag.id) for t in (document.tags or [])):
            document.tags.append(tag)
            added_count += 1

    _refresh_document_search_text(document)
    db.session.commit()

    if added_count > 0:
        flash(f'Added {added_count} tag(s).', 'success')
    else:
        flash('All tags already exist on this document.', 'info')
    return redirect(url_for('main.document_details', doc_id=int(doc_id)))


@bp.route('/document/<int:doc_id>/tags/<int:tag_id>/delete', methods=['POST'])
@login_required
def remove_document_tag(doc_id, tag_id):
    """Remove a tag from a document."""
    maybe = _require_active_org()
    if maybe is not None:
        return maybe

    org_id = _active_org_id()
    if not current_user.has_permission('documents.view', org_id=int(org_id)):
        abort(403)

    document = db.session.get(Document, int(doc_id))
    if not document or int(document.organization_id) != int(org_id):
        abort(404)

    tag = DocumentTag.query.filter_by(id=int(tag_id), organization_id=int(org_id)).first()
    if not tag:
        abort(404)

    if any(int(t.id) == int(tag.id) for t in (document.tags or [])):
        document.tags.remove(tag)
        _refresh_document_search_text(document)
        db.session.commit()
        flash('Tag removed.', 'success')
    else:
        flash('Tag is not attached to this document.', 'info')

    return redirect(url_for('main.document_details', doc_id=int(doc_id)))

@bp.route('/document/<int:doc_id>/delete', methods=['POST'])
@login_required
def delete_document(doc_id):
    """Delete a document."""
    maybe = _require_active_org()
    if maybe is not None:
        return maybe

    org_id = _active_org_id()
    if not current_user.has_permission('documents.delete', org_id=int(org_id)):
        abort(403)

    document = db.session.get(Document, int(doc_id))
    if not document:
        flash('Document not found or access denied.', 'error')
        return redirect(url_for('main.evidence_repository'))
    if document.organization_id != org_id:
        flash('Document not found or access denied.', 'error')
        return redirect(url_for('main.evidence_repository'))

    success, error_message = _soft_delete_document(document)
    if success:
        flash(f'Document "{document.filename}" deleted successfully.', 'success')
    else:
        flash(error_message or 'Error deleting document. Please try again.', 'error')
    
    return redirect(url_for('main.evidence_repository'))


def _soft_delete_document(document: Document) -> tuple[bool, str | None]:
    from app.services.azure_storage import AzureBlobStorageService

    try:
        if getattr(document, 'blob_name', None):
            storage_service = AzureBlobStorageService()
            delete_result = storage_service.delete_file(document.blob_name)
            if not delete_result.get('success'):
                raise Exception(delete_result.get('error') or 'Delete failed')
        else:
            current_app.logger.warning('Document %s has no blob_name; skipping Azure deletion', document.id)

        document.is_active = False
        db.session.commit()
        return True, None
    except Exception as exc:
        db.session.rollback()
        current_app.logger.exception('Error deleting document %s: %s', getattr(document, 'id', None), exc)
        return False, 'Error deleting document. Please try again.'


@bp.route('/documents/delete-selected', methods=['POST'])
@login_required
def delete_selected_documents():
    """Bulk delete selected repository documents after explicit confirmation."""
    maybe = _require_active_org()
    if maybe is not None:
        return maybe

    org_id = _active_org_id()
    if not current_user.has_permission('documents.delete', org_id=int(org_id)):
        abort(403)

    requested_ids = [int(v) for v in request.form.getlist('doc_ids') if str(v).isdigit()]
    if not requested_ids:
        flash('Select at least one document to delete.', 'warning')
        return redirect(url_for('main.evidence_repository'))

    confirmation_text = (request.form.get('confirmation_text') or '').strip()
    if confirmation_text != 'DELETE':
        flash('Type DELETE exactly to confirm bulk deletion.', 'warning')
        return redirect(url_for('main.evidence_repository'))

    documents = (
        Document.query
        .filter(
            Document.organization_id == int(org_id),
            Document.is_active.is_(True),
            Document.id.in_(requested_ids),
        )
        .order_by(Document.uploaded_at.desc())
        .all()
    )
    if not documents:
        flash('No valid documents found for deletion.', 'error')
        return redirect(url_for('main.evidence_repository'))

    deleted_count = 0
    failed_documents: list[str] = []
    for document in documents:
        success, _error_message = _soft_delete_document(document)
        if success:
            deleted_count += 1
        else:
            failed_documents.append(document.filename or f'Document {document.id}')

    if deleted_count > 0:
        flash(f'Deleted {deleted_count} document(s).', 'success')
    if failed_documents:
        shown = ', '.join(failed_documents[:3])
        if len(failed_documents) > 3:
            shown = f'{shown}, and {len(failed_documents) - 3} more'
        flash(f'Could not delete {len(failed_documents)} document(s): {shown}.', 'warning')

    return redirect(url_for('main.evidence_repository'))

@bp.route('/document/<int:doc_id>/details')
@login_required
def document_details(doc_id):
    """View document details."""
    maybe = _require_active_org()
    if maybe is not None:
        return maybe

    from flask import abort
    
    # Get document from database
    document = db.session.get(Document, int(doc_id))
    
    # Check if document exists and belongs to user
    org_id = _active_org_id()
    if not current_user.has_permission('documents.view', org_id=int(org_id)):
        abort(403)
    if not document:
        abort(404)
    if document.organization_id != org_id:
        abort(404)

    linked_requirements = (
        db.session.query(RequirementEvidenceLink, ComplianceRequirement, OrganizationRequirementAssessment)
        .join(ComplianceRequirement, ComplianceRequirement.id == RequirementEvidenceLink.requirement_id)
        .outerjoin(
            OrganizationRequirementAssessment,
            and_(
                OrganizationRequirementAssessment.organization_id == RequirementEvidenceLink.organization_id,
                OrganizationRequirementAssessment.requirement_id == RequirementEvidenceLink.requirement_id,
            ),
        )
        .filter(
            RequirementEvidenceLink.organization_id == int(org_id),
            RequirementEvidenceLink.document_id == int(document.id),
        )
        .order_by(RequirementEvidenceLink.linked_at.desc())
        .all()
    )

    summary_sections = _split_document_ai_summary(document.ai_summary)
    
    return render_template('main/document_details.html',
                         title=f'Document: {document.filename}',
                         document=document,
                         linked_requirements=linked_requirements,
                         summary_sections=summary_sections)


def _split_document_ai_summary(summary_text: str | None) -> dict[str, str]:
    value = (summary_text or '').replace('\r\n', '\n').strip()
    sections = {
        'why': '',
        'missing': '',
        'next_action': '',
    }
    if not value:
        return sections

    patterns = [
        ('why', r'1\)\s*Why this status'),
        ('missing', r'2\)\s*Missing evidence'),
        ('next_action', r'3\)\s*Recommended next action'),
    ]
    import re
    matches = []
    for key, pattern in patterns:
        match = re.search(pattern, value, flags=re.IGNORECASE)
        if match:
            matches.append((key, match.start(), match.end()))
    if not matches:
        sections['why'] = value
        return sections
    for index, (key, _start, end) in enumerate(matches):
        next_start = matches[index + 1][1] if index + 1 < len(matches) else len(value)
        sections[key] = value[end:next_start].strip()
    return sections


def _analyze_and_persist_document(document: Document, *, org_id: int, user_id: int) -> tuple[bool, str | None, int]:
    """Analyze one stored document and persist review/mapping fields."""
    from app.services.azure_storage import AzureBlobStorageService

    try:
        storage_service = AzureBlobStorageService()
        result = storage_service.download_file(document.blob_name)
        if not result.get('success') or not result.get('data'):
            return False, 'Could not load the document from storage for analysis.', 0

        analysis = document_analysis_service.analyze_document_bytes(
            filename=document.filename,
            raw_bytes=result.get('data') or b'',
            organization_id=int(org_id),
        )
        if not analysis.get('success'):
            return False, analysis.get('error') or 'Document analysis failed.', 0

        old_requirement_ids = {
            int(req_id)
            for req_id, in (
                db.session.query(RequirementEvidenceLink.requirement_id)
                .filter_by(organization_id=int(org_id), document_id=int(document.id))
                .all()
            )
        }

        RequirementEvidenceLink.query.filter_by(
            organization_id=int(org_id),
            document_id=int(document.id),
        ).delete(synchronize_session=False)

        matched_requirements = analysis.get('matched_requirements') or []
        new_requirement_ids = set()
        for item in matched_requirements:
            requirement_db_id = item.get('requirement_db_id')
            if not requirement_db_id:
                continue
            new_requirement_ids.add(int(requirement_db_id))
            db.session.add(
                RequirementEvidenceLink(
                    organization_id=int(org_id),
                    requirement_id=int(requirement_db_id),
                    document_id=int(document.id),
                    evidence_bucket=(item.get('evidence_bucket') or 'system')[:30],
                    rationale_note=item.get('rationale_note'),
                    linked_by_user_id=int(user_id),
                )
            )

        document.extracted_text = analysis.get('extracted_text') or None
        _refresh_document_search_text(document)
        document.ai_status = analysis.get('status')
        document.ai_confidence = analysis.get('confidence')
        document.ai_focus_area = analysis.get('focus_area')
        document.ai_question = analysis.get('question')
        document.ai_summary = analysis.get('summary')
        document.ai_provider = analysis.get('provider')
        document.ai_model = analysis.get('model_used')
        document.ai_retrieval_mode = analysis.get('retrieval_mode')
        document.ai_analysis_at = datetime.now(timezone.utc)
        db.session.commit()

        for requirement_id in old_requirement_ids.union(new_requirement_ids):
            try:
                compliance_scoring_service.recompute_requirement_assessment(
                    organization_id=int(org_id),
                    requirement_id=int(requirement_id),
                    assessed_by_user_id=int(user_id),
                    commit=False,
                )
            except Exception:
                current_app.logger.exception('Failed recomputing requirement assessment for requirement %s', requirement_id)
        db.session.commit()

        return True, None, len(matched_requirements)
    except Exception:
        db.session.rollback()
        current_app.logger.exception('Failed to analyze document %s for org %s', getattr(document, 'id', None), org_id)
        return False, 'Unexpected error while analyzing this document.', 0


@bp.route('/document/<int:doc_id>/analyze', methods=['POST'])
@login_required
def analyze_document(doc_id):
    """Run explicit AI-assisted review for a stored document."""
    maybe = _require_active_org()
    if maybe is not None:
        return maybe

    org_id = _active_org_id()
    if not current_user.has_permission('documents.view', org_id=int(org_id)):
        abort(403)

    document = _authorized_org_document_or_404(int(doc_id))

    success, error_message, matched_count = _analyze_and_persist_document(
        document,
        org_id=int(org_id),
        user_id=int(current_user.id),
    )
    if not success:
        flash(error_message or 'Document analysis failed.', 'error')
        return redirect(url_for('main.document_details', doc_id=int(document.id)))

    if matched_count:
        flash(f'Analysis complete. Mapped this document to {matched_count} requirement(s).', 'success')
    else:
        flash('Analysis complete, but no explicit requirement matches were found yet.', 'warning')
    return redirect(url_for('main.document_details', doc_id=int(document.id)))


@bp.route('/documents/analyze-selected', methods=['POST'])
@login_required
def analyze_selected_documents():
    """Run analysis for selected repository documents."""
    maybe = _require_active_org()
    if maybe is not None:
        return maybe

    org_id = _active_org_id()
    if not current_user.has_permission('documents.view', org_id=int(org_id)):
        abort(403)

    requested_ids = [int(v) for v in request.form.getlist('doc_ids') if str(v).isdigit()]
    if not requested_ids:
        flash('Select at least one document to analyze.', 'warning')
        return redirect(url_for('main.evidence_repository'))

    documents = (
        Document.query
        .filter(
            Document.organization_id == int(org_id),
            Document.is_active.is_(True),
            Document.id.in_(requested_ids),
        )
        .order_by(Document.uploaded_at.desc())
        .all()
    )
    if not documents:
        flash('No valid documents found for analysis.', 'error')
        return redirect(url_for('main.evidence_repository'))

    analyzed_count = 0
    total_mapped = 0
    failed_documents: list[str] = []
    for document in documents:
        success, _error_message, matched_count = _analyze_and_persist_document(
            document,
            org_id=int(org_id),
            user_id=int(current_user.id),
        )
        if success:
            analyzed_count += 1
            total_mapped += int(matched_count or 0)
        else:
            failed_documents.append(document.filename or f'Document {document.id}')

    if analyzed_count > 0:
        flash(
            f'Analysis completed for {analyzed_count} document(s) with {total_mapped} total requirement match(es).',
            'success',
        )
    if failed_documents:
        shown = ', '.join(failed_documents[:3])
        if len(failed_documents) > 3:
            shown = f'{shown}, and {len(failed_documents) - 3} more'
        flash(f'Could not analyze {len(failed_documents)} document(s): {shown}.', 'warning')

    return redirect(url_for('main.evidence_repository'))

@bp.route('/ai-evidence')
@login_required
def ai_evidence():
    """Legacy AI Evidence entry point; redirect to the primary repository review flow."""
    maybe = _require_active_org()
    if maybe is not None:
        return maybe

    org_id = _active_org_id()
    if not current_user.has_permission('documents.view', org_id=int(org_id)):
        abort(403)
    flash('AI evidence review now lives in the Evidence Repository and each document details page.', 'info')
    return redirect(url_for('main.evidence_repository'))


@bp.route('/organization/settings', methods=['GET', 'POST'])
@login_required
def organization_settings():
    from flask import abort, flash, make_response, request
    from app.main.forms import (
        OrganizationBillingAccessForm,
        OrganizationBillingForm,
        OrganizationMonthlyReportForm,
        OrganizationProfileSettingsForm,
    )

    maybe = _require_org_permission('org.manage')
    if maybe is not None:
        return maybe

    if not getattr(current_user, 'organization_id', None):
        flash('No organisation is associated with this account.', 'error')
        return redirect(url_for('main.dashboard'))

    organization = db.session.get(Organization, int(current_user.organization_id))
    if not organization:
        abort(404)

    profile_form = OrganizationProfileSettingsForm(obj=organization)
    billing_form = OrganizationBillingForm(obj=organization)
    monthly_report_form = OrganizationMonthlyReportForm(obj=organization)
    billing_access_form = OrganizationBillingAccessForm(
        billing_plan_code=billing_service.normalize_plan_code(organization.billing_plan_code or organization.subscription_tier),
        billing_status=((organization.billing_status or 'inactive').strip().lower() or 'inactive'),
        billing_internal_override=bool(getattr(organization, 'billing_internal_override', False)),
        billing_demo_override_enabled=bool(getattr(organization, 'billing_demo_override_until', None)),
        billing_override_reason=(organization.billing_override_reason or ''),
    )
    is_super_admin = bool(_is_super_admin_user())
    is_org_admin = bool(current_user.has_permission('users.manage', org_id=int(organization.id)))
    billing_state = billing_service.resolve_entitlements(
        organization,
        actor_email=getattr(current_user, 'email', None),
    )
    billing_catalog = billing_service.plan_catalog()

    billing_result = (request.args.get('billing') or '').strip().lower()
    if billing_result == 'success':
        flash('Checkout completed. Billing status may take a few moments to update.', 'success')
    elif billing_result == 'cancelled':
        flash('Checkout cancelled. No billing changes were made.', 'info')

    if request.method == 'POST':
        submitted = (request.form.get('form_name') or '').strip()

        if submitted == 'profile':
            if profile_form.validate_on_submit():
                organization.name = profile_form.name.data.strip()
                organization.abn = (profile_form.abn.data or '').strip() or None
                organization.acn = (profile_form.acn.data or '').strip() or None
                organization.contact_number = (profile_form.contact_number.data or '').strip() or None
                organization.address = (profile_form.address.data or '').strip() or None
                organization.contact_email = (profile_form.contact_email.data or '').strip().lower() or None

                logo_file = profile_form.logo.data
                if logo_file and getattr(logo_file, 'filename', ''):
                    success, message = _update_organization_logo(organization, logo_file)
                    if not success:
                        flash(message, 'error')
                        return render_template(
                            'main/organization_settings.html',
                            title='Organization Settings',
                            profile_form=profile_form,
                            billing_form=billing_form,
                            billing_access_form=billing_access_form,
                            monthly_report_form=monthly_report_form,
                            is_super_admin=is_super_admin,
                            is_org_admin=is_org_admin,
                            organization=organization,
                            billing_state=billing_state,
                            billing_catalog=billing_catalog,
                        )

                try:
                    db.session.commit()
                    flash('Organization profile saved.', 'success')
                    return redirect(url_for('main.organization_settings'))
                except Exception:
                    db.session.rollback()
                    flash('Failed to save organization profile. Please try again.', 'error')

        elif submitted == 'billing':
            if billing_form.validate_on_submit():
                organization.billing_email = (billing_form.billing_email.data or '').strip().lower() or None
                organization.billing_address = (billing_form.billing_address.data or '').strip() or None

                try:
                    db.session.commit()
                    flash('Billing details saved.', 'success')
                    return redirect(url_for('main.organization_settings'))
                except Exception:
                    db.session.rollback()
                    flash('Failed to save billing details. Please try again.', 'error')
        elif submitted == 'billing_access':
            if not is_super_admin:
                abort(403)

            if billing_access_form.validate_on_submit():
                organization.billing_plan_code = billing_service.normalize_plan_code(billing_access_form.billing_plan_code.data)
                organization.subscription_tier = organization.billing_plan_code.capitalize()
                organization.billing_status = ((billing_access_form.billing_status.data or 'inactive').strip().lower() or 'inactive')
                organization.billing_internal_override = bool(billing_access_form.billing_internal_override.data)
                if bool(billing_access_form.billing_demo_override_enabled.data):
                    organization.billing_demo_override_until = datetime(2099, 12, 31, tzinfo=timezone.utc)
                else:
                    organization.billing_demo_override_until = None
                organization.billing_override_reason = (billing_access_form.billing_override_reason.data or '').strip() or None

                try:
                    db.session.commit()
                    invalidate_org_switcher_context_cache(int(current_user.id), int(organization.id))
                    flash('Billing access controls updated.', 'success')
                    return redirect(url_for('main.organization_settings'))
                except Exception:
                    db.session.rollback()
                    flash('Failed to update billing access controls. Please try again.', 'error')
        elif submitted == 'monthly_reports':
            if not is_org_admin:
                abort(403)

            if monthly_report_form.validate_on_submit():
                previous_enabled = bool(getattr(organization, 'monthly_report_enabled', False))
                previous_recipient = ((getattr(organization, 'monthly_report_recipient_email', '') or '').strip().lower())

                organization.monthly_report_enabled = bool(monthly_report_form.monthly_report_enabled.data)
                organization.monthly_report_recipient_email = (
                    (monthly_report_form.monthly_report_recipient_email.data or '').strip().lower() or None
                )

                try:
                    db.session.commit()

                    new_enabled = bool(getattr(organization, 'monthly_report_enabled', False))
                    new_recipient = ((getattr(organization, 'monthly_report_recipient_email', '') or '').strip().lower())
                    should_send_setup_email = bool(
                        new_enabled
                        and new_recipient
                        and ((not previous_enabled and new_enabled) or (new_recipient != previous_recipient))
                    )
                    if should_send_setup_email:
                        try:
                            sent = notification_service.send_monthly_report_setup_confirmation(
                                recipient_email=new_recipient,
                                organization_name=(organization.name or '').strip() or 'your organisation',
                            )
                            if sent:
                                flash('Setup confirmation email sent to the monthly report recipient.', 'success')
                            else:
                                flash('Monthly report settings saved, but setup confirmation email could not be sent.', 'warning')
                        except Exception:
                            current_app.logger.exception(
                                'Failed sending monthly report setup confirmation for org %s',
                                int(organization.id),
                            )
                            flash('Monthly report settings saved, but setup confirmation email failed to send.', 'warning')

                    flash('Monthly report settings saved.', 'success')
                    return redirect(url_for('main.organization_settings'))
                except Exception:
                    db.session.rollback()
                    flash('Failed to save monthly report settings. Please try again.', 'error')
        else:
            flash('Invalid form submission.', 'error')

    return render_template(
        'main/organization_settings.html',
        title='Organization Profile',
        profile_form=profile_form,
        billing_form=billing_form,
        billing_access_form=billing_access_form,
        monthly_report_form=monthly_report_form,
        is_super_admin=is_super_admin,
        is_org_admin=is_org_admin,
        organization=organization,
        billing_state=billing_state,
        billing_catalog=billing_catalog,
    )


@bp.route('/organization/ai-controls', methods=['GET', 'POST'])
@login_required
def organization_ai_controls():
    maybe = _require_org_permission('org.manage')
    if maybe is not None:
        return maybe
    flash('AI Controls has been retired. Use Organisation Profile instead.', 'info')
    return redirect(url_for('main.organization_settings'))


@bp.route('/organization/ai-controls/retention-run', methods=['POST'])
@login_required
def organization_ai_retention_run():
    from app.main.forms import OrganizationAIUsageRetentionForm
    from datetime import datetime, timezone, timedelta

    maybe = _require_org_permission('org.manage')
    if maybe is not None:
        return maybe

    form = OrganizationAIUsageRetentionForm()
    if not form.validate_on_submit():
        flash('Invalid retention request. Check retention days.', 'error')
        return redirect(url_for('main.organization_settings'))

    org_id = _active_org_id()
    min_days = max(1, int(current_app.config.get('MIN_AUDIT_LOG_RETENTION_DAYS') or 90))
    requested_days = int(form.days.data or (current_app.config.get('AI_USAGE_RETENTION_DAYS') or 90))
    days = max(min_days, requested_days)
    dry_run = bool(form.dry_run.data)
    cutoff = datetime.now(timezone.utc) - timedelta(days=days)

    if requested_days < min_days:
        flash(f'Minimum retention floor is {min_days} days. Using {min_days} days.', 'info')

    q = (
        AIUsageEvent.query
        .filter(AIUsageEvent.organization_id == int(org_id))
        .filter(AIUsageEvent.created_at < cutoff)
    )
    candidate_rows = int(q.count())

    if dry_run:
        flash(f'Dry run: {candidate_rows} AI usage events older than {days} days would be deleted.', 'info')
        return redirect(url_for('main.organization_settings'))

    try:
        deleted = int(q.delete(synchronize_session=False) or 0)
        db.session.commit()
        flash(f'Retention cleanup complete. Deleted {deleted} AI usage events older than {days} days.', 'success')
    except Exception:
        db.session.rollback()
        current_app.logger.exception('Failed to run AI usage retention cleanup for org %s', org_id)
        flash('Retention cleanup failed. Please try again.', 'error')

    return redirect(url_for('main.organization_settings'))


@bp.route('/organization/ai-controls/usage.csv')
@login_required
def organization_ai_usage_csv():
    maybe = _require_org_permission('org.manage')
    if maybe is not None:
        return maybe

    org_id = _active_org_id()
    event_filter = (request.args.get('event') or '').strip()
    time_range = (request.args.get('time_range') or 'all').strip().lower()

    from datetime import datetime, timezone, timedelta

    now = datetime.now(timezone.utc)
    range_map = {
        '24h': timedelta(hours=24),
        '7d': timedelta(days=7),
        '30d': timedelta(days=30),
        'all': None,
    }
    selected_delta = range_map.get(time_range, None)
    start_time = (now - selected_delta) if selected_delta is not None else None

    events_query = AIUsageEvent.query.filter_by(organization_id=int(org_id))
    if event_filter:
        events_query = events_query.filter(AIUsageEvent.event == event_filter)
    if start_time is not None:
        events_query = events_query.filter(AIUsageEvent.created_at >= start_time)

    events = (
        events_query
        .order_by(AIUsageEvent.created_at.desc())
        .limit(5000)
        .all()
    )

    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow([
        'timestamp_utc',
        'event',
        'mode',
        'provider',
        'model',
        'user_id',
        'prompt_tokens',
        'completion_tokens',
        'total_tokens',
        'latency_ms',
    ])

    for event in events:
        writer.writerow([
            event.created_at.isoformat() if event.created_at else '',
            event.event or '',
            event.mode or '',
            event.provider or '',
            event.model or '',
            event.user_id or '',
            int(event.prompt_tokens or 0),
            int(event.completion_tokens or 0),
            int(event.total_tokens or 0),
            int(event.latency_ms or 0),
        ])

    response = make_response(output.getvalue())
    response.headers['Content-Type'] = 'text/csv; charset=utf-8'
    response.headers['Content-Disposition'] = 'attachment; filename=ai_usage_events.csv'
    return response


@bp.route('/organization/logo')
@login_required
def organization_logo():
    from flask import abort, send_file
    import io
    org_id = getattr(current_user, 'organization_id', None)
    if not org_id:
        abort(404)

    organization = db.session.get(Organization, int(org_id))
    if not organization or not organization.logo_blob_name:
        abort(404)

    # Strong cache validators based on blob name (changes on upload).
    etag = f'W/"orglogo-{int(org_id)}-{organization.logo_blob_name}"'
    req_version = (request.args.get('v') or '').strip()
    inm = request.headers.get('If-None-Match')
    if _etag_matches_if_none_match(inm, etag):
        resp = make_response('', 304)
        resp.headers['ETag'] = etag
        if req_version and req_version == (organization.logo_blob_name or ''):
            resp.headers['Cache-Control'] = 'private, max-age=31536000, immutable'
        else:
            resp.headers['Cache-Control'] = 'private, max-age=300'
        current_app.logger.info('Org logo 304 (etag match) org_id=%s', org_id)
        return resp

    # Cache logo bytes in-memory to avoid repeated Azure fetches.
    try:
        logo_cache_seconds = int((current_app.config.get('ORG_LOGO_CACHE_SECONDS') or 300))
    except Exception:
        logo_cache_seconds = 300

    t0 = time.monotonic()
    cached = _get_cached_org_logo(int(org_id), organization.logo_blob_name)
    if cached:
        blob_data, cached_type = cached
        content_type = cached_type or organization.logo_content_type
        current_app.logger.info('Org logo served from memory cache org_id=%s', org_id)
    else:
        disk_cached = _get_disk_cached_org_logo(int(org_id), organization.logo_blob_name)
        if disk_cached:
            blob_data, disk_type = disk_cached
            content_type = disk_type or organization.logo_content_type
            _set_cached_org_logo(int(org_id), organization.logo_blob_name, blob_data, content_type, ttl_seconds=logo_cache_seconds)
            current_app.logger.info('Org logo served from disk cache org_id=%s', org_id)
        else:
            from app.services.azure_storage_service import azure_storage_service
            # Pass org_id to ensure correct path (org_X/ prefix)
            blob_data = azure_storage_service.download_blob(organization.logo_blob_name, organization_id=int(org_id))
            if not blob_data:
                abort(404)
            content_type = organization.logo_content_type
            _set_cached_org_logo(int(org_id), organization.logo_blob_name, blob_data, content_type, ttl_seconds=logo_cache_seconds)
            _set_disk_cached_org_logo(int(org_id), organization.logo_blob_name, blob_data, content_type)
            elapsed = time.monotonic() - t0
            current_app.logger.warning('Org logo fetched from Azure org_id=%s took %.2fs', org_id, elapsed)

    file_stream = io.BytesIO(blob_data)
    file_stream.seek(0)
    resp = send_file(
        file_stream,
        mimetype=content_type or 'application/octet-stream',
        as_attachment=False,
        download_name='logo'
    )

    resp.headers['ETag'] = etag
    if req_version and req_version == (organization.logo_blob_name or ''):
        resp.headers['Cache-Control'] = 'private, max-age=31536000, immutable'
    else:
        resp.headers['Cache-Control'] = 'private, max-age=300'
    return resp

@bp.route('/organization/<int:org_id>/logo')
@login_required
def organization_logo_by_id(org_id):
    """Serve logo for any organization the user is a member of."""
    from flask import abort, send_file
    import io

    # Check user has access to this org
    membership = (
        OrganizationMembership.query
        .filter_by(user_id=int(current_user.id), organization_id=int(org_id), is_active=True)
        .first()
    )
    if not membership:
        abort(404)

    organization = db.session.get(Organization, int(org_id))
    if not organization or not organization.logo_blob_name:
        abort(404)

    etag = f'W/"orglogo-{int(org_id)}-{organization.logo_blob_name}"'
    req_version = (request.args.get('v') or '').strip()
    inm = request.headers.get('If-None-Match')
    if _etag_matches_if_none_match(inm, etag):
        resp = make_response('', 304)
        resp.headers['ETag'] = etag
        if req_version and req_version == (organization.logo_blob_name or ''):
            resp.headers['Cache-Control'] = 'private, max-age=31536000, immutable'
        else:
            resp.headers['Cache-Control'] = 'private, max-age=300'
        current_app.logger.info('Org logo(by_id) 304 (etag match) org_id=%s', org_id)
        return resp

    try:
        logo_cache_seconds = int((current_app.config.get('ORG_LOGO_CACHE_SECONDS') or 300))
    except Exception:
        logo_cache_seconds = 300

    t0 = time.monotonic()
    cached = _get_cached_org_logo(int(org_id), organization.logo_blob_name)
    if cached:
        blob_data, cached_type = cached
        content_type = cached_type or organization.logo_content_type
        current_app.logger.info('Org logo(by_id) served from memory cache org_id=%s', org_id)
    else:
        disk_cached = _get_disk_cached_org_logo(int(org_id), organization.logo_blob_name)
        if disk_cached:
            blob_data, disk_type = disk_cached
            content_type = disk_type or organization.logo_content_type
            _set_cached_org_logo(int(org_id), organization.logo_blob_name, blob_data, content_type, ttl_seconds=logo_cache_seconds)
            current_app.logger.info('Org logo(by_id) served from disk cache org_id=%s', org_id)
        else:
            from app.services.azure_storage_service import azure_storage_service
            # Pass org_id to ensure correct path (org_X/ prefix)
            blob_data = azure_storage_service.download_blob(organization.logo_blob_name, organization_id=int(org_id))
            if not blob_data:
                abort(404)
            content_type = organization.logo_content_type
            _set_cached_org_logo(int(org_id), organization.logo_blob_name, blob_data, content_type, ttl_seconds=logo_cache_seconds)
            _set_disk_cached_org_logo(int(org_id), organization.logo_blob_name, blob_data, content_type)
            elapsed = time.monotonic() - t0
            current_app.logger.warning('Org logo(by_id) fetched from Azure org_id=%s took %.2fs', org_id, elapsed)

    file_stream = io.BytesIO(blob_data)
    file_stream.seek(0)
    resp = send_file(
        file_stream,
        mimetype=content_type or 'application/octet-stream',
        as_attachment=False,
        download_name=f'{organization.name}_logo'
    )

    resp.headers['ETag'] = etag
    if req_version and req_version == (organization.logo_blob_name or ''):
        resp.headers['Cache-Control'] = 'private, max-age=31536000, immutable'
    else:
        resp.headers['Cache-Control'] = 'private, max-age=300'
    return resp

@bp.route('/ai-evidence/<int:entry_id>')
@login_required
def ai_evidence_detail(entry_id):
    """Legacy AI Evidence detail entry point; redirect to canonical document details."""
    document = _authorized_org_document_or_404(int(entry_id))
    flash('AI evidence detail has moved into the document details view.', 'info')
    return redirect(url_for('main.document_details', doc_id=int(document.id)))

@bp.route('/gap-analysis')
@login_required
def gap_analysis():
    """Legacy compliance dashboard route; redirects to AI Review workspace."""
    maybe = _require_active_org()
    if maybe is not None:
        return maybe

    org_id = _active_org_id()
    if not current_user.has_permission('documents.view', org_id=int(org_id)):
        abort(403)
    flash('Gap Analysis has moved into AI Review workspaces.', 'info')
    return redirect(url_for('main.ai_demo'))


@bp.route('/compliance-requirements')
@login_required
def compliance_requirements():
    """Requirements workboard with evidence-bucket coverage and due-state tracking."""
    maybe = _require_active_org()
    if maybe is not None:
        return maybe

    org_id = _active_org_id()
    if not current_user.has_permission('documents.view', org_id=int(org_id)):
        abort(403)

    q = (request.args.get('q') or '').strip()
    status_filter = _normalize_computed_flag_filter(request.args.get('status') or '')
    module_filter = (request.args.get('module') or '').strip()
    bucket_filter = (request.args.get('bucket') or '').strip().lower()
    due_filter = (request.args.get('due') or '').strip().lower()
    show_requirements_panel = (request.args.get('show_requirements_panel') or '').strip().lower() in {'1', 'true', 'yes'}

    page = _clamp_int(request.args.get('page', 1), default=1, minimum=1, maximum=10_000)
    per_page = 25

    base_query = (
        db.session.query(ComplianceRequirement, OrganizationRequirementAssessment)
        .join(ComplianceFrameworkVersion, ComplianceFrameworkVersion.id == ComplianceRequirement.framework_version_id)
        .outerjoin(
            OrganizationRequirementAssessment,
            and_(
                OrganizationRequirementAssessment.organization_id == int(org_id),
                OrganizationRequirementAssessment.requirement_id == ComplianceRequirement.id,
            ),
        )
        .filter(
            ComplianceFrameworkVersion.is_active.is_(True),
            or_(
                ComplianceFrameworkVersion.organization_id.is_(None),
                ComplianceFrameworkVersion.organization_id == int(org_id),
            ),
        )
    )

    if q:
        like = f'%{q}%'
        base_query = base_query.filter(
            or_(
                ComplianceRequirement.requirement_id.ilike(like),
                ComplianceRequirement.module_name.ilike(like),
                ComplianceRequirement.quality_indicator_code.ilike(like),
                ComplianceRequirement.quality_indicator_text.ilike(like),
                ComplianceRequirement.outcome_code.ilike(like),
                ComplianceRequirement.outcome_text.ilike(like),
                ComplianceRequirement.evidence_owner_role.ilike(like),
            )
        )

    if module_filter:
        base_query = base_query.filter(ComplianceRequirement.module_name == module_filter)

    if status_filter:
        base_query = base_query.filter(
            OrganizationRequirementAssessment.computed_flag.in_(_computed_flag_filter_values(status_filter))
        )

    raw_rows = base_query.order_by(
        ComplianceRequirement.module_name.asc().nullslast(),
        ComplianceRequirement.requirement_id.asc(),
    ).all()

    bucket_rows = (
        db.session.query(
            RequirementEvidenceLink.requirement_id,
            RequirementEvidenceLink.evidence_bucket,
            func.count(RequirementEvidenceLink.id),
        )
        .filter(RequirementEvidenceLink.organization_id == int(org_id))
        .group_by(RequirementEvidenceLink.requirement_id, RequirementEvidenceLink.evidence_bucket)
        .all()
    )

    bucket_counts_by_requirement: dict[int, dict[str, int]] = {}
    for requirement_id, evidence_bucket, total_links in bucket_rows:
        if requirement_id is None:
            continue
        rid = int(requirement_id)
        bucket_counts_by_requirement.setdefault(rid, {})[(evidence_bucket or '').strip().lower()] = int(total_links or 0)

    now_dt = datetime.now(timezone.utc)
    work_rows: list[dict] = []
    module_options = set()

    for requirement, assessment in raw_rows:
        module_name = (requirement.module_name or '').strip() or 'General'
        module_options.add(module_name)

        req_id = int(requirement.id)
        bucket_counts = {
            'system': int(bucket_counts_by_requirement.get(req_id, {}).get('system', 0)),
            'implementation': int(bucket_counts_by_requirement.get(req_id, {}).get('implementation', 0)),
            'workforce': int(bucket_counts_by_requirement.get(req_id, {}).get('workforce', 0)),
            'participant': int(bucket_counts_by_requirement.get(req_id, {}).get('participant', 0)),
        }

        if bucket_filter in {'system', 'implementation', 'workforce', 'participant'} and bucket_counts.get(bucket_filter, 0) <= 0:
            continue

        required_buckets = _required_buckets_for_requirement(requirement)
        linked_required_count = sum(1 for bucket in required_buckets if int(bucket_counts.get(bucket, 0)) > 0)
        bucket_coverage_pct = int(round((linked_required_count / max(1, len(required_buckets))) * 100))

        due_meta = _requirement_due_meta(requirement=requirement, assessment=assessment, now_dt=now_dt)
        due_days = due_meta.get('days_left')
        if due_filter == 'overdue' and not (due_days is not None and int(due_days) < 0):
            continue
        if due_filter == '30d' and not (due_days is not None and 0 <= int(due_days) <= 30):
            continue
        if due_filter == 'unscheduled' and due_days is not None:
            continue

        work_rows.append(
            {
                'requirement': requirement,
                'assessment': assessment,
                'module_name': module_name,
                'plain_title': _requirement_plain_title(requirement),
                'display_code': _requirement_display_code(requirement),
                'owner_role': (requirement.evidence_owner_role or '').strip() or 'Not assigned',
                'bucket_counts': bucket_counts,
                'required_buckets': required_buckets,
                'bucket_coverage_pct': int(bucket_coverage_pct),
                'due': due_meta,
            }
        )

    total_rows = len(work_rows)
    total_pages = max(1, (total_rows + per_page - 1) // per_page)
    if page > total_pages:
        page = total_pages
    start_index = (page - 1) * per_page
    end_index = start_index + per_page
    paged_rows = work_rows[start_index:end_index]

    overdue_count = sum(1 for row in work_rows if row.get('due', {}).get('days_left') is not None and int(row.get('due', {}).get('days_left') or 0) < 0)
    fully_linked_count = sum(1 for row in work_rows if int(row.get('bucket_coverage_pct') or 0) >= 100)
    risk_count = sum(
        1
        for row in work_rows
        if ((row.get('assessment').computed_flag if row.get('assessment') else '') or '').strip() in {'Critical gap', 'High risk gap', 'red', 'amber'}
    )

    summary = {
        'total_requirements': int(total_rows),
        'fully_linked_requirements': int(fully_linked_count),
        'overdue_reviews': int(overdue_count),
        'at_risk_requirements': int(risk_count),
    }

    lookback_30d = now_dt - timedelta(days=30)
    monthly_snapshot = {
        'uploads_30d': int(
            Document.query
            .filter(
                Document.organization_id == int(org_id),
                Document.is_active.is_(True),
                Document.uploaded_at >= lookback_30d,
            )
            .count()
        ),
        'assessments_30d': int(
            OrganizationRequirementAssessment.query
            .filter(
                OrganizationRequirementAssessment.organization_id == int(org_id),
                OrganizationRequirementAssessment.last_assessed_at.isnot(None),
                OrganizationRequirementAssessment.last_assessed_at >= lookback_30d,
            )
            .count()
        ),
        'links_30d': int(
            RequirementEvidenceLink.query
            .filter(
                RequirementEvidenceLink.organization_id == int(org_id),
                RequirementEvidenceLink.linked_at >= lookback_30d,
            )
            .count()
        ),
    }

    doc_rows = (
        db.session.query(
            Document.id,
            Document.filename,
            Document.uploaded_at,
            func.max(RequirementEvidenceLink.linked_at),
            func.count(RequirementEvidenceLink.id),
        )
        .outerjoin(
            RequirementEvidenceLink,
            and_(
                RequirementEvidenceLink.document_id == Document.id,
                RequirementEvidenceLink.organization_id == int(org_id),
            ),
        )
        .filter(
            Document.organization_id == int(org_id),
            Document.is_active.is_(True),
        )
        .group_by(Document.id, Document.filename, Document.uploaded_at)
        .order_by(func.max(RequirementEvidenceLink.linked_at).desc().nullslast(), Document.uploaded_at.desc())
        .limit(25)
        .all()
    )

    doc_requirement_rows = (
        db.session.query(
            RequirementEvidenceLink.document_id,
            RequirementEvidenceLink.linked_at,
            RequirementEvidenceLink.rationale_note,
            ComplianceRequirement.requirement_id,
            ComplianceRequirement.quality_indicator_text,
            ComplianceRequirement.outcome_text,
        )
        .join(ComplianceRequirement, ComplianceRequirement.id == RequirementEvidenceLink.requirement_id)
        .filter(RequirementEvidenceLink.organization_id == int(org_id))
        .order_by(RequirementEvidenceLink.linked_at.desc().nullslast())
        .all()
    )

    req_preview_by_doc: dict[int, list[dict]] = {}
    for doc_id, _linked_at, rationale_note, req_code, qi_text, outcome_text in doc_requirement_rows:
        did = int(doc_id or 0)
        if not did:
            continue
        entries = req_preview_by_doc.setdefault(did, [])
        display_value = ' '.join((req_code or '').split()).strip()
        if not display_value:
            text_source = ' '.join((qi_text or outcome_text or '').split()).strip()
            if text_source:
                display_value = text_source[:42].rstrip() + ('...' if len(text_source) > 42 else '')
            else:
                display_value = 'Requirement'

        existing_codes = {(entry.get('code') or '').strip() for entry in entries}
        if display_value in existing_codes:
            continue

        reason_text = ' '.join((rationale_note or '').split()).strip()
        if not reason_text:
            reason_source = ' '.join((qi_text or outcome_text or '').split()).strip()
            if reason_source:
                reason_text = reason_source[:180].rstrip() + ('...' if len(reason_source) > 180 else '')
            else:
                reason_text = 'Linked by document evidence matching for this requirement.'

        if len(entries) < 6:
            entries.append(
                {
                    'code': display_value,
                    'reason': reason_text,
                }
            )

    document_link_overview: list[dict] = []
    for doc_id, filename, uploaded_at, last_linked_at, link_count in doc_rows:
        did = int(doc_id or 0)
        document_link_overview.append(
            {
                'document_id': did,
                'filename': (filename or '').strip() or f'Document #{did}',
                'uploaded_at': uploaded_at,
                'last_linked_at': last_linked_at,
                'linked_count': int(link_count or 0),
                'requirements_preview': req_preview_by_doc.get(did, []),
            }
        )

    owner_action_queue: list[dict] = []
    for row in work_rows:
        assessment = row.get('assessment')
        due_days = row.get('due', {}).get('days_left')
        flag = ((assessment.computed_flag if assessment else '') or '').strip().lower()
        coverage_pct = int(row.get('bucket_coverage_pct') or 0)

        reasons: list[str] = []
        priority_score = 0

        if due_days is not None and int(due_days) < 0:
            reasons.append('Overdue review')
            priority_score += 60 + min(30, abs(int(due_days)))

        if flag in {'critical gap', 'red'}:
            reasons.append('Critical gap')
            priority_score += 100
        elif flag in {'high risk gap', 'amber'}:
            reasons.append('High risk gap')
            priority_score += 70

        if coverage_pct < 100:
            reasons.append('Evidence coverage incomplete')
            priority_score += 40

        if not reasons:
            continue

        owner_action_queue.append(
            {
                'priority_score': int(priority_score),
                'requirement_id': row.get('requirement').requirement_id,
                'plain_title': row.get('plain_title') or _requirement_plain_title(row.get('requirement')),
                'display_code': row.get('display_code') or _requirement_display_code(row.get('requirement')),
                'module_name': row.get('module_name'),
                'owner_role': row.get('owner_role'),
                'reason': ', '.join(reasons),
                'due_label': row.get('due', {}).get('label') or 'No due date',
                'status': (assessment.computed_flag if assessment else '') or 'Not assessed',
                'action_url': url_for('main.compliance_requirement_detail', requirement_db_id=int(row.get('requirement').id)),
            }
        )

    owner_action_queue.sort(key=lambda item: int(item.get('priority_score') or 0), reverse=True)
    owner_action_queue = owner_action_queue[:8]

    upcoming_review_queue: list[dict] = []
    for row in work_rows:
        due_days = row.get('due', {}).get('days_left')
        if due_days is None:
            continue
        if int(due_days) < 0 or int(due_days) > 45:
            continue

        upcoming_review_queue.append(
            {
                'days_left': int(due_days),
                'requirement_id': row.get('requirement').requirement_id,
                'plain_title': row.get('plain_title') or _requirement_plain_title(row.get('requirement')),
                'display_code': row.get('display_code') or _requirement_display_code(row.get('requirement')),
                'module_name': row.get('module_name'),
                'owner_role': row.get('owner_role'),
                'due_label': row.get('due', {}).get('label') or '',
                'coverage_pct': int(row.get('bucket_coverage_pct') or 0),
                'action_url': url_for('main.compliance_requirement_detail', requirement_db_id=int(row.get('requirement').id)),
            }
        )

    upcoming_review_queue.sort(key=lambda item: int(item.get('days_left') or 0))
    upcoming_review_queue = upcoming_review_queue[:8]

    pagination = {
        'page': int(page),
        'pages': int(total_pages),
        'has_prev': bool(page > 1),
        'has_next': bool(page < total_pages),
        'prev_num': int(max(1, page - 1)),
        'next_num': int(min(total_pages, page + 1)),
    }

    return render_template(
        'main/compliance_requirements.html',
        title='Requirements Workboard',
        rows=paged_rows,
        q=q,
        status_filter=status_filter,
        module_filter=module_filter,
        bucket_filter=bucket_filter,
        due_filter=due_filter,
        show_requirements_panel=show_requirements_panel,
        module_options=sorted(module_options),
        summary=summary,
        monthly_snapshot=monthly_snapshot,
        document_link_overview=document_link_overview,
        owner_action_queue=owner_action_queue,
        upcoming_review_queue=upcoming_review_queue,
        pagination=pagination,
    )


def _normalize_computed_flag_filter(value: str) -> str:
    normalized = (value or '').strip().lower().replace('_', ' ')
    mapping = {
        'critical gap': 'Critical gap',
        'high risk gap': 'High risk gap',
        'ok': 'OK',
        'mature': 'Mature',
        'red': 'Critical gap',
        'amber': 'High risk gap',
        'green': 'OK',
    }
    return mapping.get(normalized, '')


def _computed_flag_filter_values(canonical_flag: str) -> list[str]:
    mapping = {
        'Critical gap': ['Critical gap', 'red', 'Red'],
        'High risk gap': ['High risk gap', 'amber', 'Amber'],
        'OK': ['OK', 'ok', 'green', 'Green'],
        'Mature': ['Mature', 'mature'],
    }
    return mapping.get(canonical_flag, [canonical_flag])


def _requirement_due_meta(*, requirement: ComplianceRequirement, assessment: OrganizationRequirementAssessment | None, now_dt: datetime) -> dict:
    review_days = _review_frequency_to_days(getattr(requirement, 'review_frequency', None))
    if not review_days:
        return {
            'days_left': None,
            'label': 'No review schedule',
            'tone': 'secondary',
        }

    if not assessment or not assessment.last_assessed_at:
        return {
            'days_left': None,
            'label': 'Awaiting first assessment',
            'tone': 'secondary',
        }

    base_dt = assessment.last_assessed_at
    if base_dt.tzinfo is None:
        base_dt = base_dt.replace(tzinfo=timezone.utc)

    due_dt = base_dt + timedelta(days=int(review_days))
    days_left = int((due_dt.date() - now_dt.date()).days)
    if days_left < 0:
        return {
            'days_left': days_left,
            'label': f'Overdue by {abs(days_left)} day(s)',
            'tone': 'danger',
        }
    if days_left == 0:
        return {
            'days_left': 0,
            'label': 'Due today',
            'tone': 'warning',
        }
    if days_left <= 30:
        return {
            'days_left': days_left,
            'label': f'Due in {days_left} day(s)',
            'tone': 'warning',
        }
    return {
        'days_left': days_left,
        'label': f'Due in {days_left} day(s)',
        'tone': 'success',
    }


def _required_buckets_for_requirement(requirement: ComplianceRequirement) -> list[str]:
    required = ['system', 'implementation']
    if bool(getattr(requirement, 'requires_workforce_evidence', False)):
        required.append('workforce')
    if bool(getattr(requirement, 'requires_participant_evidence', False)):
        required.append('participant')
    return required


def _requirement_display_code(requirement: ComplianceRequirement | None) -> str:
    if not requirement:
        return 'Requirement'
    return (
        (getattr(requirement, 'requirement_id', None) or '')
        or (getattr(requirement, 'quality_indicator_code', None) or '')
        or (getattr(requirement, 'outcome_code', None) or '')
        or 'Requirement'
    ).strip()


def _requirement_plain_title(requirement: ComplianceRequirement | None, *, max_chars: int = 120) -> str:
    if not requirement:
        return 'Compliance requirement'

    candidates = [
        getattr(requirement, 'quality_indicator_text', None),
        getattr(requirement, 'outcome_text', None),
        getattr(requirement, 'standard_name', None),
    ]

    source = ''
    for raw in candidates:
        text = ' '.join(str(raw or '').split()).strip()
        if text:
            source = text
            break

    if not source:
        return _requirement_display_code(requirement)

    sentence = source.split('. ')[0].strip() or source
    if len(sentence) > max_chars:
        sentence = sentence[: max_chars - 1].rstrip() + '...'
    return sentence


def _requirement_bucket_counts(*, org_id: int, requirement_id: int) -> dict[str, int]:
    counts = {'system': 0, 'implementation': 0, 'workforce': 0, 'participant': 0}
    rows = (
        db.session.query(
            RequirementEvidenceLink.evidence_bucket,
            func.count(RequirementEvidenceLink.id),
        )
        .filter(
            RequirementEvidenceLink.organization_id == int(org_id),
            RequirementEvidenceLink.requirement_id == int(requirement_id),
        )
        .group_by(RequirementEvidenceLink.evidence_bucket)
        .all()
    )
    for bucket, total in rows:
        key = (bucket or '').strip().lower()
        if key in counts:
            counts[key] = int(total or 0)
    return counts


def _pick_bucket_for_requirement(*, requirement: ComplianceRequirement, bucket_counts: dict[str, int], preferred_bucket: str | None = None) -> str:
    required = _required_buckets_for_requirement(requirement)
    preferred = (preferred_bucket or '').strip().lower()
    allowed = {'system', 'implementation', 'workforce', 'participant'}

    if preferred in required and int(bucket_counts.get(preferred, 0)) <= 0:
        return preferred

    for bucket in required:
        if int(bucket_counts.get(bucket, 0)) <= 0:
            return bucket

    if preferred in allowed:
        return preferred

    return 'implementation'


def _auto_link_from_analyzed_documents(
    *,
    org_id: int,
    user_id: int,
    target_requirement_id: int | None = None,
    target_document_id: int | None = None,
) -> dict[str, int]:
    """Auto-create requirement evidence links from already analyzed/extracted documents."""
    docs_query = (
        Document.query
        .filter(
            Document.organization_id == int(org_id),
            Document.is_active.is_(True),
            Document.extracted_text.isnot(None),
        )
        .order_by(Document.ai_analysis_at.desc().nullslast(), Document.uploaded_at.desc())
    )
    if target_document_id is not None:
        docs_query = docs_query.filter(Document.id == int(target_document_id))
    docs = docs_query.all()

    existing_query = RequirementEvidenceLink.query.filter_by(organization_id=int(org_id))
    if target_requirement_id is not None:
        existing_query = existing_query.filter_by(requirement_id=int(target_requirement_id))

    existing_links = {
        (int(link.requirement_id), int(link.document_id), (link.evidence_bucket or '').strip().lower())
        for link in existing_query.all()
    }

    bucket_counts_cache: dict[int, dict[str, int]] = {}
    touched_requirements: set[int] = set()
    links_added = 0
    docs_scanned = 0

    for document in docs:
        text = (document.extracted_text or '').strip()
        if not text:
            continue

        docs_scanned += 1
        matched = document_analysis_service._match_requirements(
            text=text,
            filename=document.filename or '',
            organization_id=int(org_id),
            top_k=5,
        )

        for item in matched:
            requirement_id = int(item.get('requirement_db_id') or 0)
            if not requirement_id:
                continue
            if target_requirement_id is not None and int(target_requirement_id) != int(requirement_id):
                continue

            requirement = db.session.get(ComplianceRequirement, int(requirement_id))
            if requirement is None:
                continue

            counts = bucket_counts_cache.get(int(requirement_id))
            if counts is None:
                counts = _requirement_bucket_counts(org_id=int(org_id), requirement_id=int(requirement_id))
                bucket_counts_cache[int(requirement_id)] = counts

            selected_bucket = _pick_bucket_for_requirement(
                requirement=requirement,
                bucket_counts=counts,
                preferred_bucket=item.get('evidence_bucket') or '',
            )
            dedupe_key = (int(requirement_id), int(document.id), selected_bucket)
            if dedupe_key in existing_links:
                continue

            db.session.add(
                RequirementEvidenceLink(
                    organization_id=int(org_id),
                    requirement_id=int(requirement_id),
                    document_id=int(document.id),
                    evidence_bucket=selected_bucket,
                    rationale_note=(item.get('rationale_note') or None),
                    linked_by_user_id=int(user_id),
                )
            )
            existing_links.add(dedupe_key)
            counts[selected_bucket] = int(counts.get(selected_bucket, 0)) + 1
            links_added += 1
            touched_requirements.add(int(requirement_id))

    if touched_requirements:
        for requirement_id in touched_requirements:
            try:
                compliance_scoring_service.recompute_requirement_assessment(
                    organization_id=int(org_id),
                    requirement_id=int(requirement_id),
                    assessed_by_user_id=int(user_id),
                    commit=False,
                )
            except Exception:
                current_app.logger.exception('Failed to recompute assessment for requirement %s', requirement_id)
        db.session.commit()

    return {
        'docs_scanned': int(docs_scanned),
        'links_added': int(links_added),
        'requirements_updated': int(len(touched_requirements)),
    }


def _get_org_visible_requirement_or_404(requirement_db_id: int, org_id: int):
    requirement = (
        db.session.query(ComplianceRequirement)
        .join(ComplianceFrameworkVersion, ComplianceFrameworkVersion.id == ComplianceRequirement.framework_version_id)
        .filter(
            ComplianceRequirement.id == int(requirement_db_id),
            ComplianceFrameworkVersion.is_active.is_(True),
            or_(
                ComplianceFrameworkVersion.organization_id.is_(None),
                ComplianceFrameworkVersion.organization_id == int(org_id),
            ),
        )
        .first()
    )
    if not requirement:
        abort(404)
    return requirement


@bp.route('/compliance-requirements/<int:requirement_db_id>')
@login_required
def compliance_requirement_detail(requirement_db_id):
    """Requirement detail with linked evidence documents."""
    maybe = _require_active_org()
    if maybe is not None:
        return maybe

    org_id = _active_org_id()
    if not current_user.has_permission('documents.view', org_id=int(org_id)):
        abort(403)

    requirement = _get_org_visible_requirement_or_404(requirement_db_id=int(requirement_db_id), org_id=int(org_id))

    assessment = OrganizationRequirementAssessment.query.filter_by(
        organization_id=int(org_id),
        requirement_id=int(requirement.id),
    ).first()

    linked_evidence = (
        RequirementEvidenceLink.query
        .filter_by(organization_id=int(org_id), requirement_id=int(requirement.id))
        .order_by(RequirementEvidenceLink.linked_at.desc())
        .all()
    )

    linked_doc_ids = {int(link.document_id) for link in linked_evidence}
    bucket_counts = _requirement_bucket_counts(org_id=int(org_id), requirement_id=int(requirement.id))
    required_buckets = _required_buckets_for_requirement(requirement)
    missing_required_buckets = [bucket for bucket in required_buckets if int(bucket_counts.get(bucket, 0)) <= 0]
    available_documents = (
        Document.query
        .filter_by(organization_id=int(org_id), is_active=True)
        .order_by(Document.uploaded_at.desc())
        .all()
    )

    return render_template(
        'main/compliance_requirement_detail.html',
        title=_requirement_plain_title(requirement),
        requirement=requirement,
        requirement_plain_title=_requirement_plain_title(requirement),
        requirement_display_code=_requirement_display_code(requirement),
        assessment=assessment,
        linked_evidence=linked_evidence,
        available_documents=available_documents,
        linked_doc_ids=linked_doc_ids,
        required_buckets=required_buckets,
        bucket_counts=bucket_counts,
        missing_required_buckets=missing_required_buckets,
    )


@bp.route('/compliance-requirements/auto-link', methods=['POST'])
@login_required
def compliance_requirements_auto_link():
    """Easy-mode action: auto-link all analyzed documents into requirements."""
    maybe = _require_active_org()
    if maybe is not None:
        return maybe

    org_id = _active_org_id()
    if not current_user.has_permission('documents.view', org_id=int(org_id)):
        abort(403)

    try:
        result = _auto_link_from_analyzed_documents(org_id=int(org_id), user_id=int(current_user.id))
    except Exception:
        db.session.rollback()
        current_app.logger.exception('Bulk auto-link failed for org %s', org_id)
        flash('Auto-link failed. Please try again.', 'error')
        return redirect(url_for('main.compliance_requirements'))

    if int(result.get('links_added') or 0) > 0:
        flash(
            f"Auto-linked {int(result.get('links_added') or 0)} evidence item(s) across {int(result.get('requirements_updated') or 0)} requirement(s).",
            'success',
        )
    else:
        flash('No new auto-links were found. Analyze more documents first or link manually.', 'info')

    return redirect(url_for('main.compliance_requirements'))


@bp.route('/compliance-requirements/auto-link-document/<int:document_id>', methods=['POST'])
@login_required
def compliance_requirements_auto_link_document(document_id):
    """Auto-link analyzed evidence from one document into matching requirements."""
    maybe = _require_active_org()
    if maybe is not None:
        return maybe

    org_id = _active_org_id()
    if not current_user.has_permission('documents.view', org_id=int(org_id)):
        abort(403)

    document = db.session.get(Document, int(document_id))
    if not document or int(document.organization_id) != int(org_id) or not bool(document.is_active):
        flash('Document not found.', 'error')
        return redirect(url_for('main.compliance_requirements'))

    try:
        result = _auto_link_from_analyzed_documents(
            org_id=int(org_id),
            user_id=int(current_user.id),
            target_document_id=int(document.id),
        )
    except Exception:
        db.session.rollback()
        current_app.logger.exception('Document auto-link failed for doc %s (org %s)', int(document_id), int(org_id))
        flash('Auto-link failed for this document. Please try again.', 'error')
        return redirect(url_for('main.compliance_requirements'))

    if int(result.get('links_added') or 0) > 0:
        flash(
            f"Auto-linked {int(result.get('links_added') or 0)} evidence item(s) from {document.filename}.",
            'success',
        )
    else:
        flash('No new links found for this document. Analyze it first or review links manually.', 'info')

    return redirect(url_for('main.compliance_requirements'))


@bp.route('/compliance-requirements/<int:requirement_db_id>/auto-link', methods=['POST'])
@login_required
def compliance_requirement_auto_link(requirement_db_id):
    """Easy-mode action: auto-link analyzed documents for one requirement."""
    maybe = _require_active_org()
    if maybe is not None:
        return maybe

    org_id = _active_org_id()
    if not current_user.has_permission('documents.view', org_id=int(org_id)):
        abort(403)

    requirement = _get_org_visible_requirement_or_404(requirement_db_id=int(requirement_db_id), org_id=int(org_id))

    try:
        result = _auto_link_from_analyzed_documents(
            org_id=int(org_id),
            user_id=int(current_user.id),
            target_requirement_id=int(requirement.id),
        )
    except Exception:
        db.session.rollback()
        current_app.logger.exception('Requirement auto-link failed for requirement %s', requirement_db_id)
        flash('Auto-link failed for this requirement. Please try again.', 'error')
        return redirect(url_for('main.compliance_requirement_detail', requirement_db_id=int(requirement.id)))

    if int(result.get('links_added') or 0) > 0:
        flash(
            f"Auto-linked {int(result.get('links_added') or 0)} evidence item(s) for requirement {requirement.requirement_id}.",
            'success',
        )
    else:
        flash('No new auto-links found for this requirement. Try analyzing more documents first.', 'info')

    return redirect(url_for('main.compliance_requirement_detail', requirement_db_id=int(requirement.id)))


@bp.route('/compliance-requirements/<int:requirement_db_id>/link', methods=['POST'])
@login_required
def compliance_requirement_link_evidence(requirement_db_id):
    """Link a document to a requirement evidence bucket."""
    maybe = _require_active_org()
    if maybe is not None:
        return maybe

    org_id = _active_org_id()
    if not current_user.has_permission('documents.view', org_id=int(org_id)):
        abort(403)

    requirement = _get_org_visible_requirement_or_404(requirement_db_id=int(requirement_db_id), org_id=int(org_id))

    document_id = request.form.get('document_id', type=int)
    evidence_bucket = (request.form.get('evidence_bucket') or '').strip().lower()
    rationale_note = (request.form.get('rationale_note') or '').strip() or None

    allowed_buckets = {'system', 'implementation', 'workforce', 'participant'}

    document = Document.query.filter_by(
        id=document_id,
        organization_id=int(org_id),
        is_active=True,
    ).first()
    if not document:
        abort(404)

    if evidence_bucket in {'', 'auto'}:
        counts = _requirement_bucket_counts(org_id=int(org_id), requirement_id=int(requirement.id))
        evidence_bucket = _pick_bucket_for_requirement(
            requirement=requirement,
            bucket_counts=counts,
            preferred_bucket='',
        )
    elif evidence_bucket not in allowed_buckets:
        flash('Please choose a valid evidence bucket.', 'error')
        return redirect(url_for('main.compliance_requirement_detail', requirement_db_id=int(requirement.id)))

    existing = RequirementEvidenceLink.query.filter_by(
        organization_id=int(org_id),
        requirement_id=int(requirement.id),
        document_id=int(document.id),
        evidence_bucket=evidence_bucket,
    ).first()
    if existing:
        flash('This evidence link already exists.', 'info')
        return redirect(url_for('main.compliance_requirement_detail', requirement_db_id=int(requirement.id)))

    link = RequirementEvidenceLink(
        organization_id=int(org_id),
        requirement_id=int(requirement.id),
        document_id=int(document.id),
        evidence_bucket=evidence_bucket,
        rationale_note=rationale_note,
        linked_by_user_id=int(current_user.id),
    )
    db.session.add(link)
    compliance_scoring_service.recompute_requirement_assessment(
        organization_id=int(org_id),
        requirement_id=int(requirement.id),
        assessed_by_user_id=int(current_user.id),
        commit=True,
    )

    flash('Evidence linked to requirement.', 'success')
    return redirect(url_for('main.compliance_requirement_detail', requirement_db_id=int(requirement.id)))


@bp.route('/compliance-requirements/<int:requirement_db_id>/unlink/<int:link_id>', methods=['POST'])
@login_required
def compliance_requirement_unlink_evidence(requirement_db_id, link_id):
    """Remove a requirement evidence link."""
    maybe = _require_active_org()
    if maybe is not None:
        return maybe

    org_id = _active_org_id()
    if not current_user.has_permission('documents.view', org_id=int(org_id)):
        abort(403)

    requirement = _get_org_visible_requirement_or_404(requirement_db_id=int(requirement_db_id), org_id=int(org_id))

    link = RequirementEvidenceLink.query.filter_by(
        id=int(link_id),
        organization_id=int(org_id),
        requirement_id=int(requirement.id),
    ).first()
    if not link:
        abort(404)

    db.session.delete(link)
    compliance_scoring_service.recompute_requirement_assessment(
        organization_id=int(org_id),
        requirement_id=int(requirement.id),
        assessed_by_user_id=int(current_user.id),
        commit=True,
    )

    flash('Evidence link removed.', 'success')
    return redirect(url_for('main.compliance_requirement_detail', requirement_db_id=int(requirement.id)))


def _extract_demo_document_text(file_storage) -> tuple[str, str | None]:
    """Extract text from an uploaded file for demo-only analysis (in-memory)."""
    filename = (getattr(file_storage, 'filename', '') or '').strip().lower()
    if not filename:
        return '', 'A file is required.'

    try:
        raw_bytes = file_storage.read()
    except Exception:
        return '', 'Unable to read uploaded file.'

    if not raw_bytes:
        return '', 'Uploaded file is empty.'

    max_size = 8 * 1024 * 1024
    if len(raw_bytes) > max_size:
        return '', 'File is too large for demo analysis (max 8MB).'

    try:
        if filename.endswith('.txt'):
            text = raw_bytes.decode('utf-8', errors='ignore')
            return text.strip(), None

        if filename.endswith('.pdf'):
            from pypdf import PdfReader

            reader = PdfReader(io.BytesIO(raw_bytes))
            pages = []
            for page in reader.pages:
                pages.append((page.extract_text() or '').strip())
            return '\n\n'.join([p for p in pages if p]).strip(), None

        if filename.endswith('.docx'):
            from docx import Document as DocxDocument

            doc = DocxDocument(io.BytesIO(raw_bytes))
            paragraphs = [(p.text or '').strip() for p in doc.paragraphs]
            return '\n\n'.join([p for p in paragraphs if p]).strip(), None
    except Exception:
        return '', 'Unable to parse file. Use TXT, PDF, or DOCX for demo analysis.'

    return '', 'Unsupported file type. Use TXT, PDF, or DOCX.'


def _tokenize_demo_text(text: str) -> list[str]:
    import re

    stop_words = {
        'the', 'and', 'for', 'with', 'that', 'this', 'from', 'have', 'your', 'what', 'when', 'where',
        'which', 'into', 'about', 'they', 'them', 'their', 'will', 'would', 'should', 'there', 'here',
        'been', 'also', 'only', 'than', 'then', 'each', 'very', 'more', 'most', 'such', 'some', 'using',
        'does', 'did', 'not', 'are', 'was', 'were', 'is', 'it', 'to', 'of', 'in', 'on', 'by', 'as',
    }
    terms = [t for t in re.findall(r"[a-z0-9]+", (text or '').lower()) if len(t) >= 3 and t not in stop_words]
    unique = []
    seen = set()
    for t in terms:
        if t in seen:
            continue
        seen.add(t)
        unique.append(t)
    return unique[:30]


def _is_demo_noise_block(text: str) -> bool:
    import re

    value = (text or '').strip()
    if not value:
        return True

    lower = value.lower()
    noise_phrases = {
        'uncontrolled document',
        'document control',
        'sharepoint',
        'file path',
        'all rights reserved',
        'printed copies',
        'this document is uncontrolled',
        'authorised by',
        'approved by',
        'reviewed by',
        'controlled copy',
        'version history',
    }
    if any(phrase in lower for phrase in noise_phrases):
        return True

    if re.match(r'^(page\s+\d+|version\b|review date\b|effective date\b|document owner\b)', lower):
        return True

    if re.search(r'[a-z]:\\|/sites/|https?://|\.pdf\b|\.docx\b|\.xlsx\b', lower):
        return True

    alpha_chars = sum(1 for ch in value if ch.isalpha())
    separator_chars = sum(1 for ch in value if ch in '._/\\|')
    if alpha_chars < 20:
        return True
    if separator_chars >= 8 and separator_chars > max(6, alpha_chars // 3):
        return True

    return False


def _demo_candidate_blocks(document_text: str) -> list[str]:
    import re

    value = (document_text or '').replace('\r\n', '\n').strip()
    if not value:
        return []

    paragraphs = [part.strip() for part in re.split(r'\n\s*\n+', value) if (part or '').strip()]
    lines = [part.strip() for part in value.split('\n') if (part or '').strip()]

    candidates = []
    seen = set()
    for part in paragraphs + lines:
        cleaned = re.sub(r'\s+', ' ', (part or '').strip())
        key = cleaned.lower()
        if len(cleaned) < 30 or key in seen or _is_demo_noise_block(cleaned):
            continue
        seen.add(key)
        candidates.append(cleaned)
    return candidates


def _looks_like_demo_template(document_text: str) -> bool:
    import re

    lower = (document_text or '').lower()
    if 'template' in lower:
        return True

    placeholder_count = len(re.findall(r'\[[^\]]{2,40}\]|<[^>]{2,40}>', document_text or ''))
    return placeholder_count >= 3


def _rank_demo_snippets(document_text: str, query_text: str, top_k: int = 4) -> list[dict]:
    blocks = _demo_candidate_blocks(document_text)

    if not blocks:
        return []

    query_terms = _tokenize_demo_text(query_text)
    substantive_terms = {
        'participant', 'consent', 'complaint', 'complaints', 'feedback', 'privacy', 'dignity', 'rights',
        'agreement', 'support', 'supports', 'planning', 'assessment', 'culture', 'diversity', 'risk',
        'review', 'incident', 'governance', 'monitoring', 'confidentiality',
    }
    scored = []
    for block in blocks:
        lower = block.lower()
        score = 0.0
        matched_terms = 0
        for term in query_terms:
            if term in lower:
                score += 1.0
                matched_terms += 1
        if matched_terms <= 0:
            continue
        if any(term in lower for term in substantive_terms):
            score += 0.5
        if len(block) >= 140:
            score += 0.25
        if score > 0:
            scored.append((score, block))

    scored.sort(key=lambda item: item[0], reverse=True)
    selected = scored[:max(1, min(int(top_k or 4), 6))]
    return [
        {
            'score': max(1, int(round(score))),
            'text': (text[:420].rstrip() + '...') if len(text) > 420 else text,
        }
        for score, text in selected
    ]


def _derive_demo_status(document_text: str, query_text: str, snippets: list[dict], *, mode: str = 'balanced') -> tuple[str, float]:
    query_terms = _tokenize_demo_text(query_text)
    if not query_terms:
        query_terms = _tokenize_demo_text('compliance evidence policy process review audit monitoring')

    lower_text = (document_text or '').lower()
    doc_len = len((document_text or '').strip())
    snippet_count = len(snippets or [])
    hits = sum(1 for term in query_terms if term in lower_text)
    coverage = (float(hits) / float(len(query_terms))) if query_terms else 0.0
    score = (coverage * 0.65) + (min(float(snippet_count) / 4.0, 1.0) * 0.2) + (min(float(doc_len) / 1800.0, 1.0) * 0.15)
    confidence = max(0.0, min(score, 1.0))

    normalized_mode = (mode or 'balanced').strip().lower()
    if normalized_mode not in {'strict', 'balanced'}:
        normalized_mode = 'balanced'
    is_template_like = _looks_like_demo_template(document_text)

    # Strict mode is intentionally conservative to avoid over-claiming readiness.
    if normalized_mode == 'strict':
        # In strict mode, we intentionally deflate confidence to reflect higher audit caution.
        confidence = max(0.0, min(confidence * 0.90, 1.0))
        if doc_len < 220 or coverage < 0.12:
            return 'Critical gap', round(confidence, 3)
        if score >= 0.90 and snippet_count >= 4 and doc_len >= 900:
            return 'Mature', round(confidence, 3)
        if score >= 0.58 and snippet_count >= 2:
            return 'OK', round(confidence, 3)
        if score >= 0.30:
            return 'High risk gap', round(confidence, 3)
        return 'Critical gap', round(confidence, 3)

    # Balanced mode reduces false-critical outcomes for short but relevant policy docs.
    confidence = max(0.0, min(confidence * 1.05, 1.0))
    if doc_len < 120 and coverage < 0.20:
        return 'Critical gap', round(confidence, 3)
    if score >= 0.72 and snippet_count >= 3:
        return 'Mature', round(confidence, 3)
    if score >= 0.48 and snippet_count >= 2:
        return 'OK', round(confidence, 3)
    if is_template_like and coverage >= 0.16 and snippet_count >= 2:
        return 'OK', round(min(confidence, 0.58), 3)
    if score >= 0.22:
        return 'High risk gap', round(confidence, 3)
    return 'Critical gap', round(confidence, 3)


def _build_demo_rag_query(question_text: str, document_text: str) -> str:
    lower = f'{question_text or ""} {(document_text or "")[:2500]}'.lower()
    expansions = []
    topic_bundles = [
        ({'complaint', 'complaints', 'feedback', 'grievance'}, 'complaints management resolution participant feedback'),
        ({'consent', 'privacy', 'dignity', 'confidential', 'confidentiality'}, 'participant rights privacy dignity information management consent'),
        ({'service agreement', 'agreement', 'service terms', 'supports'}, 'service agreements with participants provision of supports responsive support provision'),
        ({'support plan', 'care plan', 'assessment', 'planning', 'goals'}, 'support planning participant needs preferences goals risk assessments'),
        ({'culture', 'cultural', 'diversity', 'inclusive', 'inclusion', 'beliefs', 'values'}, 'person centred supports individual values beliefs participant rights cultural inclusion'),
        ({'risk', 'governance', 'quality', 'incident', 'information management'}, 'governance operational management risk management quality management information management'),
    ]

    for keywords, expansion in topic_bundles:
        if any(keyword in lower for keyword in keywords):
            expansions.append(expansion)

    if not expansions:
        return question_text

    return ' '.join([question_text.strip()] + expansions).strip()


def _ensure_demo_summary_text(
    summary_text: str | None,
    *,
    status: str,
    analysis_mode: str,
    snippets: list[dict],
    citations: list[dict],
) -> str:
    value = (summary_text or '').replace('\r\n', '\n').strip()
    if value and _is_complete_demo_summary(value):
        return value
    if value:
        coerced = _coerce_demo_summary_text(value)
        if coerced and _is_complete_demo_summary(coerced):
            return coerced

    first_snippet = ((snippets or [{}])[0] or {}).get('text') if snippets else None
    first_citation = ((citations or [{}])[0] or {}) if citations else {}
    citation_ref = ''
    if first_citation:
        citation_ref = f"{first_citation.get('source_id', 'ndis')} p.{first_citation.get('page_number') or '?'}"

    parts = [
        f"Proposed status: {status} (mode: {analysis_mode}).",
        (f"Top supporting evidence: {first_snippet}" if first_snippet else 'No strong evidence snippet was extracted from the uploaded file.'),
        (f"Top NDIS citation: {citation_ref}." if citation_ref else 'No NDIS citation was retrieved for this run.'),
        'Recommendation: review missing controls and add clearer operational evidence (ownership, timelines, monitoring, and participant communication).',
    ]
    return '\n\n'.join(parts)


def _normalize_demo_summary_text(text: str | None) -> str:
    value = (text or '').replace('\r\n', '\n').strip()
    if not value:
        return ''

    # Remove common code-fence wrappers if the model returns markdown blocks.
    if value.startswith('```') and value.endswith('```'):
        lines = value.split('\n')
        if len(lines) >= 3:
            value = '\n'.join(lines[1:-1]).strip()

    if 'END_SUMMARY' in value:
        value = value.split('END_SUMMARY', 1)[0].strip()

    return value


def _is_complete_demo_summary(text: str | None) -> bool:
    value = _normalize_demo_summary_text(text)
    if len(value) < 80:
        return False

    lower = value.lower()
    has_why = ('1)' in lower and 'why this status' in lower) or ('why this status' in lower)
    has_missing = ('2)' in lower and 'missing evidence' in lower) or ('missing evidence' in lower)
    has_action = ('3)' in lower and 'recommended next action' in lower) or ('recommended next action' in lower)
    return bool(has_why and has_missing and has_action)


def _coerce_demo_summary_text(text: str | None) -> str | None:
    """Coerce partially structured LLM output into required 3-section demo format."""
    value = _normalize_demo_summary_text(text)
    if not value:
        return None
    if _is_complete_demo_summary(value):
        return value

    import re

    lines = [ln.strip(' -\t') for ln in value.split('\n') if ln.strip()]
    if not lines:
        return None

    def _as_sentence(fragment: str) -> str:
        cleaned = re.sub(r'\s+', ' ', (fragment or '').strip())
        cleaned = re.sub(r'[.?!]+$', '', cleaned).strip()
        if not cleaned:
            return ''
        return f'{cleaned}.'

    def _find_idx(patterns: list[str]) -> int | None:
        for i, ln in enumerate(lines):
            low = ln.lower()
            if any(re.search(p, low) for p in patterns):
                return i
        return None

    idx_why = _find_idx([r'why this status', r'rationale', r'reason'])
    idx_missing = _find_idx([r'missing evidence', r'gap', r'limitation', r'missing'])
    idx_action = _find_idx([r'recommended next action', r'next action', r'recommend', r'action'])

    def _slice(start: int | None, end: int | None) -> str:
        if start is None:
            return ''
        a = max(0, start + 1)
        b = end if end is not None else len(lines)
        return ' '.join(lines[a:b]).strip(' .')

    why = _slice(idx_why, idx_missing if idx_missing is not None else idx_action)
    missing = _slice(idx_missing, idx_action)
    action = _slice(idx_action, None)

    # If headings are missing, heuristically split by sentences.
    if not (why and missing and action):
        sentences = [s.strip() for s in re.split(r'(?<=[.!?])\s+', value) if s.strip()]
        if not why:
            why = ' '.join(sentences[:2]).strip(' .')
        if not missing:
            missing = ' '.join(sentences[2:4]).strip(' .')
        if not action:
            action = ' '.join(sentences[4:]).strip(' .')

    if not why:
        why = 'Status is based on detected evidence relevance and citation overlap.'
    if not missing:
        missing = 'Specific control ownership, timelines, and monitoring detail are not fully evidenced.'
    if not action:
        action = 'Add explicit control owners, due dates, review cadence, and participant communication pathways.'

    why = _as_sentence(why)
    missing = _as_sentence(missing)
    action = _as_sentence(action)

    coerced = (
        '1) Why this status\n'
        f'{why}\n\n'
        '2) Missing evidence\n'
        f'{missing}\n\n'
        '3) Recommended next action\n'
        f'{action}'
    )
    return _normalize_demo_summary_text(coerced)


def _azure_openai_demo_summary(*, status: str, question: str, snippets: list[dict], citations: list[dict]) -> tuple[str | None, str | None, str | None]:
    endpoint = (current_app.config.get('AZURE_OPENAI_ENDPOINT') or '').strip().rstrip('/')
    api_key = (current_app.config.get('AZURE_OPENAI_API_KEY') or '').strip()
    api_version = (current_app.config.get('AZURE_OPENAI_API_VERSION') or '2024-10-21').strip()
    timeout_seconds = int(current_app.config.get('AZURE_OPENAI_TIMEOUT_SECONDS') or 30)
    max_output_tokens = int(current_app.config.get('AZURE_OPENAI_SUMMARY_MAX_OUTPUT_TOKENS') or 450)

    deployment = (
        (current_app.config.get('AZURE_OPENAI_CHAT_DEPLOYMENT_MINI') or '').strip()
        or (current_app.config.get('AZURE_OPENAI_CHAT_DEPLOYMENT') or '').strip()
        or (current_app.config.get('AZURE_OPENAI_CHAT_DEPLOYMENT_WRITER') or '').strip()
    )

    if not endpoint or not api_key or not deployment:
        return None, 'Azure OpenAI is not fully configured. Using deterministic explanation.', None

    evidence_points = '\n'.join([f"- {item.get('text', '')}" for item in snippets[:4]]) or '- No strong document snippets found.'
    citation_points = '\n'.join([
        f"- {c.get('source_id', 'ndis')} p.{c.get('page_number') or '?'}: {c.get('text', '')}"
        for c in citations[:3]
    ]) or '- No NDIS citations retrieved.'

    prompt = (
        'You are an NDIS compliance demo assistant. '
        'Given the proposed status, document evidence snippets, and NDIS citations, produce a concise explanation. '
        'Return plain text only (no markdown, no bullet symbols, no code fences). '
        'Use this exact structure with all sections present:\n'
        '1) Why this status\n'
        '2) Missing evidence\n'
        '3) Recommended next action\n'
        'End with the line END_SUMMARY\n\n'
        f'Proposed status: {status}\n'
        f'Question: {question}\n\n'
        f'Document evidence snippets:\n{evidence_points}\n\n'
        f'NDIS citations:\n{citation_points}'
    )

    try:
        import requests

        url = f"{endpoint}/openai/deployments/{deployment}/chat/completions?api-version={api_version}"
        response = requests.post(
            url,
            headers={
                'Content-Type': 'application/json',
                'api-key': api_key,
            },
            json={
                'messages': [
                    {'role': 'system', 'content': 'You are precise and practical. Do not claim final legal compliance.'},
                    {'role': 'user', 'content': prompt},
                ],
                'temperature': 0.1,
                'max_tokens': max_output_tokens,
            },
            timeout=timeout_seconds,
        )

        if response.status_code >= 400:
            snippet = ((response.text or '').strip()[:140])
            return None, f'Azure OpenAI failed ({response.status_code}){": " + snippet if snippet else ""}. Using deterministic explanation.', None

        payload = response.json() if response.content else {}
        choices = payload.get('choices') or []
        if not choices:
            return None, 'Azure OpenAI returned no choices. Using deterministic explanation.', None

        message = ((choices[0] or {}).get('message') or {}).get('content') or ''
        answer = _normalize_demo_summary_text(message)
        if not answer:
            return None, 'Azure OpenAI returned empty content. Using deterministic explanation.', None
        if not _is_complete_demo_summary(answer):
            coerced = _coerce_demo_summary_text(answer)
            if coerced and _is_complete_demo_summary(coerced):
                return coerced, None, deployment
            return None, 'Azure OpenAI returned incomplete summary format. Using deterministic explanation.', None
        return answer, None, deployment
    except Exception as exc:
        current_app.logger.exception('Azure OpenAI request failed')
        reason = f'{type(exc).__name__}: {exc}'.strip(': ')
        return None, f'Azure OpenAI request failed ({reason}). Using deterministic explanation.', None


def _openrouter_demo_summary(*, status: str, question: str, snippets: list[dict], citations: list[dict]) -> tuple[str | None, str | None, str | None]:
    api_key = (current_app.config.get('OPENROUTER_API_KEY') or '').strip()
    configured_model = (current_app.config.get('OPENROUTER_MODEL') or '').strip() or 'mistralai/mistral-7b-instruct:free'
    fallback_models = [configured_model, 'openrouter/auto']
    token_budgets = [700, 350, 180]
    # Preserve order but remove duplicates.
    model_candidates = []
    seen = set()
    for model_name in fallback_models:
        key = (model_name or '').strip().lower()
        if not key or key in seen:
            continue
        seen.add(key)
        model_candidates.append(model_name)

    if not api_key:
        return None, 'OPENROUTER_API_KEY is not set. Using deterministic explanation.', None

    evidence_points = '\n'.join([f"- {item.get('text', '')}" for item in snippets[:4]]) or '- No strong document snippets found.'
    citation_points = '\n'.join([
        f"- {c.get('source_id', 'ndis')} p.{c.get('page_number') or '?'}: {c.get('text', '')}"
        for c in citations[:3]
    ]) or '- No NDIS citations retrieved.'

    prompt = (
        'You are an NDIS compliance demo assistant. '
        'Given the proposed status, document evidence snippets, and NDIS citations, produce a concise explanation. '
        'Return plain text only (no markdown, no bullet symbols, no code fences). '
        'Use this exact structure with all sections present:\n'
        '1) Why this status\n'
        '2) Missing evidence\n'
        '3) Recommended next action\n'
        'End with the line END_SUMMARY\n\n'
        f'Proposed status: {status}\n'
        f'Question: {question}\n\n'
        f'Document evidence snippets:\n{evidence_points}\n\n'
        f'NDIS citations:\n{citation_points}'
    )

    try:
        import requests

        last_warning = None
        attempt_messages = []
        retriable_statuses = {400, 402, 404, 408, 409, 425, 429, 500, 502, 503, 504}
        for model in model_candidates:
            for max_tokens in token_budgets:
                response = requests.post(
                    'https://openrouter.ai/api/v1/chat/completions',
                    headers={
                        'Authorization': f'Bearer {api_key}',
                        'Content-Type': 'application/json',
                        'HTTP-Referer': 'https://cenaris.local',
                        'X-Title': 'Cenaris AI Demo',
                    },
                    json={
                        'model': model,
                        'messages': [
                            {'role': 'system', 'content': 'You are precise and practical. Do not claim final legal compliance.'},
                            {'role': 'user', 'content': prompt},
                        ],
                        'temperature': 0.1,
                        'max_tokens': max_tokens,
                    },
                    timeout=20,
                )

                if response.status_code >= 400:
                    snippet = ((response.text or '').strip()[:140])
                    last_warning = f'OpenRouter model {model} failed ({response.status_code}){": " + snippet if snippet else ""}.'
                    attempt_messages.append(f'{model}:{max_tokens} -> {response.status_code}')
                    if response.status_code in retriable_statuses:
                        continue
                    return None, f'{last_warning} Using deterministic explanation.', None

                payload = response.json() if response.content else {}
                choices = payload.get('choices') or []
                if not choices:
                    last_warning = f'OpenRouter model {model} returned no choices.'
                    attempt_messages.append(f'{model}:{max_tokens} -> no choices')
                    continue

                message = ((choices[0] or {}).get('message') or {}).get('content') or ''
                answer = _normalize_demo_summary_text(message)
                if not answer:
                    last_warning = f'OpenRouter model {model} returned empty content.'
                    attempt_messages.append(f'{model}:{max_tokens} -> empty content')
                    continue
                if not _is_complete_demo_summary(answer):
                    coerced = _coerce_demo_summary_text(answer)
                    if coerced and _is_complete_demo_summary(coerced):
                        return coerced, None, model
                    last_warning = f'OpenRouter model {model} returned incomplete summary format.'
                    attempt_messages.append(f'{model}:{max_tokens} -> incomplete format')
                    continue
                return answer, None, model

        if last_warning:
            attempts = '; '.join(attempt_messages[-3:]) if attempt_messages else 'unknown'
            return None, f'OpenRouter summary unavailable ({attempts}). Using deterministic explanation.', None
        return None, 'OpenRouter call failed. Using deterministic explanation.', None
    except Exception as exc:
        current_app.logger.exception('OpenRouter request failed')
        reason = f'{type(exc).__name__}: {exc}'.strip(': ')
        return None, f'OpenRouter request failed ({reason}). Using deterministic explanation.', None


def _llm_demo_summary(*, status: str, question: str, snippets: list[dict], citations: list[dict]) -> tuple[str | None, str | None, str | None, str]:
    azure_enabled = bool(
        (current_app.config.get('AZURE_OPENAI_ENDPOINT') or '').strip()
        and (current_app.config.get('AZURE_OPENAI_API_KEY') or '').strip()
    )

    if azure_enabled:
        summary, warning, model = _azure_openai_demo_summary(
            status=status,
            question=question,
            snippets=snippets,
            citations=citations,
        )
        return summary, warning, model, 'azure-openai'

    summary, warning, model = _openrouter_demo_summary(
        status=status,
        question=question,
        snippets=snippets,
        citations=citations,
    )
    return summary, warning, model, 'openrouter'


def _assistant_is_org_admin(org_id: int) -> bool:
    try:
        membership = (
            OrganizationMembership.query
            .filter_by(organization_id=int(org_id), user_id=int(current_user.id), is_active=True)
            .first()
        )
    except Exception:
        return False
    role_name = ((membership.role if membership else '') or '').strip().lower()
    return role_name in {'admin', 'owner'}


def _assistant_display_name() -> str:
    first = ((getattr(current_user, 'first_name', None) or '')).strip()
    if first:
        return first
    full_name = ((getattr(current_user, 'full_name', None) or '')).strip()
    if full_name:
        return full_name.split()[0]
    email = ((getattr(current_user, 'email', None) or '')).strip()
    if email and '@' in email:
        return email.split('@', 1)[0]
    return ''


def _assistant_default_actions() -> list[dict]:
    return [
        {
            'id': 'open_requirements',
            'kind': 'navigate',
            'label': 'Open Requirements Workboard',
            'url': url_for('main.compliance_requirements'),
        },
        {
            'id': 'open_ai_review',
            'kind': 'navigate',
            'label': 'Open AI Review',
            'url': url_for('main.ai_demo'),
        },
        {
            'id': 'open_repository',
            'kind': 'navigate',
            'label': 'Open Evidence Repository',
            'url': url_for('main.evidence_repository'),
        },
        {
            'id': 'open_upload',
            'kind': 'navigate',
            'label': 'Open Upload Section',
            'url': url_for('main.dashboard', _anchor='upload-documents', open_upload=1),
        },
        {
            'id': 'open_dashboard',
            'kind': 'navigate',
            'label': 'Open Dashboard',
            'url': url_for('main.dashboard'),
        },
        {
            'id': 'open_profile',
            'kind': 'navigate',
            'label': 'Open My Profile',
            'url': url_for('main.profile'),
        },
    ]


def _assistant_is_action_intent(query_text: str) -> bool:
    lower = (query_text or '').strip().lower()
    if not lower:
        return False
    command_prefixes = (
        'open ',
        'show ',
        'go to ',
        'take me to ',
        'navigate to ',
    )
    if lower.startswith(command_prefixes):
        return True
    if ('mark' in lower and 'notification' in lower and 'read' in lower) or ('clear notifications' in lower):
        return True
    return False


def _assistant_feature_context_text() -> str:
    return (
        'Cenaris product capability map:\n'
        '- Dashboard: operational starting point with readiness overview and shortcuts.\n'
        '- Evidence Repository: upload, preview, tag, and manage evidence documents.\n'
        '- AI Review Workspace: analyze documents, produce summaries, snippets, and NDIS-style citations.\n'
        '- Requirements Workboard: map evidence to requirements, track statuses, and due-state.\n'
        '- Policy Studio: create policies from scratch or optional document context; supports Word export.\n'
        '- Analytics Dashboard: readiness trends and framework analytics visualizations.\n'
        '- Organisation Profile: organisation settings, billing, and plans preview navigation.\n'
        '- Notifications: admin-focused operational alerts with mark-as-read actions.\n'
        '\n'
        'Rules:\n'
        '- Answer Cenaris-related product/workflow questions in detail.\n'
        '- Do not invent features not listed or not inferable from this context.\n'
        '- If asked about non-Cenaris topics, politely redirect to Cenaris help scope.\n'
        '- Prefer practical, step-by-step guidance when user asks how-to questions.'
    )


def _assistant_generate_ai_reply(
    *,
    org_id: int,
    query_text: str,
    fallback_reply: str,
    user_display_name: str,
    doc_count: int,
    reviewed_doc_count: int,
    dashboard_summary: dict,
    bridge: dict,
) -> str | None:
    if not bool(current_app.config.get('ASSISTANT_CHAT_USE_LLM')):
        return None

    query = (query_text or '').strip()
    if not query:
        return None

    feature_map = _assistant_feature_context_text()
    compliance_rate = round(float((dashboard_summary or {}).get('compliance_rate') or 0), 1)
    linked_requirements = int((bridge or {}).get('linked_requirements') or 0)
    pending_assessment = int((bridge or {}).get('linked_requirements_pending_assessment') or 0)

    user_prompt = (
        f"User question: {query}\n\n"
        f"User display name: {user_display_name or 'Unknown'}\n"
        f"Org snapshot: uploads={int(doc_count)}, ai_reviewed={int(reviewed_doc_count)}, "
        f"compliance_rate={compliance_rate}%, linked_requirements={linked_requirements}, "
        f"pending_assessment={pending_assessment}.\n\n"
        f"Fallback guidance (if needed): {fallback_reply}\n\n"
        f"{feature_map}\n\n"
        "Response style:\n"
        "- Speak like a helpful human product specialist.\n"
        "- Use the user's first name once at the start when available.\n"
        "- Use short sections and concise bullets for readability.\n"
        "- Keep the response practical and action-oriented."
    )

    endpoint = (current_app.config.get('AZURE_OPENAI_ENDPOINT') or '').strip().rstrip('/')
    api_key = (current_app.config.get('AZURE_OPENAI_API_KEY') or '').strip()
    api_version = (current_app.config.get('AZURE_OPENAI_API_VERSION') or '2024-10-21').strip()
    deployment = (
        (current_app.config.get('AZURE_OPENAI_ASSISTANT_DEPLOYMENT') or '').strip()
        or (current_app.config.get('AZURE_OPENAI_CHAT_DEPLOYMENT_MINI') or '').strip()
        or (current_app.config.get('AZURE_OPENAI_CHAT_DEPLOYMENT') or '').strip()
    )
    timeout_seconds = int(current_app.config.get('AZURE_OPENAI_TIMEOUT_SECONDS') or 30)
    max_tokens = _clamp_int(
        current_app.config.get('ASSISTANT_CHAT_MAX_OUTPUT_TOKENS') or 550,
        default=550,
        minimum=180,
        maximum=1200,
    )
    temperature = float(current_app.config.get('ASSISTANT_CHAT_TEMPERATURE') or 0.2)

    started = time.perf_counter()

    # Prefer Azure OpenAI when configured.
    if endpoint and api_key and deployment:
        try:
            import requests

            url = f"{endpoint}/openai/deployments/{deployment}/chat/completions?api-version={api_version}"
            response = requests.post(
                url,
                headers={
                    'Content-Type': 'application/json',
                    'api-key': api_key,
                },
                json={
                    'messages': [
                        {
                            'role': 'system',
                            'content': 'You are Cenaris Compass, an in-product assistant for Cenaris only. Be accurate, detailed, and practical.',
                        },
                        {'role': 'user', 'content': user_prompt},
                    ],
                    'temperature': max(0.0, min(1.0, temperature)),
                    'max_tokens': int(max_tokens),
                },
                timeout=timeout_seconds,
            )

            if response.status_code < 400:
                payload = response.json() if response.content else {}
                choices = payload.get('choices') or []
                if choices:
                    reply = (((choices[0] or {}).get('message') or {}).get('content') or '').strip()
                    if reply:
                        usage = payload.get('usage') or {}
                        _log_ai_call(
                            'assistant_chat',
                            org_id=int(org_id),
                            mode='llm',
                            provider='azure-openai',
                            model=deployment,
                            usage={
                                'prompt_tokens': int(usage.get('prompt_tokens') or 0),
                                'completion_tokens': int(usage.get('completion_tokens') or 0),
                                'total_tokens': int(usage.get('total_tokens') or 0),
                            },
                            latency_ms=int((time.perf_counter() - started) * 1000),
                        )
                        return reply
        except Exception:
            current_app.logger.exception('Assistant Azure LLM request failed; falling back')

    # Optional fallback: OpenRouter when configured.
    openrouter_key = (current_app.config.get('OPENROUTER_API_KEY') or '').strip()
    openrouter_model = (current_app.config.get('OPENROUTER_ASSISTANT_MODEL') or '').strip() or 'openrouter/auto'
    if not openrouter_key:
        return None

    try:
        import requests

        response = requests.post(
            'https://openrouter.ai/api/v1/chat/completions',
            headers={
                'Authorization': f'Bearer {openrouter_key}',
                'Content-Type': 'application/json',
                'HTTP-Referer': 'https://cenaris.local',
                'X-Title': 'Cenaris Compass Assistant',
            },
            json={
                'model': openrouter_model,
                'messages': [
                    {
                        'role': 'system',
                        'content': 'You are Cenaris Compass, an in-product assistant for Cenaris only. Be accurate, detailed, and practical.',
                    },
                    {'role': 'user', 'content': user_prompt},
                ],
                'temperature': max(0.0, min(1.0, temperature)),
                'max_tokens': int(max_tokens),
            },
            timeout=timeout_seconds,
        )

        if response.status_code >= 400:
            return None
        payload = response.json() if response.content else {}
        choices = payload.get('choices') or []
        if not choices:
            return None
        reply = (((choices[0] or {}).get('message') or {}).get('content') or '').strip()
        if not reply:
            return None

        usage = payload.get('usage') or {}
        _log_ai_call(
            'assistant_chat',
            org_id=int(org_id),
            mode='llm',
            provider='openrouter',
            model=openrouter_model,
            usage={
                'prompt_tokens': int(usage.get('prompt_tokens') or 0),
                'completion_tokens': int(usage.get('completion_tokens') or 0),
                'total_tokens': int(usage.get('total_tokens') or 0),
            },
            latency_ms=int((time.perf_counter() - started) * 1000),
        )
        return reply
    except Exception:
        current_app.logger.exception('Assistant OpenRouter request failed; falling back')
        return None


def _assistant_compose_response(*, org_id: int, query_text: str) -> tuple[str, list[dict]]:
    lower = (query_text or '').strip().lower()
    if not lower:
        return (
            'Ask me things like: "Open requirements workboard", "Open AI review", or "Mark all notifications as read".',
            _assistant_default_actions(),
        )

    doc_count = (
        Document.query
        .filter(Document.organization_id == int(org_id), Document.is_active.is_(True))
        .count()
    )
    reviewed_doc_count = (
        Document.query
        .filter(
            Document.organization_id == int(org_id),
            Document.is_active.is_(True),
            Document.ai_analysis_at.isnot(None),
        )
        .count()
    )
    dashboard_summary = (analytics_service.build_dashboard_payload(organization_id=int(org_id)).get('summary') or {})
    bridge = _build_dashboard_bridge_stats(org_id=int(org_id))
    is_admin = _assistant_is_org_admin(int(org_id))

    if ('readiness' in lower) or ('overall readiness' in lower) or ('why' in lower and 'zero' in lower):
        return (
            (
                f'Readiness is requirement-based, not upload-only. You currently have {int(doc_count)} uploaded file(s), '
                f'{int(reviewed_doc_count)} AI-reviewed file(s), and a requirement compliance rate of '
                f'{round(float(dashboard_summary.get("compliance_rate") or 0), 1)}%. '
                f'Reviewed but not linked documents: {int(bridge.get("reviewed_docs_not_linked") or 0)}. '
                'AI Review already auto-maps likely requirement links. To increase readiness, verify/adjust those mappings and confirm requirement assessments.'
            ),
            [
                {
                    'id': 'open_requirements',
                    'kind': 'navigate',
                    'label': 'Open Requirements',
                    'url': url_for('main.compliance_requirements'),
                },
                {
                    'id': 'open_repository',
                    'kind': 'navigate',
                    'label': 'Open Evidence Repository',
                    'url': url_for('main.evidence_repository'),
                },
                {
                    'id': 'open_analytics',
                    'kind': 'navigate',
                    'label': 'Open Analytics',
                    'url': url_for('main.analytics_dashboard'),
                },
            ],
        )

    if (
        'journey' in lower
        or ('start' in lower and ('where' in lower or 'first' in lower or 'begin' in lower))
        or ('step by step' in lower)
        or ('full flow' in lower)
    ):
        return (
            (
                'Use Requirements as your main workflow: upload evidence -> run AI review -> verify links -> confirm assessment status -> export reports. '
                'The requirements workboard now includes a simple document-link view first, with advanced planning hidden by default.'
            ),
            [
                {
                    'id': 'open_requirements',
                    'kind': 'navigate',
                    'label': 'Open Requirements',
                    'url': url_for('main.compliance_requirements'),
                },
                {
                    'id': 'open_ai_review',
                    'kind': 'navigate',
                    'label': 'Open AI Review',
                    'url': url_for('main.ai_demo'),
                },
            ],
        )

    if 'deadline' in lower or 'due' in lower:
        return (
            (
                'Deadlines are auto-calculated per requirement: last assessed date + review frequency. '
                'A deadline appears only when a requirement has linked evidence and a completed assessment. '
                'If no completed assessment exists, no deadline is shown.'
            ),
            [
                {
                    'id': 'open_dashboard',
                    'kind': 'navigate',
                    'label': 'Open Dashboard',
                    'url': url_for('main.dashboard'),
                },
                {
                    'id': 'open_ai_review',
                    'kind': 'navigate',
                    'label': 'Open AI Review',
                    'url': url_for('main.ai_demo'),
                },
            ],
        )

    if (
        'link' in lower
        and ('evidence' in lower or 'requirement' in lower or 'them' in lower)
    ) or ('use of it' in lower) or ('how will it help' in lower):
        linked_count = int(bridge.get('linked_requirements') or 0)
        pending_count = int(bridge.get('linked_requirements_pending_assessment') or 0)
        return (
            (
                'In this app, AI Review already tries to auto-link each reviewed document to likely requirements. '
                'You still verify those links so the evidence is trustworthy for audit use. '
                'Why it helps: 1) readiness score moves from requirement evidence, 2) gaps become specific and actionable, '
                '3) exports and reviews become audit-ready.\n\n'
                f'Current org snapshot: {linked_count} requirement(s) already linked to evidence, '
                f'{pending_count} linked requirement(s) still waiting assessment update.\n\n'
                'Quick steps: open Evidence Repository -> open a reviewed document -> check mapped requirements -> keep/edit links -> review requirement status.'
            ),
            [
                {
                    'id': 'open_repository',
                    'kind': 'navigate',
                    'label': 'Open Evidence Repository',
                    'url': url_for('main.evidence_repository'),
                },
                {
                    'id': 'open_requirements',
                    'kind': 'navigate',
                    'label': 'Open Compliance Requirements',
                    'url': url_for('main.compliance_requirements'),
                },
                {
                    'id': 'open_dashboard',
                    'kind': 'navigate',
                    'label': 'Open Dashboard',
                    'url': url_for('main.dashboard'),
                },
            ],
        )

    if 'how' in lower and ('website' in lower or 'app' in lower or 'work' in lower):
        return (
            (
                'Quick app flow: 1) Upload evidence in Repository. 2) Run AI Review for each document. '
                '3) Link evidence to requirements. 4) Review requirement status and gaps. 5) Track trends in Analytics and export reports.'
            ),
            [
                {
                    'id': 'open_upload',
                    'kind': 'navigate',
                    'label': 'Open Upload Section',
                    'url': url_for('main.dashboard', _anchor='upload-documents', open_upload=1),
                },
                {
                    'id': 'open_ai_review',
                    'kind': 'navigate',
                    'label': 'Open AI Review',
                    'url': url_for('main.ai_demo'),
                },
                {
                    'id': 'open_analytics',
                    'kind': 'navigate',
                    'label': 'Open Analytics',
                    'url': url_for('main.analytics_dashboard'),
                },
            ],
        )

    if 'settings' in lower or 'organisation' in lower or 'organization' in lower:
        return (
            'Use Organisation Profile for business details and Team Management for role/member administration.',
            [
                {
                    'id': 'open_org_settings',
                    'kind': 'navigate',
                    'label': 'Open Organisation Profile',
                    'url': url_for('main.organization_settings'),
                },
                {
                    'id': 'open_team_mgmt',
                    'kind': 'navigate',
                    'label': 'Open Team Management',
                    'url': url_for('main.org_admin_dashboard'),
                },
            ],
        )

    if ('mark' in lower and 'notification' in lower and 'read' in lower) or ('clear notifications' in lower):
        if is_admin:
            return (
                'I can do that for you. Confirm the action below and I will mark all notifications as read.',
                [
                    {
                        'id': 'mark_all_notifications_read',
                        'kind': 'execute',
                        'label': 'Mark all notifications as read',
                        'confirm': 'Mark all unread notifications as read for your organisation?'
                    },
                    {
                        'id': 'open_notifications',
                        'kind': 'navigate',
                        'label': 'Open Notifications',
                        'url': url_for('main.notifications'),
                    },
                ],
            )
        return (
            'Viewing notifications is admin-only in this workspace. Ask an org admin to run this action, or open your dashboard instead.',
            [
                {
                    'id': 'open_dashboard',
                    'kind': 'navigate',
                    'label': 'Open Dashboard',
                    'url': url_for('main.dashboard'),
                }
            ],
        )

    if 'notification' in lower:
        if is_admin:
            return (
                'Open Notifications to review recent events and unread items. I can also mark all as read for you.',
                [
                    {
                        'id': 'open_notifications',
                        'kind': 'navigate',
                        'label': 'Open Notifications',
                        'url': url_for('main.notifications'),
                    },
                    {
                        'id': 'open_notifications_unread',
                        'kind': 'navigate',
                        'label': 'Open Unread Notifications',
                        'url': url_for('main.notifications', status='unread'),
                    },
                    {
                        'id': 'mark_all_notifications_read',
                        'kind': 'execute',
                        'label': 'Mark all notifications as read',
                        'confirm': 'Mark all unread notifications as read for your organisation?'
                    },
                ],
            )
        return (
            'Notifications are available to organisation admins in this app.',
            _assistant_default_actions(),
        )

    if 'ai review' in lower or 'analy' in lower or 'citation' in lower or 'rag' in lower:
        return (
            'Use AI Review to analyze repository documents and get evidence snippets plus NDIS citations. Pick a document, choose a question preset, then run analysis.',
            [
                {
                    'id': 'open_ai_review',
                    'kind': 'navigate',
                    'label': 'Open AI Review',
                    'url': url_for('main.ai_demo'),
                }
            ],
        )

    if 'upload' in lower or 'repository' in lower or 'document' in lower or 'evidence' in lower:
        return (
            f'You currently have {int(doc_count)} active repository document(s). Open Evidence Repository to upload, tag, preview, and launch AI Review from a document row.',
            [
                {
                    'id': 'open_repository',
                    'kind': 'navigate',
                    'label': 'Open Evidence Repository',
                    'url': url_for('main.evidence_repository'),
                },
                {
                    'id': 'open_upload',
                    'kind': 'navigate',
                    'label': 'Open Upload Section',
                    'url': url_for('main.dashboard', _anchor='upload-documents', open_upload=1),
                },
                {
                    'id': 'open_ai_review',
                    'kind': 'navigate',
                    'label': 'Open AI Review',
                    'url': url_for('main.ai_demo'),
                },
            ],
        )

    if (
        'profile' in lower
        or ('change' in lower and 'name' in lower)
        or ('update' in lower and 'name' in lower)
        or ('password' in lower)
        or ('my name' in lower)
    ):
        return (
            'For name updates, open My Profile and edit your first/last name, then save. For password changes, use the forgot-password flow from the sign-in page if you need to reset credentials.',
            [
                {
                    'id': 'open_profile',
                    'kind': 'navigate',
                    'label': 'Open My Profile',
                    'url': url_for('main.profile'),
                },
                {
                    'id': 'open_help',
                    'kind': 'navigate',
                    'label': 'Open Help',
                    'url': url_for('main.help'),
                },
            ],
        )

    if 'analytics' in lower or 'report' in lower or 'trend' in lower:
        return (
            'Open Analytics to view readiness trends, requirement status, and downloadable reports.',
            [
                {
                    'id': 'open_analytics',
                    'kind': 'navigate',
                    'label': 'Open Analytics',
                    'url': url_for('main.analytics_dashboard'),
                }
            ],
        )

    if 'dashboard' in lower or 'home' in lower:
        return (
            'Dashboard is the best place to start. From there you can upload files, review compliance progress, and jump into AI tools.',
            [
                {
                    'id': 'open_dashboard',
                    'kind': 'navigate',
                    'label': 'Open Dashboard',
                    'url': url_for('main.dashboard'),
                }
            ],
        )

    return (
        'I can help with app navigation and simple admin automations. Try: "Open AI Review", "How do I upload evidence?", "Show analytics", or "Mark all notifications as read".',
        _assistant_default_actions(),
    )


@bp.route('/api/assistant/chat', methods=['POST'])
@login_required
@limiter.limit('20 per minute', key_func=_ai_rate_limit_key)
def assistant_chat_api():
    """Simple in-app assistant for product Q&A and low-risk automations."""
    maybe = _require_active_org()
    if maybe is not None:
        return jsonify({'success': False, 'error': 'No active organization'}), 400

    org_id = _active_org_id()
    payload = request.get_json(silent=True) or {}
    message = _limit_text((payload.get('message') or '').strip(), max_chars=800)
    action_id = (payload.get('action_id') or '').strip()
    execute_flag = str(payload.get('execute') or '').strip().lower() in {'1', 'true', 'yes', 'on'}

    # Allow execution calls with action_id only.
    if execute_flag and action_id == 'mark_all_notifications_read':
        if not _assistant_is_org_admin(int(org_id)):
            return jsonify(
                {
                    'success': True,
                    'reply': 'I can only run this action for organisation admins.',
                    'actions': [
                        {
                            'id': 'open_dashboard',
                            'kind': 'navigate',
                            'label': 'Open Dashboard',
                            'url': url_for('main.dashboard'),
                        }
                    ],
                    'executed_action': {'id': action_id, 'success': False},
                }
            )

        marked = notification_service.mark_all_read(organization_id=int(org_id), user_id=int(current_user.id))
        invalidate_org_switcher_context_cache(int(current_user.id), int(org_id))
        return jsonify(
            {
                'success': True,
                'reply': f'Done. I marked {int(marked)} notification(s) as read.',
                'actions': [
                    {
                        'id': 'open_notifications',
                        'kind': 'navigate',
                        'label': 'Open Notifications',
                        'url': url_for('main.notifications'),
                    }
                ],
                'executed_action': {'id': action_id, 'success': True, 'count': int(marked)},
            }
        )

    reply, actions = _assistant_compose_response(org_id=int(org_id), query_text=message)
    assistant_mode = 'rules'

    if message and not _assistant_is_action_intent(message):
        try:
            display_name = _assistant_display_name()
            doc_count = (
                Document.query
                .filter(Document.organization_id == int(org_id), Document.is_active.is_(True))
                .count()
            )
            reviewed_doc_count = (
                Document.query
                .filter(
                    Document.organization_id == int(org_id),
                    Document.is_active.is_(True),
                    Document.ai_analysis_at.isnot(None),
                )
                .count()
            )
            dashboard_summary = (analytics_service.build_dashboard_payload(organization_id=int(org_id)).get('summary') or {})
            bridge = _build_dashboard_bridge_stats(org_id=int(org_id))
            ai_reply = _assistant_generate_ai_reply(
                org_id=int(org_id),
                query_text=message,
                fallback_reply=reply,
                user_display_name=display_name,
                doc_count=int(doc_count),
                reviewed_doc_count=int(reviewed_doc_count),
                dashboard_summary=dashboard_summary,
                bridge=bridge,
            )
            if ai_reply:
                reply = _limit_text(ai_reply, max_chars=3000)
                assistant_mode = 'llm'
        except Exception:
            current_app.logger.exception('Assistant AI enrich failed; using rules reply')

    return jsonify(
        {
            'success': True,
            'reply': reply,
            'actions': actions,
            'executed_action': None,
            'assistant_mode': assistant_mode,
        }
    )


@bp.route('/ai-demo')
@login_required
def ai_demo():
    """Dedicated AI review workspace for repository documents."""
    maybe = _require_active_org()
    if maybe is not None:
        return maybe

    org_id = _active_org_id()
    if not current_user.has_permission('documents.view', org_id=int(org_id)):
        abort(403)

    requested_ids = [int(v) for v in request.args.getlist('doc_ids') if str(v).isdigit()]
    selected_documents = (
        Document.query
        .filter(
            Document.organization_id == int(org_id),
            Document.is_active.is_(True),
            Document.id.in_(requested_ids),
        )
        .order_by(Document.uploaded_at.desc())
        .all()
        if requested_ids
        else []
    )

    if not selected_documents:
        selected_documents = (
            Document.query
            .filter(
                Document.organization_id == int(org_id),
                Document.is_active.is_(True),
            )
            .order_by(Document.uploaded_at.desc())
            .limit(50)
            .all()
        )

    selected_document_ids = [int(document.id) for document in selected_documents]
    selected_document_meta: dict[str, dict] = {}
    for document in selected_documents:
        analyzed = bool(
            getattr(document, 'ai_analysis_at', None)
            and (
                ' '.join((getattr(document, 'ai_summary', '') or '').split()).strip()
                or ' '.join((getattr(document, 'ai_status', '') or '').split()).strip()
            )
        )
        preview_checklist = _build_checklist_from_analysis(
            {
                'status': document.ai_status,
                'summary': document.ai_summary,
                'focus_area': document.ai_focus_area or 'General compliance coverage',
            }
        ) if analyzed else {'items': [], 'summary': {'clear': [], 'partial': [], 'gap': []}}

        analyzed_at = getattr(document, 'ai_analysis_at', None)
        if analyzed_at and analyzed_at.tzinfo is None:
            analyzed_at = analyzed_at.replace(tzinfo=timezone.utc)

        selected_document_meta[str(int(document.id))] = {
            'filename': (document.filename or '').strip(),
            'analyzed': bool(analyzed),
            'status': (document.ai_status or '').strip(),
            'confidence': float(document.ai_confidence or 0),
            'question': (document.ai_question or '').strip(),
            'summary': (document.ai_summary or '').strip(),
            'provider': (document.ai_provider or '').strip(),
            'model': (document.ai_model or '').strip(),
            'analysis_at': analyzed_at.isoformat() if analyzed_at else '',
            'checklist': preview_checklist,
        }

    active_doc_id_raw = (request.args.get('active_doc_id') or '').strip()
    active_doc_id = int(active_doc_id_raw) if active_doc_id_raw.isdigit() else None
    if active_doc_id not in selected_document_ids:
        active_doc_id = selected_document_ids[0] if selected_document_ids else None

    autostart = (request.args.get('autostart') or '').strip().lower() in {'1', 'true', 'yes'}

    corpus_path = current_app.config.get('NDIS_RAG_CORPUS_PATH') or 'data/rag/ndis/ndis_chunks.jsonl'
    corpus_abs = os.path.abspath(os.path.join(current_app.root_path, os.pardir, corpus_path))
    try:
        from app.models import DemoAnalysisResult
        raw_recent_analyses = (
            DemoAnalysisResult.query
            .filter_by(organization_id=int(org_id))
            .order_by(DemoAnalysisResult.created_at.desc())
            .limit(40)
            .all()
        )
        seen_filenames: set[str] = set()
        recent_analyses = []
        for item in raw_recent_analyses:
            filename_key = (item.filename or '').strip().lower() or f'id-{int(item.id)}'
            if filename_key in seen_filenames:
                continue
            seen_filenames.add(filename_key)
            recent_analyses.append(item)
            if len(recent_analyses) >= 10:
                break
    except Exception:
        recent_analyses = []
    return render_template(
        'main/ai_demo.html',
        title='AI Review Workspace',
        has_ndis_corpus=bool(os.path.exists(corpus_abs)),
        recent_analyses=recent_analyses,
        selected_documents=selected_documents,
        selected_document_meta=selected_document_meta,
        selected_document_ids=selected_document_ids,
        active_doc_id=active_doc_id,
        autostart=autostart,
    )


@bp.route('/policy-studio')
@login_required
def policy_studio():
    """Standalone policy drafting studio with document-wide generation support."""
    maybe = _require_active_org()
    if maybe is not None:
        return maybe

    org_id = _active_org_id()
    if not current_user.has_permission('documents.view', org_id=int(org_id)):
        abort(403)

    documents = (
        Document.query
        .filter(
            Document.organization_id == int(org_id),
            Document.is_active.is_(True),
        )
        .order_by(Document.uploaded_at.desc())
        .limit(200)
        .all()
    )

    requirements = (
        ComplianceRequirement.query
        .join(ComplianceFrameworkVersion, ComplianceFrameworkVersion.id == ComplianceRequirement.framework_version_id)
        .filter(
            ComplianceFrameworkVersion.is_active.is_(True),
            or_(
                ComplianceFrameworkVersion.organization_id.is_(None),
                ComplianceFrameworkVersion.organization_id == int(org_id),
            ),
        )
        .order_by(
            ComplianceRequirement.requirement_id.asc(),
            ComplianceRequirement.quality_indicator_code.asc(),
            ComplianceRequirement.id.asc(),
        )
        .limit(500)
        .all()
    )
    requirement_options = [
        {
            'value': _requirement_display_code(requirement),
            'label': f"{_requirement_display_code(requirement)} - {_requirement_plain_title(requirement, max_chars=90)}",
        }
        for requirement in requirements
    ]

    prefill_scope = (request.args.get('scope') or 'linked').strip().lower()
    if prefill_scope not in {'linked', 'single'}:
        prefill_scope = 'linked'

    prefill_requirement = (request.args.get('requirement_id') or '').strip()
    known_requirement_values = {item['value'] for item in requirement_options}
    if prefill_requirement not in known_requirement_values:
        prefill_requirement = ''

    prefill_document_id_raw = (request.args.get('document_id') or '').strip()
    prefill_document_id = int(prefill_document_id_raw) if prefill_document_id_raw.isdigit() else None
    if prefill_document_id is not None and all(int(doc.id) != int(prefill_document_id) for doc in documents):
        prefill_document_id = None

    prefill_source = (request.args.get('source') or '').strip().lower()
    if prefill_source not in {'scratch', 'document'}:
        prefill_source = 'document' if prefill_document_id is not None else 'scratch'

    return render_template(
        'main/policy_studio.html',
        title='Policy Studio',
        documents=documents,
        requirement_options=requirement_options,
        prefill_scope=prefill_scope,
        prefill_requirement=prefill_requirement,
        prefill_document_id=prefill_document_id,
        prefill_source=prefill_source,
    )


@bp.route('/plans-preview')
@login_required
def plans_preview():
    """Client-facing feature matrix preview for plan discussions."""
    maybe = _require_active_org()
    if maybe is not None:
        return maybe
    organization = db.session.get(Organization, int(_active_org_id()))
    billing_state = billing_service.resolve_entitlements(
        organization,
        actor_email=getattr(current_user, 'email', None),
    ) if organization else None
    billing_catalog = billing_service.plan_catalog()
    is_super_admin = bool(_is_super_admin_user())
    return render_template(
        'main/plans_preview.html',
        title='Plans Preview',
        billing_state=billing_state,
        billing_catalog=billing_catalog,
        is_super_admin=is_super_admin,
    )


@bp.route('/plans-preview/select-plan', methods=['POST'])
@login_required
def plans_preview_select_plan():
    maybe = _require_org_permission('org.manage')
    if maybe is not None:
        return maybe

    if not _is_super_admin_user():
        abort(403)

    org_id = _active_org_id()
    organization = db.session.get(Organization, int(org_id)) if org_id else None
    if not organization:
        abort(404)

    selected_plan = billing_service.normalize_plan_code(request.form.get('plan_code'))
    organization.billing_plan_code = selected_plan
    organization.subscription_tier = selected_plan.capitalize()
    # Test-mode selection should allow immediate in-app gating checks without live billing.
    organization.billing_status = 'active'
    organization.billing_internal_override = True
    organization.billing_override_reason = 'Manual plan selection for testing (super admin)'
    try:
        db.session.commit()
        invalidate_org_switcher_context_cache(int(current_user.id), int(org_id))
        flash(f'Test plan set to {selected_plan.capitalize()}. Feature access has been updated.', 'success')
    except Exception:
        db.session.rollback()
        flash('Could not update test plan. Please try again.', 'error')

    return redirect(url_for('main.plans_preview'))


@bp.route('/billing/checkout', methods=['POST'])
@login_required
def billing_checkout():
    maybe = _require_org_permission('org.manage')
    if maybe is not None:
        return maybe

    org_id = _active_org_id()
    organization = db.session.get(Organization, int(org_id)) if org_id else None
    if not organization:
        flash('Organisation not found.', 'error')
        return redirect(url_for('main.organization_settings'))

    plan_code = (request.form.get('plan_code') or '').strip().lower()
    try:
        checkout_url = billing_service.create_checkout_session(organization, plan_code=plan_code)
        return redirect(checkout_url)
    except Exception as exc:
        current_app.logger.exception('Unable to create Stripe checkout session for org %s', int(organization.id))
        flash(f'Could not start checkout: {exc}', 'error')
        return redirect(url_for('main.organization_settings'))


@bp.route('/billing/portal', methods=['POST'])
@login_required
def billing_portal():
    maybe = _require_org_permission('org.manage')
    if maybe is not None:
        return maybe

    org_id = _active_org_id()
    organization = db.session.get(Organization, int(org_id)) if org_id else None
    if not organization:
        flash('Organisation not found.', 'error')
        return redirect(url_for('main.organization_settings'))

    try:
        portal_url = billing_service.create_portal_session(organization)
        return redirect(portal_url)
    except Exception as exc:
        current_app.logger.exception('Unable to create Stripe portal session for org %s', int(organization.id))
        flash(f'Could not open billing portal: {exc}', 'error')
        return redirect(url_for('main.organization_settings'))


@bp.route('/billing/webhook', methods=['POST'])
def billing_webhook():
    payload = request.get_data(cache=False)
    signature = request.headers.get('Stripe-Signature')

    try:
        event = billing_service.verify_webhook(payload, signature)
        billing_service.apply_webhook_event(event)
        return jsonify({'received': True}), 200
    except ValueError as exc:
        current_app.logger.warning('Stripe webhook rejected: %s', exc)
        return jsonify({'received': False, 'error': str(exc)}), 400
    except Exception:
        current_app.logger.exception('Stripe webhook processing failed')
        return jsonify({'received': False, 'error': 'Webhook processing failed.'}), 500


@bp.route('/api/ai/demo/analyze', methods=['POST'])
@login_required
@limiter.limit('8 per minute', key_func=_ai_rate_limit_key)
def ai_demo_analyze_api():
    """Analyze a stored repository document in the AI workspace."""
    maybe = _require_active_org()
    if maybe is not None:
        return jsonify({'success': False, 'error': 'No active organization'}), 400

    maybe_plan = _require_plan_feature('ai_tagging')
    if maybe_plan is not None:
        return jsonify({'success': False, 'error': 'This action requires Scale plan or above.'}), 403

    org_id = _active_org_id()
    if not current_user.has_permission('documents.view', org_id=int(org_id)):
        return jsonify({'success': False, 'error': 'Forbidden'}), 403

    stored_doc_id = (request.form.get('stored_doc_id') or '').strip()
    question = _limit_text((request.form.get('question') or '').strip(), max_chars=700)
    # Demo mode is intentionally fixed to balanced for consistent client-facing behavior.
    analysis_mode = 'balanced'
    # Checkbox semantics: explicit truthy value means reuse previous analysis.
    # Unchecked checkboxes are omitted from form payload and must resolve to False.
    reuse_last_raw = str(request.form.get('reuse_last') or '').strip().lower()
    reuse_last = reuse_last_raw in {'1', 'true', 'yes', 'on'}

    if not question:
        question = 'Assess this document against NDIS-style compliance evidence expectations.'

    if not stored_doc_id or not stored_doc_id.isdigit():
        return jsonify({'success': False, 'error': 'Choose a repository document first.'}), 400

    document = _authorized_org_document_or_404(int(stored_doc_id))
    source_filename = (document.filename or '').strip()
    has_previous_analysis = bool(
        getattr(document, 'ai_analysis_at', None)
        and (
            ' '.join((getattr(document, 'ai_summary', '') or '').split()).strip()
            or ' '.join((getattr(document, 'ai_status', '') or '').split()).strip()
        )
    )
    force_reanalyze = bool(has_previous_analysis and not reuse_last)

    normalized_question = ' '.join((question or '').split()).lower()
    cached_question = ' '.join(((document.ai_question or '')).split()).lower()
    if (
        reuse_last
        and not force_reanalyze
        and normalized_question
        and normalized_question == cached_question
        and (document.ai_summary or '').strip()
        and (document.ai_status or '').strip()
    ):
        reused_warning = 'Reused the latest stored analysis for this document and question.'
        checklist = _build_checklist_from_analysis(
            {
                'status': document.ai_status,
                'summary': document.ai_summary,
                'focus_area': document.ai_focus_area or 'General compliance coverage',
            }
        )
        return jsonify(
            {
                'success': True,
                'status': document.ai_status,
                'confidence': float(document.ai_confidence or 0),
                'summary': document.ai_summary,
                'checklist': checklist,
                'snippets': [],
                'citations': [],
                'warnings': [reused_warning],
                'warning_items': [{'source': 'cache', 'message': reused_warning}],
                'meta': {
                    'provider': (document.ai_provider or 'cached').strip() or 'cached',
                    'model': (document.ai_model or 'cached').strip() or 'cached',
                    'analysis_mode': analysis_mode,
                    'scoring_version': 'demo-v3',
                    'retrieval_mode': (document.ai_retrieval_mode or 'cached').strip() or 'cached',
                    'document_chars': len(document.extracted_text or ''),
                    'temporary_processing_only': False,
                    'filename': source_filename,
                    'source': 'repository',
                    'stored_doc_id': int(stored_doc_id) if stored_doc_id.isdigit() else None,
                    'cached_result': True,
                    'question': (document.ai_question or '').strip(),
                    'analysis_at': document.ai_analysis_at.isoformat() if document.ai_analysis_at else '',
                },
            }
        )

    from app.services.azure_storage import AzureBlobStorageService

    storage_service = AzureBlobStorageService()
    result = storage_service.download_file(document.blob_name)
    if not result.get('success') or not result.get('data'):
        return jsonify({'success': False, 'error': 'Could not load stored document for AI review.'}), 400

    doc_text, extraction_error = document_analysis_service.extract_text_from_bytes(source_filename, result.get('data') or b'')
    if extraction_error:
        return jsonify({'success': False, 'error': extraction_error}), 400

    snippets = _rank_demo_snippets(doc_text, question, top_k=4)
    status, confidence = _derive_demo_status(doc_text, question, snippets, mode=analysis_mode)

    rag_citations = []
    rag_warning = None
    warning_items: list[dict] = []
    _demo_retrieval_mode = 'lexical'
    rag_query_text = _build_demo_rag_query(question, doc_text)
    corpus_path = current_app.config.get('NDIS_RAG_CORPUS_PATH') or 'data/rag/ndis/ndis_chunks.jsonl'
    corpus_abs = os.path.abspath(os.path.join(current_app.root_path, os.pardir, corpus_path))
    try:
        rag_result = rag_query_service.query(corpus_path=corpus_abs, query_text=rag_query_text, requirement_id='', top_k=3)
        _demo_retrieval_mode = getattr(rag_result, 'retrieval_mode', 'lexical')
        rag_citations = [
            {
                'chunk_id': c.chunk_id,
                'source_id': c.source_id,
                'page_number': c.page_number,
                'score': c.score,
                'text': c.text,
            }
            for c in rag_result.citations
        ]
        if _demo_retrieval_mode == 'lexical' and rag_citations:
            warning_items.append({'source': 'rag', 'message': 'Semantic embeddings computing on first run — using keyword matching now. Reload to use hybrid mode.'})
    except FileNotFoundError:
        rag_warning = 'NDIS corpus is not built yet; showing document-only analysis.'
        warning_items.append({'source': 'rag', 'message': rag_warning})
    except Exception:
        rag_warning = 'Could not retrieve NDIS citations; showing document-only analysis.'
        warning_items.append({'source': 'rag', 'message': rag_warning})

    ai_summary, llm_warning, used_model, llm_provider = _llm_demo_summary(
        status=status,
        question=question,
        snippets=snippets,
        citations=rag_citations,
    )
    if llm_warning:
        warning_items.append({'source': 'llm', 'message': llm_warning})

    ai_summary = _ensure_demo_summary_text(
        ai_summary,
        status=status,
        analysis_mode=analysis_mode,
        snippets=snippets,
        citations=rag_citations,
    )

    provider_label = llm_provider if not llm_warning else 'deterministic'
    if llm_provider == 'azure-openai':
        model_label = used_model or (current_app.config.get('AZURE_OPENAI_CHAT_DEPLOYMENT_MINI') or current_app.config.get('AZURE_OPENAI_CHAT_DEPLOYMENT') or current_app.config.get('AZURE_OPENAI_CHAT_DEPLOYMENT_WRITER') or 'azure-openai')
    else:
        model_label = used_model or current_app.config.get('OPENROUTER_MODEL') or 'mistralai/mistral-7b-instruct:free'

    # Persist result (non-blocking — never fail the API on a DB error)
    analyzed_at = datetime.now(timezone.utc)
    try:
        from app.models import DemoAnalysisResult

        document.extracted_text = doc_text or None
        _refresh_document_search_text(document)
        document.ai_status = status
        document.ai_confidence = confidence
        document.ai_focus_area = 'General compliance coverage'
        document.ai_question = question
        document.ai_summary = ai_summary
        document.ai_provider = provider_label
        document.ai_model = model_label
        document.ai_retrieval_mode = _demo_retrieval_mode
        document.ai_analysis_at = analyzed_at

        record = DemoAnalysisResult(
            organization_id=int(org_id),
            user_id=int(current_user.id),
            filename=(source_filename or '')[:255] or None,
            question=(question or '')[:700] or None,
            status=status,
            confidence=confidence,
            analysis_mode=analysis_mode,
            summary=ai_summary,
            snippet_count=len(snippets),
            citation_count=len(rag_citations),
            provider=provider_label,
            model_used=(model_label or '')[:120] or None,
            retrieval_mode=_demo_retrieval_mode,
        )
        db.session.add(record)
        db.session.commit()
    except Exception:
        db.session.rollback()
        current_app.logger.exception('Failed to persist demo analysis result')

    matched_requirements = document_analysis_service._match_requirements(
        text=doc_text,
        filename=source_filename,
        organization_id=int(org_id),
    )
    checklist = _build_checklist_from_analysis(
        {
            'matched_requirements': matched_requirements,
            'status': status,
            'summary': ai_summary,
            'focus_area': 'General compliance coverage',
        }
    )

    warnings = [w.get('message', '') for w in warning_items if (w.get('message') or '').strip()]
    return jsonify(
        {
            'success': True,
            'status': status,
            'confidence': confidence,
            'summary': ai_summary,
            'checklist': checklist,
            'snippets': snippets,
            'citations': rag_citations,
            'warnings': warnings,
            'warning_items': warning_items,
            'meta': {
                'provider': provider_label,
                'model': model_label,
                'analysis_mode': analysis_mode,
                'scoring_version': 'demo-v3',
                'retrieval_mode': _demo_retrieval_mode,
                'document_chars': len(doc_text or ''),
                'temporary_processing_only': False,
                'filename': source_filename,
                'source': 'repository',
                'stored_doc_id': int(stored_doc_id) if stored_doc_id.isdigit() else None,
                'cached_result': False,
                'question': question,
                'analysis_at': analyzed_at.isoformat(),
            },
        }
    )


def _checklist_status_from_score(score: float | None) -> str:
    value = float(score or 0)
    if value >= 0.28:
        return 'Clear'
    if value >= 0.18:
        return 'Partial'
    return 'Gap'


def _checklist_status_from_overall(status: str | None) -> str:
    value = (status or '').strip().lower()
    if value in {'mature', 'ok'}:
        return 'Clear'
    if 'critical' in value or 'high risk' in value:
        return 'Gap'
    return 'Partial'


def _build_checklist_from_analysis(analysis: dict) -> dict:
    def _compact_text(value: str, *, max_chars: int = 140) -> str:
        text = ' '.join((value or '').split()).strip()
        if not text:
            return ''
        sentence = text.split('. ')[0].strip() or text
        if len(sentence) > max_chars:
            sentence = sentence[: max_chars - 1].rstrip() + '...'
        return sentence

    def format_note(*, status: str, missing: str, rationale: str) -> str:
        missing_value = (missing or '').strip()
        rationale_value = (rationale or '').strip()
        if status == 'Clear':
            if rationale_value:
                return f"Covered: {_compact_text(rationale_value, max_chars=130)}"
            return 'Covered by clear document evidence.'
        if missing_value:
            return f"Missing: {_compact_text(missing_value, max_chars=130)}"
        if rationale_value:
            return f"Partially addressed: {_compact_text(rationale_value, max_chars=130)}"
        return 'Needs clearer evidence in this document.'

    items = []
    matched = analysis.get('matched_requirements') or []
    for item in matched:
        score = item.get('score')
        status = _checklist_status_from_score(score)
        label = (item.get('label') or item.get('requirement_id') or 'Requirement').strip()
        note = format_note(
            status=status,
            missing=(item.get('missing_evidence') or ''),
            rationale=(item.get('rationale_note') or ''),
        )
        items.append(
            {
                'label': label,
                'status': status,
                'note': note,
                'score': score,
                'requirement_id': item.get('requirement_id'),
            }
        )

    default_areas = [
        ('Policy scope and intent', 'Define what this policy covers and where it applies.'),
        ('Operational evidence records', 'Provide logs, registers, and implementation proof.'),
        ('Review cadence and ownership', 'Name owners and confirm review and approval dates.'),
    ]

    def add_fallback_items(*, base_status: str) -> None:
        existing_labels = {' '.join((item.get('label') or '').lower().split()) for item in items}
        for area_label, area_note in default_areas:
            if len(items) >= 3:
                break
            normalized = ' '.join(area_label.lower().split())
            if normalized in existing_labels:
                continue
            items.append(
                {
                    'label': area_label,
                    'status': base_status,
                    'note': area_note,
                    'score': None,
                    'requirement_id': None,
                }
            )
            existing_labels.add(normalized)

    if not items:
        overall_status = _checklist_status_from_overall(analysis.get('status'))
        label = (analysis.get('focus_area') or 'General compliance coverage').strip()
        note = _compact_text(
            (analysis.get('summary') or 'Use more detailed evidence to improve assessment coverage.').strip(),
            max_chars=150,
        )
        items.append(
            {
                'label': label,
                'status': overall_status,
                'note': note,
                'score': None,
                'requirement_id': None,
            }
        )
        add_fallback_items(base_status=overall_status)
    elif len(items) < 3:
        add_fallback_items(base_status=_checklist_status_from_overall(analysis.get('status')))

    summary = {'clear': [], 'partial': [], 'gap': []}
    for item in items:
        status_key = (item.get('status') or '').lower()
        if status_key == 'clear':
            summary['clear'].append(item.get('label') or '')
        elif status_key == 'partial':
            summary['partial'].append(item.get('label') or '')
        else:
            summary['gap'].append(item.get('label') or '')

    return {
        'items': items,
        'summary': summary,
        'thresholds': {
            'clear_min_score': 0.28,
            'partial_min_score': 0.18,
        },
    }


@bp.route('/ai-summary-test')
@login_required
def ai_summary_test():
    """Legacy route retained for compatibility and redirected to AI Review."""
    maybe = _require_active_org()
    if maybe is not None:
        return maybe

    maybe_plan = _require_plan_feature('ai_tagging')
    if maybe_plan is not None:
        return maybe_plan

    flash('AI Summary Test has been merged into AI Review.', 'info')
    return redirect(url_for('main.ai_demo'))


@bp.route('/api/ai/summary-checklist', methods=['POST'])
@login_required
@limiter.limit('8 per minute', key_func=_ai_rate_limit_key)
def ai_summary_checklist_api():
    """Analyze a stored repository document and return a checklist summary."""
    maybe = _require_active_org()
    if maybe is not None:
        return jsonify({'success': False, 'error': 'No active organization'}), 400

    org_id = _active_org_id()
    if not current_user.has_permission('documents.view', org_id=int(org_id)):
        return jsonify({'success': False, 'error': 'Forbidden'}), 403

    stored_doc_id = (request.form.get('stored_doc_id') or '').strip()
    if not stored_doc_id or not stored_doc_id.isdigit():
        return jsonify({'success': False, 'error': 'Choose a repository document first.'}), 400

    from app.services.azure_storage import AzureBlobStorageService

    document = _authorized_org_document_or_404(int(stored_doc_id))
    storage_service = AzureBlobStorageService()
    result = storage_service.download_file(document.blob_name)
    if not result.get('success') or not result.get('data'):
        return jsonify({'success': False, 'error': 'Could not load stored document for AI review.'}), 400

    source_filename = (document.filename or '').strip()
    doc_text, extraction_error = document_analysis_service.extract_text_from_bytes(source_filename, result.get('data') or b'')
    if extraction_error:
        return jsonify({'success': False, 'error': extraction_error}), 400

    question = 'Assess this document against NDIS-style compliance evidence expectations.'
    snippets = _rank_demo_snippets(doc_text, question, top_k=4)
    status, confidence = _derive_demo_status(doc_text, question, snippets, mode='balanced')

    rag_citations = []
    warning_items: list[dict] = []
    _demo_retrieval_mode = 'lexical'
    rag_query_text = _build_demo_rag_query(question, doc_text)
    corpus_path = current_app.config.get('NDIS_RAG_CORPUS_PATH') or 'data/rag/ndis/ndis_chunks.jsonl'
    corpus_abs = os.path.abspath(os.path.join(current_app.root_path, os.pardir, corpus_path))
    try:
        rag_result = rag_query_service.query(corpus_path=corpus_abs, query_text=rag_query_text, requirement_id='', top_k=3)
        _demo_retrieval_mode = getattr(rag_result, 'retrieval_mode', 'lexical')
        rag_citations = [
            {
                'chunk_id': c.chunk_id,
                'source_id': c.source_id,
                'page_number': c.page_number,
                'score': c.score,
                'text': c.text,
            }
            for c in rag_result.citations
        ]
        if _demo_retrieval_mode == 'lexical' and rag_citations:
            warning_items.append({'source': 'rag', 'message': 'Semantic embeddings computing on first run — using keyword matching now. Reload to use hybrid mode.'})
    except FileNotFoundError:
        warning_items.append({'source': 'rag', 'message': 'NDIS corpus is not built yet; showing document-only analysis.'})
    except Exception:
        warning_items.append({'source': 'rag', 'message': 'Could not retrieve NDIS citations; showing document-only analysis.'})

    ai_summary, llm_warning, used_model, llm_provider = _llm_demo_summary(
        status=status,
        question=question,
        snippets=snippets,
        citations=rag_citations,
    )
    if llm_warning:
        warning_items.append({'source': 'llm', 'message': llm_warning})

    ai_summary = _ensure_demo_summary_text(
        ai_summary,
        status=status,
        analysis_mode='balanced',
        snippets=snippets,
        citations=rag_citations,
    )

    provider_label = llm_provider if not llm_warning else 'deterministic'
    if llm_provider == 'azure-openai':
        model_label = used_model or (current_app.config.get('AZURE_OPENAI_CHAT_DEPLOYMENT_MINI') or current_app.config.get('AZURE_OPENAI_CHAT_DEPLOYMENT') or current_app.config.get('AZURE_OPENAI_CHAT_DEPLOYMENT_WRITER') or 'azure-openai')
    else:
        model_label = used_model or current_app.config.get('OPENROUTER_MODEL') or 'mistralai/mistral-7b-instruct:free'

    matched_requirements = document_analysis_service._match_requirements(
        text=doc_text,
        filename=source_filename,
        organization_id=int(org_id),
    )
    checklist = _build_checklist_from_analysis(
        {
            'matched_requirements': matched_requirements,
            'status': status,
            'summary': ai_summary,
            'focus_area': 'General compliance coverage',
        }
    )
    warnings = [w.get('message', '') for w in warning_items if (w.get('message') or '').strip()]

    return jsonify(
        {
            'success': True,
            'status': status,
            'confidence': confidence,
            'summary': ai_summary,
            'checklist': checklist,
            'snippets': snippets,
            'citations': rag_citations,
            'warnings': warnings,
            'warning_items': warning_items,
            'meta': {
                'provider': provider_label,
                'model': model_label,
                'retrieval_mode': _demo_retrieval_mode,
                'document_chars': len(doc_text or ''),
                'filename': source_filename,
                'source': 'repository',
                'stored_doc_id': int(stored_doc_id),
            },
        }
    )


def _openrouter_demo_followup(*, document_name: str, initial_question: str, initial_summary: str, followup_question: str, citations: list[dict]) -> tuple[str | None, str | None]:
    api_key = (current_app.config.get('OPENROUTER_API_KEY') or '').strip()
    configured_model = (current_app.config.get('OPENROUTER_MODEL') or '').strip() or 'mistralai/mistral-7b-instruct:free'
    fallback_models = [configured_model, 'openrouter/auto']
    token_budgets = [500, 260, 160]
    model_candidates = []
    seen = set()
    for model_name in fallback_models:
        key = (model_name or '').strip().lower()
        if not key or key in seen:
            continue
        seen.add(key)
        model_candidates.append(model_name)
    if not api_key:
        return None, 'OPENROUTER_API_KEY is not set. Using deterministic follow-up response.'

    citation_points = '\n'.join([
        f"- {item.get('source_id', 'ndis')} p.{item.get('page_number') or '?'}: {item.get('text', '')}"
        for item in citations[:3]
    ]) or '- No citations available.'

    prompt = (
        'You are an NDIS evidence review assistant. '
        'Answer the follow-up question using the initial review and citations below. '
        'Keep the response practical and concise for a compliance user.\n\n'
        f'Document: {document_name}\n'
        f'Initial review question: {initial_question}\n\n'
        f'Initial review summary:\n{initial_summary}\n\n'
        f'Citations:\n{citation_points}\n\n'
        f'Follow-up question: {followup_question}\n\n'
        'Return plain text only.'
    )

    try:
        import requests

        retriable_statuses = {400, 402, 404, 408, 409, 425, 429, 500, 502, 503, 504}
        last_warning = None
        for model in model_candidates:
            for max_tokens in token_budgets:
                response = requests.post(
                    'https://openrouter.ai/api/v1/chat/completions',
                    headers={
                        'Authorization': f'Bearer {api_key}',
                        'Content-Type': 'application/json',
                        'HTTP-Referer': 'https://cenaris.local',
                        'X-Title': 'Cenaris AI Review Follow-up',
                    },
                    json={
                        'model': model,
                        'messages': [
                            {'role': 'system', 'content': 'Be practical, specific, and avoid legal conclusions.'},
                            {'role': 'user', 'content': prompt},
                        ],
                        'temperature': 0.2,
                        'max_tokens': max_tokens,
                    },
                    timeout=20,
                )
                if response.status_code >= 400:
                    snippet = ((response.text or '').strip()[:140])
                    last_warning = f'OpenRouter follow-up unavailable ({response.status_code}){": " + snippet if snippet else ""}.'
                    if response.status_code in retriable_statuses:
                        continue
                    return None, f'{last_warning} Using deterministic response.'

                payload = response.json() if response.content else {}
                choices = payload.get('choices') or []
                if not choices:
                    last_warning = 'OpenRouter follow-up returned no choices.'
                    continue

                text = (((choices[0] or {}).get('message') or {}).get('content') or '').strip()
                if not text:
                    last_warning = 'OpenRouter follow-up returned empty content.'
                    continue
                return text, None

        if last_warning:
            return None, f'{last_warning} Using deterministic response.'
        return None, 'OpenRouter follow-up unavailable. Using deterministic response.'
    except Exception:
        return None, 'OpenRouter follow-up request failed. Using deterministic response.'


@bp.route('/api/ai/demo/followup', methods=['POST'])
@login_required
@limiter.limit('12 per minute', key_func=_ai_rate_limit_key)
def ai_demo_followup_api():
    """Handle follow-up questions for an analyzed repository document workspace."""
    maybe = _require_active_org()
    if maybe is not None:
        return jsonify({'success': False, 'error': 'No active organization'}), 400

    org_id = _active_org_id()
    if not current_user.has_permission('documents.view', org_id=int(org_id)):
        return jsonify({'success': False, 'error': 'Forbidden'}), 403

    payload = request.get_json(silent=True) or {}
    stored_doc_id = str(payload.get('stored_doc_id') or '').strip()
    followup_question = _limit_text((payload.get('question') or '').strip(), max_chars=700)
    base_question = _limit_text((payload.get('base_question') or '').strip(), max_chars=700)
    base_summary = _limit_text((payload.get('base_summary') or '').strip(), max_chars=3200)
    citations = payload.get('citations') or []
    if not isinstance(citations, list):
        citations = []

    if not stored_doc_id.isdigit():
        return jsonify({'success': False, 'error': 'Invalid repository document selection.'}), 400
    if not followup_question:
        return jsonify({'success': False, 'error': 'Enter a follow-up question first.'}), 400

    document = _authorized_org_document_or_404(int(stored_doc_id))

    answer, warning = _openrouter_demo_followup(
        document_name=(document.filename or 'Document'),
        initial_question=base_question or (document.ai_question or ''),
        initial_summary=base_summary or (document.ai_summary or ''),
        followup_question=followup_question,
        citations=citations[:3],
    )

    if not answer:
        summary_text = (base_summary or document.ai_summary or '').strip()
        fallback = summary_text if summary_text else 'No prior summary is available yet. Run analysis first, then ask a follow-up question.'
        answer = (
            'Using deterministic follow-up guidance:\n\n'
            f'Current review context:\n{fallback}\n\n'
            f'Follow-up request: {followup_question}\n\n'
            'Recommended next step: link this follow-up question to a specific missing evidence item, then re-run analysis after updating the document.'
        )

    return jsonify(
        {
            'success': True,
            'answer': answer,
            'warning': warning,
            'meta': {
                'stored_doc_id': int(stored_doc_id),
                'filename': document.filename,
            },
        }
    )


@bp.route('/api/rag/query', methods=['POST'])
@login_required
@limiter.limit(_rag_rate_limit_value, key_func=_ai_rate_limit_key)
def rag_query_api():
    """Retrieve relevant NDIS corpus citations for a question/requirement."""
    maybe = _require_active_org()
    if maybe is not None:
        return jsonify({'success': False, 'error': 'No active organization'}), 400

    org_id = _active_org_id()
    if not current_user.has_permission('documents.view', org_id=int(org_id)):
        return jsonify({'success': False, 'error': 'Forbidden'}), 403

    payload = request.get_json(silent=True) or {}
    max_query_chars = int(_effective_ai_setting(org_id, 'max_query_chars', current_app.config.get('AI_MAX_QUERY_CHARS') or 1200))
    max_top_k = int(_effective_ai_setting(org_id, 'max_top_k', current_app.config.get('AI_MAX_TOP_K') or 5))
    max_citation_text_chars = int(_effective_ai_setting(org_id, 'max_citation_text_chars', current_app.config.get('AI_MAX_CITATION_TEXT_CHARS') or 600))
    max_answer_chars = int(_effective_ai_setting(org_id, 'max_answer_chars', current_app.config.get('AI_MAX_ANSWER_CHARS') or 2000))
    started = time.perf_counter()

    query_text = _limit_text((payload.get('query') or '').strip(), max_chars=max_query_chars)
    requirement_id = (payload.get('requirement_id') or '').strip()
    top_k = _clamp_int(payload.get('top_k', 3), default=3, minimum=1, maximum=max_top_k)

    if not query_text and not requirement_id:
        return jsonify({'success': False, 'error': 'query or requirement_id is required'}), 400

    corpus_path = current_app.config.get('NDIS_RAG_CORPUS_PATH') or 'data/rag/ndis/ndis_chunks.jsonl'
    corpus_abs = os.path.abspath(os.path.join(current_app.root_path, os.pardir, corpus_path))
    try:
        result = rag_query_service.query(
            corpus_path=corpus_abs,
            query_text=query_text,
            requirement_id=requirement_id,
            top_k=top_k,
        )
    except FileNotFoundError:
        return jsonify({'success': False, 'error': 'RAG corpus is not built yet'}), 503
    except ValueError as e:
        return jsonify({'success': False, 'error': str(e)}), 400
    except Exception:
        current_app.logger.exception('RAG query failed')
        return jsonify({'success': False, 'error': 'Failed to process RAG query'}), 500

    response_payload = {
        'success': True,
        'answer': _limit_text(result.answer, max_chars=max_answer_chars),
        'citations': [
            {
                'chunk_id': c.chunk_id,
                'source_id': c.source_id,
                'page_number': c.page_number,
                'score': c.score,
                'text': _limit_text(c.text, max_chars=max_citation_text_chars),
            }
            for c in result.citations
        ],
        'limits': {
            'top_k': top_k,
            'max_query_chars': max_query_chars,
            'max_citation_text_chars': max_citation_text_chars,
            'max_answer_chars': max_answer_chars,
        },
    }

    latency_ms = int((time.perf_counter() - started) * 1000)
    _log_ai_call(
        'rag_query',
        org_id=int(org_id),
        mode='retrieval',
        provider='local',
        model='lexical-rag',
        usage={'prompt_tokens': 0, 'completion_tokens': 0, 'total_tokens': 0},
        latency_ms=latency_ms,
    )

    return jsonify(
        response_payload
    )


@bp.route('/api/policy/draft', methods=['POST'])
@login_required
@limiter.limit(_policy_rate_limit_value, key_func=_ai_rate_limit_key)
def policy_draft_api():
    """Generate a policy draft using LLM mode (if enabled) with deterministic fallback."""
    maybe = _require_active_org()
    if maybe is not None:
        return jsonify({'success': False, 'error': 'No active organization'}), 400

    org_id = _active_org_id()
    if not current_user.has_permission('documents.view', org_id=int(org_id)):
        return jsonify({'success': False, 'error': 'Forbidden'}), 403

    payload = request.get_json(silent=True) or {}
    max_query_chars = int(_effective_ai_setting(org_id, 'max_query_chars', current_app.config.get('AI_MAX_QUERY_CHARS') or 1200))
    max_top_k = int(_effective_ai_setting(org_id, 'max_top_k', current_app.config.get('AI_MAX_TOP_K') or 5))
    max_citation_text_chars = int(_effective_ai_setting(org_id, 'max_citation_text_chars', current_app.config.get('AI_MAX_CITATION_TEXT_CHARS') or 600))
    max_draft_chars = int(_effective_ai_setting(org_id, 'max_policy_draft_chars', current_app.config.get('AI_MAX_POLICY_DRAFT_CHARS') or 6000))

    policy_type = (payload.get('policy_type') or '').strip()
    query_text = _limit_text((payload.get('query') or '').strip(), max_chars=max_query_chars)
    requirement_id = (payload.get('requirement_id') or '').strip()
    document_id_raw = str(payload.get('document_id') or '').strip()
    document_id = int(document_id_raw) if document_id_raw.isdigit() else None
    requirement_scope = (payload.get('requirement_scope') or 'linked').strip().lower()
    if requirement_scope not in {'linked', 'single'}:
        requirement_scope = 'linked'
    top_k = _clamp_int(payload.get('top_k', 3), default=3, minimum=1, maximum=max_top_k)
    output_mode = (payload.get('output_mode') or 'full_draft').strip().lower()
    audience = _limit_text((payload.get('audience') or '').strip(), max_chars=160)
    policy_tone = _limit_text((payload.get('policy_tone') or '').strip(), max_chars=120)
    strictness = _limit_text((payload.get('strictness') or '').strip(), max_chars=120)
    organization_size = _limit_text((payload.get('organization_size') or '').strip(), max_chars=120)
    context_brief = _limit_text((payload.get('context_brief') or '').strip(), max_chars=max_query_chars)

    if output_mode not in {'template', 'template_plus', 'full_draft'}:
        output_mode = 'full_draft'

    audience = audience or 'Leadership team and frontline workers'
    policy_tone = policy_tone or 'Plain-English'
    strictness = strictness or 'Balanced'
    organization_size = organization_size or 'Small provider'

    if not policy_type:
        return jsonify({'success': False, 'error': 'policy_type is required'}), 400
    if not query_text and not requirement_id and document_id is None:
        query_text = f"Create a comprehensive {policy_type} from scratch with practical procedures, ownership, and review cadence."

    selected_document = None
    linked_requirement_codes: list[str] = []
    linked_requirement_labels: list[str] = []
    if document_id is not None:
        selected_document = db.session.get(Document, int(document_id))
        if (
            not selected_document
            or int(getattr(selected_document, 'organization_id', 0) or 0) != int(org_id)
            or not bool(getattr(selected_document, 'is_active', True))
        ):
            return jsonify({'success': False, 'error': 'document_id is invalid for this organization'}), 400

        if requirement_scope == 'linked':
            linked_requirements = (
                ComplianceRequirement.query
                .join(RequirementEvidenceLink, RequirementEvidenceLink.requirement_id == ComplianceRequirement.id)
                .filter(
                    RequirementEvidenceLink.organization_id == int(org_id),
                    RequirementEvidenceLink.document_id == int(selected_document.id),
                )
                .order_by(ComplianceRequirement.requirement_id.asc(), ComplianceRequirement.id.asc())
                .all()
            )

            seen_codes: set[str] = set()
            for requirement in linked_requirements:
                code = _requirement_display_code(requirement)
                if not code or code in seen_codes:
                    continue
                seen_codes.add(code)
                linked_requirement_codes.append(code)
                linked_requirement_labels.append(f"{code} - {_requirement_plain_title(requirement, max_chars=90)}")

            if linked_requirement_codes and not query_text:
                code_list = ', '.join(linked_requirement_codes[:25])
                query_text = f"Create a comprehensive {policy_type} that covers linked requirements: {code_list}."

            if not linked_requirement_codes:
                # Fall back to document-driven drafting even when no explicit links exist yet.
                if not query_text:
                    query_text = f"Create a comprehensive {policy_type} based on the selected evidence document."
        elif not query_text and not requirement_id:
            query_text = f"Create a comprehensive {policy_type} for the selected requirement."

    if not query_text and requirement_id:
        query_text = f"Create a {policy_type} for requirement {requirement_id}."

    corpus_path = current_app.config.get('NDIS_RAG_CORPUS_PATH') or 'data/rag/ndis/ndis_chunks.jsonl'
    corpus_path = os.path.abspath(os.path.join(current_app.root_path, os.pardir, corpus_path))

    try:
        rag_result = rag_query_service.query(
            corpus_path=corpus_path,
            query_text=query_text,
            requirement_id=requirement_id,
            top_k=top_k,
        )
    except FileNotFoundError:
        return jsonify({'success': False, 'error': 'RAG corpus is not built yet'}), 503
    except ValueError as e:
        return jsonify({'success': False, 'error': str(e)}), 400
    except Exception:
        current_app.logger.exception('Policy draft retrieval failed')
        return jsonify({'success': False, 'error': 'Failed to retrieve citations for policy draft'}), 500

    organization = db.session.get(Organization, int(org_id))
    prompt_path = current_app.config.get('NDIS_POLICY_PROMPT_PATH') or 'app/ai/prompts/ndis_policy_system_prompt.txt'
    prompt_path = os.path.abspath(os.path.join(current_app.root_path, os.pardir, prompt_path))

    citations = [
        {
            'chunk_id': c.chunk_id,
            'source_id': c.source_id,
            'page_number': c.page_number,
            'score': c.score,
            'text': _limit_text(c.text, max_chars=max_citation_text_chars),
        }
        for c in rag_result.citations
    ]

    context_parts = []
    if context_brief:
        context_parts.append(context_brief)
    if selected_document is not None:
        context_parts.append(f"Selected evidence document: {selected_document.filename}")
        extracted = _limit_text(
            ' '.join((getattr(selected_document, 'extracted_text', '') or '').split()),
            max_chars=max(1200, int(max_query_chars)),
        )
        if extracted:
            context_parts.append('Document excerpt:\n' + extracted)
    if linked_requirement_labels:
        context_parts.append(
            'Linked requirements to cover:\n' + '\n'.join(f"- {label}" for label in linked_requirement_labels[:40])
        )
    context_brief = _limit_text('\n\n'.join(context_parts), max_chars=max(2400, int(max_query_chars) * 2))

    draft_result = policy_draft_service.build_draft(
        policy_type=policy_type,
        organization_name=(organization.name if organization else 'Organisation'),
        requirement_id=requirement_id,
        user_goal=query_text,
        citations=citations,
        prompt_path=prompt_path,
        output_mode=output_mode,
        audience=audience,
        policy_tone=policy_tone,
        strictness=strictness,
        organization_size=organization_size,
        context_brief=context_brief,
    )

    draft_text = _limit_text(draft_result.draft_text, max_chars=max_draft_chars)
    disclaimer_text = draft_result.disclaimer
    draft_mode = 'deterministic'
    provider = 'local'
    model = 'deterministic'
    usage: dict[str, int] = {'prompt_tokens': 0, 'completion_tokens': 0, 'total_tokens': 0}
    warnings: list[str] = []

    use_llm = bool(_effective_ai_setting(org_id, 'policy_draft_use_llm', current_app.config.get('POLICY_DRAFT_USE_LLM')))
    ai_env = (current_app.config.get('AI_ENVIRONMENT') or 'development').strip().lower()
    allow_in_dev = bool(current_app.config.get('AI_POLICY_LLM_ALLOW_IN_DEVELOPMENT'))
    if use_llm and ai_env != 'production' and not allow_in_dev:
        warnings.append('LLM mode is disabled outside production by configuration. Used deterministic fallback.')
        use_llm = False

    started = time.perf_counter()
    if use_llm:
        if not azure_openai_policy_service.is_configured(current_app.config):
            warnings.append('LLM mode enabled but Azure OpenAI is not fully configured. Used deterministic fallback.')
        else:
            try:
                llm_result = azure_openai_policy_service.generate_policy_draft(
                    config=current_app.config,
                    policy_type=policy_type,
                    organization_name=(organization.name if organization else 'Organisation'),
                    requirement_id=requirement_id,
                    user_goal=query_text,
                    citations=citations,
                    prompt_path=prompt_path,
                    output_mode=output_mode,
                    audience=audience,
                    policy_tone=policy_tone,
                    strictness=strictness,
                    organization_size=organization_size,
                    context_brief=context_brief,
                )
                draft_text = _limit_text(llm_result.draft_text, max_chars=max_draft_chars)
                disclaimer_text = llm_result.disclaimer
                usage = llm_result.usage
                draft_mode = 'llm'
                provider = 'azure-openai'
                model = llm_result.deployment
            except Exception:
                current_app.logger.exception('Policy draft LLM generation failed; using deterministic fallback')
                warnings.append('LLM generation failed. Used deterministic fallback.')

    latency_ms = int((time.perf_counter() - started) * 1000)
    _log_ai_call(
        'policy_draft',
        org_id=int(org_id),
        mode=draft_mode,
        provider=provider,
        model=model,
        usage=usage,
        latency_ms=latency_ms,
    )

    return jsonify(
        {
            'success': True,
            'draft_text': draft_text,
            'disclaimer': disclaimer_text,
            'draft_mode': draft_mode,
            'llm': {
                'provider': provider,
                'model': model,
                'usage': usage,
                'latency_ms': latency_ms,
            },
            'warnings': warnings,
            'citations': citations,
            'inputs': {
                'output_mode': output_mode,
                'audience': audience,
                'policy_tone': policy_tone,
                'strictness': strictness,
                'organization_size': organization_size,
                'source_mode': 'document' if selected_document else 'scratch',
                'document_id': int(selected_document.id) if selected_document else None,
                'requirement_scope': requirement_scope,
                'linked_requirements_count': int(len(linked_requirement_codes)),
                'context_brief_present': bool(context_brief),
            },
            'limits': {
                'top_k': top_k,
                'max_query_chars': max_query_chars,
                'max_citation_text_chars': max_citation_text_chars,
                'max_policy_draft_chars': max_draft_chars,
            },
        }
    )


@bp.route('/api/policy/export-docx', methods=['POST'])
@login_required
@limiter.limit('20 per minute', key_func=_ai_rate_limit_key)
def policy_export_docx_api():
    """Export a generated policy draft to DOCX with basic structure formatting."""
    maybe = _require_active_org()
    if maybe is not None:
        return jsonify({'success': False, 'error': 'No active organization'}), 400

    org_id = _active_org_id()
    if not current_user.has_permission('documents.view', org_id=int(org_id)):
        return jsonify({'success': False, 'error': 'Forbidden'}), 403

    payload = request.get_json(silent=True) or {}
    policy_type = _limit_text((payload.get('policy_type') or 'Policy Draft').strip(), max_chars=180) or 'Policy Draft'
    draft_text = (payload.get('draft_text') or '').strip()
    if not draft_text:
        return jsonify({'success': False, 'error': 'draft_text is required'}), 400

    lines = [line.rstrip() for line in draft_text.splitlines()]

    try:
        from docx import Document as WordDocument
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        from docx.shared import Pt

        doc = WordDocument()

        title = doc.add_paragraph()
        title_run = title.add_run(policy_type)
        title_run.bold = True
        title_run.font.size = Pt(18)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        subtitle = doc.add_paragraph()
        subtitle_run = subtitle.add_run(f"Generated on {datetime.now(timezone.utc).strftime('%d %b %Y')}")
        subtitle_run.italic = True
        subtitle_run.font.size = Pt(10)
        subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        doc.add_paragraph('')

        for raw_line in lines:
            line = (raw_line or '').strip()
            if not line:
                continue

            if line.startswith('# '):
                doc.add_heading(line[2:].strip(), level=1)
                continue
            if line.startswith('## '):
                doc.add_heading(line[3:].strip(), level=2)
                continue
            if line.startswith('### '):
                doc.add_heading(line[4:].strip(), level=3)
                continue

            if line.endswith(':') and len(line) <= 120:
                doc.add_heading(line.rstrip(':'), level=2)
                continue

            if line.startswith('- ') or line.startswith('* '):
                doc.add_paragraph(line[2:].strip(), style='List Bullet')
                continue

            if line[:2].isdigit() and line[1] == '.':
                doc.add_paragraph(line[2:].strip(), style='List Number')
                continue

            paragraph = doc.add_paragraph(line)
            paragraph.paragraph_format.space_after = Pt(8)

        output = io.BytesIO()
        doc.save(output)
        output.seek(0)

        safe_name = ''.join(ch if ch.isalnum() else '-' for ch in policy_type.lower()).strip('-') or 'policy-draft'
        filename = f"{safe_name}.docx"
        return send_file(
            output,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        )
    except Exception:
        current_app.logger.exception('Policy DOCX export failed')
        return jsonify({'success': False, 'error': 'Failed to export policy document'}), 500

@bp.route('/reports')
@login_required
def reports():
    """Reports route."""
    reports = [
        {
            'id': 1,
            'name': 'ISO 27001 Compliance Report',
            'description': 'Comprehensive assessment of ISO 27001 compliance status',
            'type': 'Compliance Assessment',
            'generated_date': '2024-10-13',
            'status': 'Complete',
            'download_url': '#'
        }
    ]
    
    return render_template('main/reports.html', 
                         title='Compliance Reports',
                         reports=reports)

@bp.route('/settings')
@login_required
def settings():
    """Settings route."""
    return render_template('main/settings.html', title='Settings')

@bp.route('/help')
@login_required
def help():
    """Help route."""
    return render_template('main/help.html', title='Help & Documentation')

@bp.route('/profile', methods=['GET', 'POST'])
@login_required
def profile():
    """User profile route."""
    from app.main.forms import UserProfileForm
    from app.models import Department

    form = UserProfileForm(obj=current_user)

    if form.validate_on_submit():
        current_user.first_name = (form.first_name.data or '').strip() or None
        current_user.last_name = (form.last_name.data or '').strip() or None

        # Keep full_name in sync if it was empty.
        if not (current_user.full_name or '').strip():
            parts = [p for p in [(current_user.first_name or ''), (current_user.last_name or '')] if p.strip()]
            current_user.full_name = ' '.join([p.strip() for p in parts]) or None

        try:
            db.session.commit()
            flash('Profile updated.', 'success')
            return redirect(url_for('main.profile'))
        except Exception as e:
            db.session.rollback()
            flash('Profile update failed. Please try again.', 'error')
            current_app.logger.error(f"Profile update failed for user {current_user.id}: {e}")

    # Get current membership and departments for self-assignment
    current_membership = None
    departments = []
    org_id = getattr(current_user, 'organization_id', None)
    if org_id:
        current_membership = (
            OrganizationMembership.query
            .filter_by(user_id=int(current_user.id), organization_id=int(org_id), is_active=True)
            .first()
        )
        departments = (
            Department.query
            .filter_by(organization_id=int(org_id))
            .order_by(Department.name.asc())
            .all()
        )

    # Get current membership and departments for self-assignment
    current_membership = None
    departments = []
    org_id = getattr(current_user, 'organization_id', None)
    if org_id:
        current_membership = (
            OrganizationMembership.query
            .filter_by(user_id=int(current_user.id), organization_id=int(org_id), is_active=True)
            .first()
        )
        departments = (
            Department.query
            .filter_by(organization_id=int(org_id))
            .order_by(Department.name.asc())
            .all()
        )

    return render_template(
        'main/profile.html',
        title='My Profile',
        form=form,
        current_membership=current_membership,
        departments=departments,
        has_local_password=bool((getattr(current_user, 'password_hash', '') or '').strip()),
    )


@bp.route('/profile/department', methods=['POST'])
@login_required
def profile_update_department():
    """Allow user to assign themselves to a department."""
    from app.models import Department

    org_id = getattr(current_user, 'organization_id', None)
    if not org_id:
        flash('No organisation associated with your account.', 'error')
        return redirect(url_for('main.profile'))

    membership = (
        OrganizationMembership.query
        .filter_by(user_id=int(current_user.id), organization_id=int(org_id), is_active=True)
        .first()
    )
    if not membership:
        flash('Membership not found.', 'error')
        return redirect(url_for('main.profile'))

    dept_id_str = (request.form.get('department_id') or '').strip()
    
    if not dept_id_str:
        # Unassign department
        membership.department_id = None
        try:
            db.session.commit()
            flash('Department unassigned.', 'success')
        except Exception as e:
            db.session.rollback()
            flash('Failed to update department.', 'error')
            current_app.logger.error(f"Failed to unassign department for user {current_user.id}: {e}")
        return redirect(url_for('main.profile'))

    try:
        dept_id = int(dept_id_str)
    except ValueError:
        flash('Invalid department selected.', 'error')
        return redirect(url_for('main.profile'))

    # Verify department belongs to the same organization
    department = db.session.get(Department, dept_id)
    if not department or department.organization_id != int(org_id):
        flash('Invalid department selected.', 'error')
        return redirect(url_for('main.profile'))

    membership.department_id = dept_id
    try:
        db.session.commit()
        flash(f'Department updated to "{department.name}".', 'success')
    except Exception as e:
        db.session.rollback()
        flash('Failed to update department.', 'error')
        current_app.logger.error(f"Failed to update department for user {current_user.id}: {e}")

    return redirect(url_for('main.profile'))


@bp.route('/profile/delete-account', methods=['POST'])
@login_required
def profile_delete_account():
    """Deactivate the current user account and end all sessions."""
    confirm_text = (request.form.get('confirm_text') or '').strip().upper()
    password = (request.form.get('password') or '').strip()

    if confirm_text != 'DELETE':
        flash('Type DELETE to confirm account deletion.', 'error')
        return redirect(url_for('main.profile'))

    user = db.session.get(User, int(current_user.id))
    if not user:
        flash('Account not found.', 'error')
        return redirect(url_for('main.profile'))

    has_local_password = bool((user.password_hash or '').strip())
    if has_local_password and not password:
        flash('Password is required for password-based accounts.', 'error')
        return redirect(url_for('main.profile'))

    if has_local_password and not user.check_password(password):
        flash('Password is incorrect.', 'error')
        return redirect(url_for('main.profile'))

    active_memberships = (
        OrganizationMembership.query
        .filter_by(user_id=int(user.id), is_active=True)
        .all()
    )

    # Safety: do not allow deleting the only active admin in any organisation.
    for membership in active_memberships:
        if not _membership_has_permission(membership, 'users.manage'):
            continue

        active_org_memberships = (
            OrganizationMembership.query
            .filter_by(organization_id=int(membership.organization_id), is_active=True)
            .all()
        )
        active_admins = sum(1 for item in active_org_memberships if _membership_has_permission(item, 'users.manage'))
        if active_admins <= 1:
            flash('You are the only active admin in at least one organisation. Promote another admin before deleting this account.', 'error')
            return redirect(url_for('main.profile'))

    try:
        for membership in active_memberships:
            membership.is_active = False

        # Free up the original email so users can re-register with the same address later.
        # Keep a traceable but non-routable tombstone value.
        tombstone_email = f"deleted+{int(user.id)}+{int(time.time())}@deleted.local"
        user.email = tombstone_email
        user.email_verified = False
        user.password_hash = None
        user.first_name = None
        user.last_name = None
        user.full_name = None

        user.is_active = False
        user.organization_id = None
        user.session_version = int(getattr(user, 'session_version', 1) or 1) + 1

        db.session.commit()
    except Exception:
        db.session.rollback()
        current_app.logger.exception('Failed to delete account for user_id=%s', getattr(current_user, 'id', None))
        flash('Could not delete account right now. Please try again.', 'error')
        return redirect(url_for('main.profile'))

    logout_user()
    session.clear()
    flash('Your account has been deleted.', 'success')
    return redirect(url_for('main.index'))


@bp.route('/profile/avatar')
@login_required
def profile_avatar():
    """Serve the current user's avatar image."""
    from flask import abort, send_file
    from app.services.azure_storage import AzureBlobStorageService
    import io

    if not getattr(current_user, 'avatar_blob_name', None):
        abort(404)

    storage_service = AzureBlobStorageService()
    result = storage_service.download_file(current_user.avatar_blob_name)
    if not result.get('success'):
        abort(404)

    blob_data = result.get('data')
    if not blob_data:
        abort(404)

    file_stream = io.BytesIO(blob_data)
    file_stream.seek(0)
    return send_file(
        file_stream,
        mimetype=getattr(current_user, 'avatar_content_type', None) or 'application/octet-stream',
        as_attachment=False,
        download_name='avatar'
    )

@bp.route('/notifications')
@login_required
def notifications():
    """Admin-only notifications center for important organisation events."""
    maybe = _require_org_admin()
    if maybe is not None:
        return maybe

    org_id = _active_org_id()
    if not org_id:
        abort(400)

    unread_only = (request.args.get('status') or '').strip().lower() == 'unread'
    notifications = notification_service.list_admin_notifications(
        organization_id=int(org_id),
        unread_only=unread_only,
        limit=150,
    )
    unread_count = notification_service.unread_count(organization_id=int(org_id))

    return render_template(
        'main/notifications.html',
        title='Notifications',
        notifications=notifications,
        unread_only=unread_only,
        unread_count=unread_count,
    )


@bp.route('/notifications/<int:notification_id>/read', methods=['POST'])
@login_required
def mark_notification_read(notification_id: int):
    maybe = _require_org_admin()
    if maybe is not None:
        return maybe

    org_id = _active_org_id()
    if not org_id:
        abort(400)

    ok = notification_service.mark_read(
        notification_id=int(notification_id),
        user_id=int(current_user.id),
        organization_id=int(org_id),
    )
    invalidate_org_switcher_context_cache(int(current_user.id), int(org_id))
    if not ok:
        flash('Notification not found.', 'error')
    return redirect(url_for('main.notifications'))


@bp.route('/notifications/mark-all-read', methods=['POST'])
@login_required
def mark_all_notifications_read():
    maybe = _require_org_admin()
    if maybe is not None:
        return maybe

    org_id = _active_org_id()
    if not org_id:
        abort(400)

    marked = notification_service.mark_all_read(organization_id=int(org_id), user_id=int(current_user.id))
    invalidate_org_switcher_context_cache(int(current_user.id), int(org_id))
    if marked > 0:
        flash(f'Marked {marked} notification(s) as read.', 'success')
    return redirect(url_for('main.notifications'))

@bp.route('/audit-export')
@login_required
def audit_export():
    """Legacy audit export route; redirects to AI Review workspace."""
    maybe = _require_active_org()
    if maybe is not None:
        return maybe

    org_id = _active_org_id()
    if not current_user.has_permission('documents.view', org_id=int(org_id)):
        abort(403)

    flash('Audit Export has moved into AI Review workspaces.', 'info')
    return redirect(url_for('main.ai_demo'))


@bp.route('/analytics')
@login_required
def analytics_dashboard():
    """Analytics dashboard for compliance trends and reporting."""
    maybe = _require_org_permission('documents.view')
    if maybe is not None:
        return maybe

    maybe_plan = _require_plan_feature('multi_site_reporting')
    if maybe_plan is not None:
        return maybe_plan

    org_id = _active_org_id()
    organization = db.session.get(Organization, int(org_id))
    if not organization:
        abort(404)

    payload = analytics_service.build_dashboard_payload(organization_id=int(org_id))
    return render_template(
        'main/analytics_dashboard.html',
        title='Analytics Dashboard',
        summary=payload.get('summary') or {},
        framework_analytics=payload.get('framework_analytics') or [],
        analytics_payload=payload,
    )


@bp.route('/analytics/export.xlsx')
@login_required
def analytics_export_xlsx():
    """Export analytics data to Excel."""
    maybe = _require_org_permission('audits.export')
    if maybe is not None:
        return maybe

    maybe_plan = _require_plan_feature('multi_site_reporting')
    if maybe_plan is not None:
        return maybe_plan

    org_id = _active_org_id()
    payload = analytics_service.build_dashboard_payload(organization_id=int(org_id))
    workbook = analytics_service.build_excel(payload)
    filename = f"analytics_dashboard_{datetime.now().strftime('%Y%m%d')}.xlsx"

    return send_file(
        workbook,
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        as_attachment=True,
        download_name=filename,
    )


@bp.route('/analytics/export.pdf')
@login_required
def analytics_export_pdf():
    """Export analytics summary to PDF."""
    maybe = _require_org_permission('audits.export')
    if maybe is not None:
        return maybe

    maybe_plan = _require_plan_feature('multi_site_reporting')
    if maybe_plan is not None:
        return maybe_plan

    org_id = _active_org_id()
    organization = db.session.get(Organization, int(org_id))
    if not organization:
        abort(404)

    payload = analytics_service.build_dashboard_payload(organization_id=int(org_id))
    pdf_buffer = analytics_service.build_pdf(payload, organization_name=(organization.name or 'Organisation'))
    filename = f"analytics_dashboard_{datetime.now().strftime('%Y%m%d')}.pdf"

    return send_file(
        pdf_buffer,
        mimetype='application/pdf',
        as_attachment=True,
        download_name=filename,
    )

# User roles route removed - functionality moved to Org Admin Dashboard

@bp.route('/debug-adls')
@login_required
def debug_adls():
    """Debug ADLS connection and data."""
    import os
    from datetime import datetime
    
    debug_info = {
        'timestamp': datetime.now().isoformat(),
        'connection_string_set': bool(os.getenv('AZURE_STORAGE_CONNECTION_STRING')),
        'user_id': current_user.id,
        'service_client_initialized': azure_data_service.service_client is not None,
    }
    
    # Try to get files
    try:
        files = azure_data_service.get_compliance_files(user_id=current_user.id)
        debug_info['files_found'] = len(files)
        debug_info['files'] = files
    except Exception as e:
        debug_info['files_error'] = str(e)
    
    # Try to get summary
    try:
        summary = azure_data_service.get_dashboard_summary(user_id=current_user.id)
        debug_info['summary'] = summary
        
        # Show raw data from files
        if summary.get('file_summaries'):
            debug_info['raw_frameworks'] = []
            for fs in summary['file_summaries']:
                debug_info['raw_frameworks'].extend(fs.get('frameworks', []))
    except Exception as e:
        debug_info['summary_error'] = str(e)
    
    return jsonify(debug_info)

@bp.route('/reports/generate/<report_type>')
@login_required
def generate_report(report_type):
    """Generate and download compliance reports."""
    maybe = _require_org_permission('audits.export')
    if maybe is not None:
        return maybe

    from flask import send_file
    from app.services.report_generator import report_generator
    from datetime import datetime

    org_id = _active_org_id()
    organization = db.session.get(Organization, int(org_id))
    if not organization:
        abort(404)

    if not organization.billing_complete():
        flash('Add billing details to generate reports.', 'warning')
        return redirect(url_for('onboarding.billing'))
    
    # Get organization data (you can customize this)
    org_data = {
        'name': organization.name,
        'abn': organization.abn or '',
        'address': organization.address or '',
        'contact_name': current_user.display_name(),
        'email': organization.contact_email or current_user.email,
        'framework': organization.industry or '',
        'audit_type': 'Initial'
    }
    
    # Get gap analysis data
    summary = azure_data_service.get_dashboard_summary(user_id=current_user.id, organization_id=org_id)
    gap_data = []
    
    if summary.get('file_summaries'):
        for file_summary in summary['file_summaries']:
            frameworks_data = file_summary.get('frameworks', [])
            for framework_data in frameworks_data:
                status = framework_data.get('status', '').strip()
                if status.lower() == 'complete':
                    display_status = 'Complete'
                elif status.lower() == 'needs review':
                    display_status = 'Needs Review'
                elif status.lower() == 'missing':
                    display_status = 'Missing'
                else:
                    display_status = status
                
                gap_data.append({
                    'requirement_name': framework_data['name'],
                    'status': display_status,
                    'completion_percentage': round(framework_data['score'], 1),  # Score is already a percentage
                    'supporting_evidence': file_summary.get('file_name', 'compliance_summary.csv'),
                    'last_updated': file_summary.get('last_updated')
                })
    
    # Calculate summary stats
    total = len(gap_data)
    met = len([g for g in gap_data if g['status'] == 'Complete'])
    pending = len([g for g in gap_data if g['status'] == 'Needs Review'])
    not_met = len([g for g in gap_data if g['status'] == 'Missing'])
    
    if gap_data:
        avg_percentage = sum([g['completion_percentage'] for g in gap_data]) / len(gap_data)
    else:
        avg_percentage = 0
    
    summary_stats = {
        'total': total,
        'met': met,
        'pending': pending,
        'not_met': not_met,
        'compliance_percentage': int(avg_percentage)
    }
    
    # Get documents for audit pack
    documents = (
        Document.query.filter_by(organization_id=int(org_id), is_active=True)
        .order_by(Document.uploaded_at.desc())
        .all()
    )
    
    # Generate appropriate report
    try:
        if report_type == 'gap-analysis':
            pdf_buffer = report_generator.generate_gap_analysis_report(org_data, gap_data, summary_stats)
            filename = f'Gap_Analysis_Report_{datetime.now().strftime("%Y%m%d")}.pdf'
        elif report_type == 'accreditation-plan':
            pdf_buffer = report_generator.generate_accreditation_plan(org_data, gap_data, summary_stats)
            filename = f'Accreditation_Plan_{datetime.now().strftime("%Y%m%d")}.pdf'
        elif report_type == 'audit-pack':
            pdf_buffer = report_generator.generate_audit_pack(org_data, gap_data, summary_stats, documents)
            filename = f'Audit_Pack_Export_{datetime.now().strftime("%Y%m%d")}.pdf'
        else:
            return "Invalid report type", 400
        
        return send_file(
            pdf_buffer,
            mimetype='application/pdf',
            as_attachment=True,
            download_name=filename
        )
    except Exception as e:
        current_app.logger.exception(f'Error generating report {report_type}: {e}')
        return f"Error generating report: {str(e)}", 500


@bp.route('/system-logs')
@login_required
def system_logs():
    """System logs viewing interface (organisation admin only)."""
    maybe = _require_active_org()
    if maybe is not None:
        return maybe

    org_id = _active_org_id()
    
    # Check if user is admin or has org.settings permission
    from app.models import OrganizationMembership
    org_member = OrganizationMembership.query.filter_by(
        user_id=current_user.id,
        organization_id=org_id
    ).first()
    
    # Allow if user is org admin or has explicit org.settings permission
    if not (org_member and (org_member.role == 'Admin' or current_user.has_permission('org.settings', org_id=int(org_id)))):
        flash('You must be an organisation administrator to view system logs.', 'error')
        return redirect(url_for('main.dashboard'))
    
    # Get filter parameters
    log_type = request.args.get('log_type', 'all')
    event_type = request.args.get('event_type', '')
    time_range = request.args.get('time_range', '24h')
    user_id_filter = request.args.get('user_id', '')
    
    # Parse time range
    from datetime import datetime, timezone, timedelta
    now = datetime.now(timezone.utc)

    # Display timestamps in the user's preferred timezone.
    # Source timestamps are stored in UTC in the database.
    display_tz_name = (getattr(current_user, 'time_zone', None) or 'Australia/Sydney').strip() or 'Australia/Sydney'
    display_tz = timezone.utc
    display_tz_label = 'UTC'
    try:
        from zoneinfo import ZoneInfo  # py>=3.9

        display_tz = ZoneInfo(display_tz_name)
        # Use current time to derive a readable label (handles DST).
        display_tz_label = (now.astimezone(display_tz).tzname() or display_tz_name)
    except Exception:
        display_tz = timezone.utc
        display_tz_label = 'UTC'
    time_ranges = {
        '1h': timedelta(hours=1),
        '24h': timedelta(hours=24),
        '7d': timedelta(days=7),
        '30d': timedelta(days=30),
    }
    time_delta = time_ranges.get(time_range, timedelta(hours=24))
    start_time = now - time_delta
    
    logs = []

    # NOTE: We currently back System Logs with persisted DB events.
    # `LoginEvent` captures login success/failure, and we scope to the active organisation
    # by selecting events for users who are members of that organisation.
    if log_type in ['all', 'security']:
        from app.models import LoginEvent
        from sqlalchemy import desc

        member_ids = [m.user_id for m in OrganizationMembership.query.filter_by(organization_id=org_id).all()]

        if member_ids:
            q = LoginEvent.query.filter(LoginEvent.user_id.in_(member_ids))
            if start_time is not None:
                q = q.filter(LoginEvent.created_at >= start_time)

            if (user_id_filter or '').strip():
                try:
                    q = q.filter(LoginEvent.user_id == int(user_id_filter))
                except Exception:
                    pass

            if (event_type or '').strip() == 'LOGIN_SUCCESS':
                q = q.filter(LoginEvent.success.is_(True))
            elif (event_type or '').strip() == 'LOGIN_FAILURE':
                q = q.filter(LoginEvent.success.is_(False))

            events = q.order_by(desc(LoginEvent.created_at)).limit(500).all()
            for evt in events:
                created_at = evt.created_at
                if created_at and getattr(created_at, 'tzinfo', None) is None:
                    created_at = created_at.replace(tzinfo=timezone.utc)

                created_at_local = None
                try:
                    created_at_local = created_at.astimezone(display_tz) if created_at else None
                except Exception:
                    created_at_local = created_at

                is_success = bool(evt.success)
                derived_event_type = 'LOGIN_SUCCESS' if is_success else 'LOGIN_FAILURE'
                derived_description = 'User logged in successfully' if is_success else 'Failed login attempt'

                user_name = None
                user_email = None
                if getattr(evt, 'user', None) is not None:
                    try:
                        user_name = evt.user.display_name()
                    except Exception:
                        user_name = None
                    user_email = getattr(evt.user, 'email', None)

                logs.append({
                    'timestamp': created_at_local.strftime('%Y-%m-%d %H:%M:%S %Z') if created_at_local else None,
                    'log_type': 'security',
                    'event_type': derived_event_type,
                    'event_description': derived_description,
                    'user_id': evt.user_id,
                    'user_name': user_name,
                    'user_email': user_email or evt.email,
                    'organization_id': org_id,
                    'ip_address': evt.ip_address,
                    'details': {
                        'email': evt.email,
                        'provider': evt.provider,
                        'success': is_success,
                        'reason': evt.reason,
                        'user_agent': evt.user_agent,
                    }
                })

    if log_type in ['all', 'ai']:
        q = AIUsageEvent.query.filter_by(organization_id=int(org_id))
        if start_time is not None:
            q = q.filter(AIUsageEvent.created_at >= start_time)
        if (user_id_filter or '').strip():
            try:
                q = q.filter(AIUsageEvent.user_id == int(user_id_filter))
            except Exception:
                pass

        if (event_type or '').strip():
            q = q.filter(AIUsageEvent.event == (event_type or '').strip())

        ai_events_rows = q.order_by(AIUsageEvent.created_at.desc()).limit(500).all()
        for evt in ai_events_rows:
            created_at = evt.created_at
            if created_at and getattr(created_at, 'tzinfo', None) is None:
                created_at = created_at.replace(tzinfo=timezone.utc)

            created_at_local = None
            try:
                created_at_local = created_at.astimezone(display_tz) if created_at else None
            except Exception:
                created_at_local = created_at

            logs.append({
                'timestamp': created_at_local.strftime('%Y-%m-%d %H:%M:%S %Z') if created_at_local else None,
                'log_type': 'ai',
                'event_type': evt.event,
                'event_description': f"{evt.mode} via {evt.provider}",
                'user_id': evt.user_id,
                'user_name': None,
                'user_email': None,
                'organization_id': org_id,
                'ip_address': None,
                'details': {
                    'mode': evt.mode,
                    'provider': evt.provider,
                    'model': evt.model,
                    'prompt_tokens': int(evt.prompt_tokens or 0),
                    'completion_tokens': int(evt.completion_tokens or 0),
                    'total_tokens': int(evt.total_tokens or 0),
                    'latency_ms': int(evt.latency_ms or 0),
                }
            })

    logs.sort(key=lambda item: item.get('timestamp') or '', reverse=True)
    
    # Statistics
    total_logs = len(logs)
    security_events = len([l for l in logs if l.get('log_type') == 'security'])
    ai_events = len([l for l in logs if l.get('log_type') == 'ai'])
    error_count = len([l for l in logs if l.get('log_type') == 'error'])
    failed_logins = len([l for l in logs if l.get('event_type') == 'LOGIN_FAILURE'])
    
    appinsights_enabled = current_app.config.get('APPINSIGHTS_ENABLED', False)
    
    return render_template('main/system_logs.html',
                         title='System Logs',
                         logs=logs,
                         log_type=log_type,
                         event_type=event_type,
                         time_range=time_range,
                         user_id=user_id_filter,
                         total_logs=total_logs,
                         security_events=security_events,
                         ai_events=ai_events,
                         error_count=error_count,
                         failed_logins=failed_logins,
                         display_tz_label=display_tz_label,
                         appinsights_enabled=appinsights_enabled)
