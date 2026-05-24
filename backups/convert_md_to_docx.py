import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def convert_md_to_docx(md_path, docx_path):
    if not os.path.exists(md_path):
        print(f"Error: {md_path} not found.")
        return

    # Read markdown content
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()
    
    # Configure page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Styles setup
    styles = doc.styles
    normal_style = styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Arial'
    normal_font.size = Pt(11)
    normal_font.color.rgb = RGBColor(0x22, 0x22, 0x22)

    in_code_block = False
    code_text = []
    in_table = False
    table_rows = []

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Handle Code Blocks
        if stripped.startswith('```'):
            if in_code_block:
                # End of code block
                in_code_block = False
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.4)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                run = p.add_run('\n'.join(code_text))
                run.font.name = 'Consolas'
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                code_text = []
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_text.append(line.rstrip('\r\n'))
            i += 1
            continue

        # Handle Tables
        if stripped.startswith('|'):
            in_table = True
            table_rows.append(stripped)
            i += 1
            continue
        elif in_table:
            # End of table
            in_table = False
            # Parse and write table
            write_word_table(doc, table_rows)
            table_rows = []
            # Fall through to process current line

        if not stripped:
            i += 1
            continue

        # Handle Headings
        if stripped.startswith('# '):
            heading_text = stripped[2:].strip()
            p = doc.add_heading(heading_text, level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            run.font.name = 'Arial'
            run.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
        elif stripped.startswith('## '):
            heading_text = stripped[3:].strip()
            p = doc.add_heading(heading_text, level=1)
            run = p.runs[0]
            run.font.name = 'Arial'
            run.font.color.rgb = RGBColor(0x00, 0x44, 0x88)
        elif stripped.startswith('### '):
            heading_text = stripped[4:].strip()
            p = doc.add_heading(heading_text, level=2)
            run = p.runs[0]
            run.font.name = 'Arial'
            run.font.color.rgb = RGBColor(0x00, 0x66, 0x99)
        elif stripped.startswith('#### '):
            heading_text = stripped[5:].strip()
            p = doc.add_heading(heading_text, level=3)
            run = p.runs[0]
            run.font.name = 'Arial'
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        # Handle Bullet Points
        elif stripped.startswith('* ') or stripped.startswith('- '):
            text = stripped[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, text)
        # Handle Blockquotes
        elif stripped.startswith('>'):
            text = stripped[1:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            run = p.add_run(text)
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        # Regular Paragraph
        else:
            p = doc.add_paragraph()
            add_formatted_text(p, stripped)

        i += 1

    # Save Document
    doc.save(docx_path)
    print(f"Success: Document saved to {docx_path}")

def write_word_table(doc, raw_rows):
    # Filter out separator row (contains dashes and pipes)
    rows = []
    for row in raw_rows:
        cleaned = [cell.strip() for cell in row.split('|')[1:-1]]
        if all(re.match(r'^:?-+:?$', cell) for cell in cleaned if cell):
            continue
        rows.append(cleaned)

    if not rows:
        return

    col_count = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = 'Light Shading Accent 1'

    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            # Safe cell assignment
            if c_idx < len(table.rows[r_idx].cells):
                cell = table.rows[r_idx].cells[c_idx]
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(val)
                run.font.name = 'Arial'
                run.font.size = Pt(9.5)
                # Bold headers
                if r_idx == 0:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

def add_formatted_text(paragraph, text):
    # Match inline code `code`, bold **bold**, italic *italic*, LaTeX inline $$math$$ or $math$
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`|\$\$.*?\$\$|\$.*?\$)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.font.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.font.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x99, 0x00, 0x33)
        elif part.startswith('$$') and part.endswith('$$'):
            run = paragraph.add_run(part[2:-2])
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x11, 0x44, 0x77)
        elif part.startswith('$') and part.endswith('$'):
            run = paragraph.add_run(part[1:-1])
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x11, 0x44, 0x77)
        else:
            paragraph.add_run(part)

if __name__ == '__main__':
    md_file = r'c:\Users\DELL\Desktop\cenaris\docs\AI_REVIEW_ENGINE_TECHNICAL_SPEC.md'
    docx_file = r'c:\Users\DELL\Desktop\cenaris\docs\AI_REVIEW_ENGINE_TECHNICAL_SPEC.docx'
    convert_md_to_docx(md_file, docx_file)
